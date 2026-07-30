from pathlib import Path

from docx import Document

from paper_formatter.compiler import LatexCompiler
from paper_formatter.benchmark import TemplateBenchmarkRunner
from paper_formatter.models import TableBlock
from paper_formatter.pipeline import DocxToLatexPipeline
from paper_formatter.renderers.latex_renderer import LatexRenderer
from paper_formatter.utils.text import latex_escape


def test_pipeline_creates_ir_tex_and_zip(tmp_path: Path) -> None:
    source = tmp_path / "article.docx"
    document = Document()
    document.add_heading("Название", level=0)
    document.add_heading("Введение", level=1)
    document.add_paragraph("Текст с символами 10% и A_B.")
    document.save(source)

    output = tmp_path / "run"
    result = DocxToLatexPipeline().run(source, output, compile_pdf=False)

    assert (output / "parsed" / "article_ir.json").exists()
    assert result.main_tex.exists()
    assert result.latex_zip.exists()
    tex = (result.main_tex.parent / "body.tex").read_text(encoding="utf-8")
    assert r"10\%" in tex
    assert r"A\_B" in tex
    main = result.main_tex.read_text(encoding="utf-8")
    assert r"\input{metadata.tex}" in main
    assert r"\input{body.tex}" in main


def test_math_unicode_and_missing_glyph_warnings_are_handled() -> None:
    normalized = LatexRenderer._math_latex("A∧B≠C")

    assert r"\land" in normalized
    assert r"\ne" in normalized
    assert any(
        "отсутствуют глифы" in warning
        for warning in LatexCompiler._log_warnings("Missing character: U+2227")
    )


def test_latex_tables_use_weighted_academic_columns() -> None:
    table = TableBlock(
        id="t1",
        header_rows=1,
        rows=[
            ["Тип вопроса", "N", "Оценка"],
            ["переформулировка", "10", "3,10"],
        ],
    )

    rendered = "\n".join(LatexRenderer()._render_table(table))

    assert r"\toprule" in rendered
    assert r"\midrule" in rendered
    assert r"\bottomrule" in rendered
    assert r"{\RaggedRight\arraybackslash\textbf{Тип вопроса}}" in rendered
    assert r"\hsize=" in rendered
    assert r"\begin{tabularx}{\linewidth}{|" not in rendered


def test_long_identifiers_have_safe_latex_breakpoints() -> None:
    escaped = latex_escape("chat_template_kwargs.enable_thinking=false")

    assert r"\_\allowbreak{}" in escaped
    assert r".\allowbreak{}" in escaped
    assert r"=\allowbreak{}" in escaped


def test_latex_title_uses_separate_pdf_title_zone(tmp_path: Path) -> None:
    source = tmp_path / "article.docx"
    document = Document()
    document.add_heading("Широкое название статьи", level=0)
    document.add_paragraph("Основной текст.")
    document.save(source)
    result = DocxToLatexPipeline().run(
        source,
        tmp_path / "run",
        compile_pdf=False,
    )
    profile = result.template_profile.model_copy(deep=True)
    profile.page.width_mm = 210.0
    profile.page.margin_left_mm = 58.56
    profile.page.margin_right_mm = 12.12
    profile.page.title_margin_left_mm = 12.35
    profile.page.title_margin_right_mm = 12.77

    main = LatexRenderer().render(result.article_ir, tmp_path / "wide", profile)
    rendered = main.read_text(encoding="utf-8")

    assert r"\hspace*{-46.21mm}" in rendered
    assert r"\begin{minipage}{184.88mm}" in rendered
    assert r"\setlength{\droptitle}{25mm}" in rendered


def test_benchmark_log_diagnostics_tracks_overfull(tmp_path: Path) -> None:
    log = tmp_path / "compile.log"
    log.write_text(
        "Overfull \\hbox (12.50pt too wide)\n"
        "Underfull \\hbox\nMissing character: U+1234",
        encoding="utf-8",
    )

    result = TemplateBenchmarkRunner._log_diagnostics(log)

    assert result["max_overfull_pt"] == 12.5
    assert result["underfull_boxes"] == 1
    assert result["missing_glyphs"] == 1
