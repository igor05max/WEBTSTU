from pathlib import Path
from unittest.mock import patch
import base64

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from paper_formatter.models import (
    ArticleIR,
    ArticleMetadata,
    EquationBlock,
    ListItemBlock,
    LocalizedText,
    ParagraphBlock,
    SectionBlock,
    TableBlock,
    TextRun,
)
from paper_formatter.parsers.docx_parser import DocxParser
from paper_formatter.renderers.docx_renderer import DocxRenderer
from paper_formatter.validator import ConversionValidator


def test_docx_parser_extracts_basic_structure(tmp_path: Path) -> None:
    source = tmp_path / "article.docx"
    document = Document()
    document.add_heading("Тестовая научная статья", level=0)
    paragraph = document.add_paragraph("Иванов И. И.")
    paragraph.style = document.styles["Subtitle"]
    document.add_paragraph("Аннотация: Проверка преобразования DOCX.")
    document.add_paragraph("Ключевые слова: DOCX; LaTeX; статья")
    document.add_heading("Введение", level=1)
    document.add_paragraph("Основной текст статьи.")
    document.add_paragraph("Пункт списка", style="List Bullet")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Параметр"
    table.cell(0, 1).text = "Значение"
    table.cell(1, 0).text = "A"
    table.cell(1, 1).text = "10"
    document.save(source)

    article = DocxParser(source, tmp_path / "assets").parse()

    assert article.metadata.titles[0].text == "Тестовая научная статья"
    assert article.metadata.authors[0].name == "Иванов И. И."
    assert article.metadata.abstracts[0].text == "Проверка преобразования DOCX."
    assert article.metadata.keywords == ["DOCX", "LaTeX", "статья"]
    assert any(isinstance(block, SectionBlock) for block in article.body)
    assert any(isinstance(block, ParagraphBlock) for block in article.body)
    assert any(isinstance(block, TableBlock) for block in article.body)
    assert any(isinstance(block, ListItemBlock) for block in article.body)


def test_docx_parser_preserves_typed_heading_number_separately(
    tmp_path: Path,
) -> None:
    source = tmp_path / "numbered-heading.docx"
    document = Document()
    document.add_heading("Тестовая статья", level=0)
    document.add_heading("2.1. Методы исследования", level=1)
    document.add_paragraph("Основной текст.")
    document.save(source)

    article = DocxParser(source, tmp_path / "assets").parse()
    section = next(
        block for block in article.body if isinstance(block, SectionBlock)
    )

    assert section.title == "Методы исследования"
    assert section.number == "2.1"


def test_docx_parser_reads_bibliography_inside_content_control(
    tmp_path: Path,
) -> None:
    source = tmp_path / "mendeley-bibliography.docx"
    document = Document()
    document.add_heading("Test article", level=0)
    document.add_paragraph("References")
    first = document.add_paragraph("First source. DOI: 10.1000/example-1")
    second = document.add_paragraph("Second source.")
    body = document._element.body
    sdt = OxmlElement("w:sdt")
    sdt_pr = OxmlElement("w:sdtPr")
    tag = OxmlElement("w:tag")
    tag.set(qn("w:val"), "MENDELEY_BIBLIOGRAPHY")
    sdt_pr.append(tag)
    content = OxmlElement("w:sdtContent")
    body.remove(first._p)
    body.remove(second._p)
    content.append(first._p)
    content.append(second._p)
    sdt.append(sdt_pr)
    sdt.append(content)
    body.insert(len(body) - 1, sdt)
    document.save(source)

    article = DocxParser(source, tmp_path / "assets").parse()

    assert [reference.text for reference in article.references] == [
        "First source. DOI: 10.1000/example-1",
        "Second source.",
    ]


def test_docx_parser_keeps_every_image_from_distinct_table_cells(
    tmp_path: Path,
) -> None:
    image = tmp_path / "pixel.png"
    image.write_bytes(
        base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVQIHWP4"
            "z8DwHwAFgAI/ScL0RQAAAABJRU5ErkJggg=="
        )
    )
    source = tmp_path / "six-images.docx"
    document = Document()
    document.add_heading("Image table", level=0)
    table = document.add_table(rows=3, cols=2)
    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].add_run().add_picture(str(image))
    document.save(source)

    article = DocxParser(source, tmp_path / "assets").parse()

    assert len(article.assets) == 6


def test_docx_parser_does_not_duplicate_text_from_merged_cells(
    tmp_path: Path,
) -> None:
    source = tmp_path / "merged-header.docx"
    document = Document()
    document.add_heading("Merged table", level=0)
    table = document.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "Mechanical characteristics"
    table.cell(0, 0).merge(table.cell(0, 2))
    table.cell(1, 0).text = "A"
    table.cell(1, 1).text = "B"
    table.cell(1, 2).text = "C"
    document.save(source)

    article = DocxParser(source, tmp_path / "assets").parse()
    parsed = next(block for block in article.body if isinstance(block, TableBlock))

    assert parsed.rows[0] == ["Mechanical characteristics", "", ""]
    assert parsed.merged_cells[0].column_span == 3


def test_docx_parser_recognizes_bilingual_publisher_front_matter(
    tmp_path: Path,
) -> None:
    source = tmp_path / "publisher-front-matter.docx"
    document = Document()

    def centered(text: str, *, bold: bool = False, size: float = 12) -> None:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)

    document.add_paragraph("6-УДК 678.01")
    centered("Тип статьи оригинальная", bold=True)
    centered("Рубрика журнала (необходимо выбрать из списка):", bold=True)
    centered("Повышение стойкости полимерного материала", bold=True, size=16)
    centered("© Е.А. Бойченкоa,b, С.М. Аншинc", bold=True)
    centered("Институт химической физики Российской академии наук")
    document.add_paragraph("Аннотация")
    document.add_paragraph("Краткое описание исследования.")
    document.add_paragraph("Ключевые слова: полимер; испытание")
    centered("Improving polymer material durability", bold=True, size=16)
    centered("© E.A. Boychenkoa,b, S.M. Anshinc", bold=True)
    centered("Research Institute of Chemical Physics")
    document.add_paragraph("Abstract")
    document.add_paragraph("A short description of the study.")
    document.add_paragraph("Keywords: polymer; testing")
    heading = document.add_paragraph("Introduction", style="List Number")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.runs[0].bold = True
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "1")
    num_pr.extend((ilvl, num_id))
    heading._p.get_or_add_pPr().append(num_pr)
    document.add_paragraph("Main body text.")
    document.save(source)

    article = DocxParser(source, tmp_path / "assets").parse()

    assert [item.text for item in article.metadata.titles] == [
        "Повышение стойкости полимерного материала",
        "Improving polymer material durability",
    ]
    assert article.metadata.udc == "678.01"
    assert article.metadata.authors[0].name == "© Е.А. Бойченкоa,b, С.М. Аншинc"
    assert article.metadata.author_variants[0].text == (
        "© E.A. Boychenkoa,b, S.M. Anshinc"
    )
    assert len(article.metadata.affiliations) == 2
    section = next(block for block in article.body if isinstance(block, SectionBlock))
    assert section.title == "Introduction"
    assert section.number == "1"


def test_docx_parser_preserves_inline_omml_inside_text_paragraph(
    tmp_path: Path,
) -> None:
    source = tmp_path / "inline-math.docx"
    document = Document()
    document.add_heading("Тестовая статья", level=0)
    document.add_heading("Методика", level=1)
    paragraph = document.add_paragraph()
    paragraph.add_run("где ")
    math = OxmlElement("m:oMath")
    math_run = OxmlElement("m:r")
    math_text = OxmlElement("m:t")
    math_text.text = "m"
    math_run.append(math_text)
    math.append(math_run)
    paragraph._p.append(math)
    paragraph.add_run(" — поисковый метод.")
    document.save(source)

    article = DocxParser(source, tmp_path / "assets").parse()
    paragraphs = [
        block for block in article.body if isinstance(block, ParagraphBlock)
    ]

    assert len(paragraphs) == 1
    assert [run.text for run in paragraphs[0].runs] == [
        "где ",
        "",
        " — поисковый метод.",
    ]
    assert paragraphs[0].runs[1].math_latex == "m"
    assert not any(isinstance(block, EquationBlock) for block in article.body)
    assert ConversionValidator()._article_counts(article)["equations"] == 1


def test_docx_parser_skips_wmf_conversion_for_preserved_mathtype_ole(
    tmp_path: Path,
) -> None:
    preview_path = tmp_path / "formula.wmf"
    ole_path = tmp_path / "formula.bin"
    preview_path.write_bytes(b"WMF preview")
    ole_path.write_bytes(b"MathType OLE payload")
    object_xml = """
    <w:object xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
              xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
              xmlns:v="urn:schemas-microsoft-com:vml"
              xmlns:o="urn:schemas-microsoft-com:office:office">
      <v:shape style="width:12pt;height:12pt">
        <v:imagedata r:id="rIdPreview"/>
      </v:shape>
      <o:OLEObject Type="Embed" ProgID="Equation.DSMT4" r:id="rIdOle"/>
    </w:object>
    """
    source = DocxRenderer().render(
        ArticleIR(
            metadata=ArticleMetadata(
                titles=[LocalizedText(language="en", text="Formula test")]
            ),
            body=[
                SectionBlock(id="s1", title="Methods", level=1),
                ParagraphBlock(
                    id="p1",
                    runs=[
                        TextRun(text="A sufficiently long body paragraph before "),
                        TextRun(
                            asset_id="formula-preview",
                            formula_image=True,
                            ole_object_xml=object_xml,
                            ole_object_asset_id="formula-ole",
                            ole_preview_asset_id="formula-preview",
                        ),
                        TextRun(text=" after"),
                    ],
                )
            ],
            assets=[
                {
                    "id": "formula-preview",
                    "path": preview_path.name,
                    "media_type": "image/x-wmf",
                },
                {
                    "id": "formula-ole",
                    "path": ole_path.name,
                    "media_type": "application/vnd.openxmlformats-officedocument.oleObject",
                },
            ],
        ),
        tmp_path / "source.docx",
        asset_root=tmp_path,
    )

    with patch(
        "paper_formatter.parsers.docx_parser.convert_metafile_to_png"
    ) as mocked_convert:
        article = DocxParser(source, tmp_path / "parsed-assets").parse()

    mocked_convert.assert_not_called()
    formula_run = next(
        run
        for block in article.body
        if isinstance(block, ParagraphBlock)
        for run in block.runs
        if run.formula_image
    )
    assert formula_run.ole_object_asset_id is not None
    assert formula_run.ole_preview_asset_id is not None
    formula_asset = next(
        asset for asset in article.assets if asset.id == formula_run.asset_id
    )
    assert Path(formula_asset.path).suffix.lower() == ".wmf"


def test_docx_parser_recognizes_numbered_equation_table(
    tmp_path: Path,
) -> None:
    source = tmp_path / "equation-table.docx"
    document = Document()
    document.add_heading("Тестовая статья", level=0)
    table = document.add_table(rows=1, cols=2)
    paragraph = table.cell(0, 0).paragraphs[0]
    math_paragraph = OxmlElement("m:oMathPara")
    math = OxmlElement("m:oMath")
    math_run = OxmlElement("m:r")
    math_text = OxmlElement("m:t")
    math_text.text = "x"
    math_run.append(math_text)
    math.append(math_run)
    math_paragraph.append(math)
    paragraph._p.append(math_paragraph)
    table.cell(0, 1).text = "(7)"
    document.save(source)

    article = DocxParser(source, tmp_path / "assets").parse()
    equations = [
        block for block in article.body if isinstance(block, EquationBlock)
    ]

    assert len(equations) == 1
    assert equations[0].latex == "x"
    assert equations[0].number == "7"
    assert not any(isinstance(block, TableBlock) for block in article.body)
