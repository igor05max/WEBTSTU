from pathlib import Path
import base64
from zipfile import ZipFile

import pytest
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.opc.part import Part
from docx.shared import Mm, Pt, RGBColor

from paper_formatter.models import (
    Affiliation,
    ArticleIR,
    ArticleMetadata,
    Author,
    FigureBlock,
    HeadingStyleProfile,
    LocalizedText,
    PageLayout,
    ParagraphBlock,
    ReferenceEntry,
    SectionBlock,
    TableBlock,
    TemplateProfile,
    TextRun,
    TypographyProfile,
)
from paper_formatter.renderers.docx_renderer import DocxRenderer
from paper_formatter.template_analyzers.docx_analyzer import DocxTemplateAnalyzer
from paper_formatter.template_analyzers.pdf_analyzer import PdfTemplateAnalyzer


def test_docx_renderer_applies_template_geometry_and_typography(
    tmp_path: Path,
) -> None:
    profile = TemplateProfile(
        page=PageLayout(
            width_mm=210.0,
            height_mm=297.0,
            margin_top_mm=26.34,
            margin_right_mm=12.12,
            margin_bottom_mm=25.07,
            margin_left_mm=58.56,
            title_margin_left_mm=12.35,
            title_margin_right_mm=12.70,
        ),
        typography=TypographyProfile(
            main_font="Palatino Linotype",
            main_size_pt=10.1,
            line_spacing=1.38,
            first_line_indent_mm=7.99,
            paragraph_space_before_pt=0.0,
            paragraph_space_after_pt=0.0,
            paragraph_alignment="justify",
            title_font="Palatino Linotype",
            title_size_pt=17.9,
            caption_size_pt=8.0,
        ),
        headings=[
            HeadingStyleProfile(
                level=1,
                font="Palatino Linotype",
                size_pt=12.0,
                space_before_pt=14.4,
                space_after_pt=6.0,
            )
        ],
    )
    article = ArticleIR(
        metadata=ArticleMetadata(
            titles=[LocalizedText(language="en", text="Template test")]
        ),
        body=[
            SectionBlock(id="s1", title="Introduction", level=1),
            ParagraphBlock(id="p1", runs=[TextRun(text="Body paragraph.")]),
            TableBlock(
                id="t1",
                rows=[["A", "B"], ["1", "2"]],
                header_rows=1,
            ),
        ],
        references=[
            ReferenceEntry(
                id="r1",
                text="First reference with a wrapped continuation line.",
            ),
            ReferenceEntry(id="r2", text="Second reference."),
        ],
    )

    output = DocxRenderer().render(article, tmp_path / "result.docx", profile=profile)
    document = Document(output)
    section = document.sections[0]
    normal = document.styles["Normal"]
    title = document.styles["Title"]
    heading = document.styles["Heading 1"]

    assert section.left_margin.mm == pytest.approx(58.56, abs=0.02)
    assert section.right_margin.mm == pytest.approx(12.12, abs=0.02)
    assert section.top_margin.mm == pytest.approx(26.34, abs=0.02)
    assert section.bottom_margin.mm == pytest.approx(25.07, abs=0.02)
    assert normal.font.name == "Palatino Linotype"
    assert normal.font.size.pt == pytest.approx(10.0, abs=0.1)
    assert normal.paragraph_format.line_spacing == pytest.approx(1.38, abs=0.01)
    assert normal.paragraph_format.first_line_indent.mm == pytest.approx(
        7.99,
        abs=0.02,
    )
    assert normal.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
    assert title.font.name == "Palatino Linotype"
    assert title.font.size.pt == pytest.approx(17.5, abs=0.1)
    assert title.font.color.rgb == RGBColor(0, 0, 0)
    assert title.element.pPr.find(qn("w:pBdr")) is None
    assert document.paragraphs[0].paragraph_format.left_indent.mm == pytest.approx(
        -46.21,
        abs=0.02,
    )
    assert document.paragraphs[0].paragraph_format.right_indent.mm == pytest.approx(
        0.58,
        abs=0.02,
    )
    assert heading.font.size.pt == pytest.approx(12.0, abs=0.1)
    assert heading.font.color.rgb == RGBColor(0, 0, 0)
    assert document.tables[0].cell(0, 0).paragraphs[
        0
    ].paragraph_format.first_line_indent.mm == pytest.approx(0.0)
    table_xml = document.tables[0]._tbl.xml
    assert 'w:type="fixed"' in table_xml
    assert "<w:tblHeader" in table_xml
    assert 'w:val="nil"' in table_xml
    assert document.tables[0].cell(1, 0).paragraphs[0].alignment == (
        WD_ALIGN_PARAGRAPH.CENTER
    )
    reference_paragraphs = document.paragraphs[-2:]
    assert [paragraph.text for paragraph in reference_paragraphs] == [
        ".      First reference with a wrapped continuation line.",
        ".      Second reference.",
    ]
    assert all(
        paragraph.style.name != "List Number"
        for paragraph in reference_paragraphs
    )
    assert reference_paragraphs[0].paragraph_format.left_indent.mm == pytest.approx(
        -38.21,
        abs=0.02,
    )
    assert reference_paragraphs[0].paragraph_format.first_line_indent.mm == (
        pytest.approx(-8.0, abs=0.02)
    )
    fields = reference_paragraphs[0]._p.xpath("./w:fldSimple")
    assert fields
    assert "SEQ PFBibliography" in fields[0].get(qn("w:instr"))


def test_pdf_profile_geometry_ignores_sidebar_and_detects_indent() -> None:
    analyzer = PdfTemplateAnalyzer()
    samples = [
        {
            "page": 1,
            "block": 1,
            "line": 0,
            "bbox": (188.0, 75.0, 559.0, 85.0),
        },
        {
            "page": 1,
            "block": 1,
            "line": 1,
            "bbox": (166.0, 89.0, 560.0, 99.0),
        },
        {
            "page": 1,
            "block": 2,
            "line": 0,
            "bbox": (188.0, 117.0, 559.0, 127.0),
        },
        {
            "page": 1,
            "block": 2,
            "line": 1,
            "bbox": (166.0, 131.0, 560.0, 141.0),
        },
    ]

    assert analyzer._modal_edge(
        [35.0, 166.1, 165.9, 166.2, 166.0],
        fallback=56.7,
    ) == 166.0
    assert analyzer._first_line_indent_mm(samples, [166.0]) == pytest.approx(
        7.76,
        abs=0.02,
    )
    assert analyzer._word_font_name("URWPalladioL-Roma") == "Palatino Linotype"


def test_docx_template_uses_multicolumn_body_not_opening_section(
    tmp_path: Path,
) -> None:
    template_path = tmp_path / "journal-template.docx"
    template = Document()
    opening = template.sections[0]
    opening.left_margin = Mm(28.5)
    opening.right_margin = Mm(28.5)
    template.styles["Normal"].font.name = "Times New Roman"
    template.styles["Normal"].font.size = Pt(12)
    title = template.add_paragraph("JOURNAL ARTICLE TITLE")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run().bold = True
    opening.header.paragraphs[0].text = "Journal running header"

    body = template.add_section(WD_SECTION.CONTINUOUS)
    body.left_margin = Mm(20)
    body.right_margin = Mm(20)
    body.top_margin = Mm(19.5)
    body.bottom_margin = Mm(30)
    columns = OxmlElement("w:cols")
    columns.set(qn("w:num"), "2")
    columns.set(qn("w:space"), str(round(5.0 * 1440 / 25.4)))
    body._sectPr.append(columns)
    page_numbering = OxmlElement("w:pgNumType")
    page_numbering.set(qn("w:start"), "9")
    body._sectPr.append(page_numbering)
    custom = template.styles.add_style(
        "ArticleHeading1",
        WD_STYLE_TYPE.PARAGRAPH,
    )
    custom.font.name = "Times New Roman"
    custom.font.size = Pt(10.5)
    custom.font.bold = True
    custom.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    template.add_paragraph("BODY", style=custom)
    template.save(template_path)

    profile = DocxTemplateAnalyzer().analyze(template_path)

    assert profile.page.columns == 2
    assert profile.page.column_gap_mm == pytest.approx(5.0, abs=0.03)
    assert profile.page.margin_left_mm == pytest.approx(20.0, abs=0.02)
    assert profile.page.title_margin_left_mm == pytest.approx(28.5, abs=0.02)
    assert profile.typography.title_size_pt == pytest.approx(12.0, abs=0.1)
    assert profile.headings[0].size_pt == pytest.approx(10.5, abs=0.1)
    assert profile.headings[0].alignment == "center"

    article = ArticleIR(
        metadata=ArticleMetadata(
            titles=[LocalizedText(language="en", text="Generated title")]
        ),
        body=[ParagraphBlock(id="p1", runs=[TextRun(text="Body text.")])],
    )
    output = DocxRenderer().render(
        article,
        tmp_path / "generated.docx",
        profile=profile,
    )
    rendered = Document(output)
    assert len(rendered.sections) == 2
    first_columns = rendered.sections[0]._sectPr.xpath("./w:cols")[0]
    body_columns = rendered.sections[1]._sectPr.xpath("./w:cols")[0]
    assert first_columns.get(qn("w:num")) == "1"
    assert body_columns.get(qn("w:num")) == "2"
    for section in rendered.sections:
        page_numbering = section._sectPr.find(qn("w:pgNumType"))
        assert page_numbering is None or page_numbering.get(qn("w:start")) is None
    assert rendered.sections[0].header.paragraphs[0].text == (
        "Journal running header"
    )


def test_docx_template_preserves_at_least_line_spacing_and_mdpi_title(
    tmp_path: Path,
) -> None:
    template_path = tmp_path / "mdpi.docx"
    template = Document()
    spacing = template.styles["Normal"].element.get_or_add_pPr().get_or_add_spacing()
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "atLeast")
    title_style = template.styles.add_style("MDPI_1.2_title", WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = "Palatino Linotype"
    title_style.font.size = Pt(18)
    title_style.font.bold = True
    template.add_paragraph("Article", style="Normal")
    template.add_paragraph("Template title", style="MDPI_1.2_title")
    template.save(template_path)

    profile = DocxTemplateAnalyzer().analyze(template_path)

    assert profile.typography.line_spacing == 1.0
    assert profile.typography.line_spacing_pt == pytest.approx(14.0)
    assert profile.typography.line_spacing_rule == "atLeast"
    assert profile.typography.title_size_pt == pytest.approx(18.0)
    assert profile.typography.title_alignment == "left"


def test_docx_template_analyzer_ignores_styles_without_names(
    tmp_path: Path,
) -> None:
    template_path = tmp_path / "publisher-template.docx"
    template = Document()
    unnamed = OxmlElement("w:style")
    unnamed.set(qn("w:type"), "paragraph")
    unnamed.set(qn("w:styleId"), "PublisherOrphanStyle")
    template.styles.element.append(unnamed)
    template.add_paragraph("Template body")
    template.save(template_path)

    profile = DocxTemplateAnalyzer().analyze(template_path)

    assert profile.source_type == "docx"


def test_docx_template_ignores_unused_legacy_indented_body_and_caption_styles(
    tmp_path: Path,
) -> None:
    template_path = tmp_path / "legacy-publisher-template.docx"
    template = Document()
    body_indent = template.styles.add_style(
        "Body Text Indent", WD_STYLE_TYPE.PARAGRAPH
    )
    body_indent.paragraph_format.right_indent = Mm(45)
    caption = template.styles["Caption"]
    caption.paragraph_format.right_indent = Mm(45)
    template.add_paragraph("Заглавие", style="Normal")
    template.add_paragraph("A long demonstrated body paragraph " * 10, style="Normal")
    template.save(template_path)

    profile = DocxTemplateAnalyzer().analyze(template_path)
    article = ArticleIR(
        metadata=ArticleMetadata(
            titles=[LocalizedText(language="ru", text="Название статьи")]
        ),
        body=[
            ParagraphBlock(id="p1", runs=[TextRun(text="Основной текст")]),
            FigureBlock(id="f1", asset_id="missing", caption="Подпись"),
        ],
    )
    output = DocxRenderer().render(
        article,
        tmp_path / "result.docx",
        profile=profile,
    )
    result = Document(output)

    assert profile.evidence["docx_style_body"] == "Normal"
    assert profile.evidence["docx_style_title"] == "Normal"
    assert profile.evidence["docx_style_figure_caption"] == "Normal"
    body = next(p for p in result.paragraphs if p.text == "Основной текст")
    figure_caption = next(p for p in result.paragraphs if p.text.endswith("Подпись"))
    assert body.paragraph_format.right_indent is None
    assert figure_caption.paragraph_format.right_indent is None


def test_docx_journal_template_keeps_bilingual_front_matter_and_figure_group(
    tmp_path: Path,
) -> None:
    template_path = tmp_path / "bilingual-journal.docx"
    template = Document()
    template.settings.odd_and_even_pages_header_footer = True
    template.sections[0].header.paragraphs[0].text = "Journal. 2026. Vol. 11"
    template.sections[0].even_page_header.paragraphs[0].text = (
        "2021;1(21):00-00    Journal"
    )
    template.sections[0].footer.paragraphs[0].text = "Surname I.O."
    template.add_paragraph("УДК 000    DOI: 10.1000/example")
    title = template.add_paragraph("Заглавие")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.name = "Times New Roman"
    title_run.font.size = Pt(14)
    template.add_paragraph("Abstract.")
    template.add_paragraph("Заглавие")
    template.add_paragraph("Аннотация.")
    template.add_paragraph(
        "For citation: Journal. 2026;11(3):000-000. DOI: 10.1000/example"
    )
    body_section = template.add_section(WD_SECTION.CONTINUOUS)
    columns = OxmlElement("w:cols")
    columns.set(qn("w:num"), "2")
    body_section._sectPr.append(columns)
    template.save(template_path)

    image_path = tmp_path / "pixel.png"
    image_path.write_bytes(
        base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVQIHWP4"
            "z8DwHwAFgAI/ScL0RQAAAABJRU5ErkJggg=="
        )
    )
    profile = DocxTemplateAnalyzer().analyze(template_path)
    article = ArticleIR(
        metadata=ArticleMetadata(
            titles=[
                LocalizedText(language="ru", text="Русское название"),
                LocalizedText(language="en", text="English title"),
            ],
            authors=[Author(id="a1", name="© Е.А. Авторa, Е.Б. Второйb")],
            author_variants=[
                LocalizedText(language="en", text="© E.A. Authora, E.B. Secondb")
            ],
            affiliations=[
                Affiliation(id="af1", name="a Российский университет"),
                Affiliation(id="af2", name="a Research University"),
            ],
            abstracts=[
                LocalizedText(language="ru", text="Русская аннотация."),
                LocalizedText(language="en", text="English abstract."),
            ],
            keywords=["материал", "material"],
            udc="123.4",
        ),
        body=[
            SectionBlock(id="s1", title="Introduction", level=1),
            FigureBlock(id="f1", asset_id="img1", group_id="g1"),
            FigureBlock(
                id="f2",
                asset_id="img2",
                group_id="g1",
                caption="Two panels",
            ),
        ],
        assets=[
            {"id": "img1", "path": "pixel.png", "media_type": "image/png"},
            {"id": "img2", "path": "pixel.png", "media_type": "image/png"},
        ],
    )

    output = DocxRenderer().render(
        article,
        tmp_path / "bilingual-result.docx",
        profile=profile,
        asset_root=tmp_path,
    )
    result = Document(output)
    texts = [paragraph.text for paragraph in result.paragraphs]

    assert profile.evidence["bilingual_front_matter"] is True
    assert profile.evidence["identifier_before_title"] is True
    assert profile.typography.title_font == "Times New Roman"
    assert texts.index("УДК: 123.4") < texts.index("Русское название")
    assert texts.index("English title") < texts.index("Introduction")
    assert len(result.tables) == 1
    assert len(result.tables[0].rows) == 1
    assert "2026;11(3):000-000" in result.sections[0].even_page_header.paragraphs[0].text
    assert "Author E.A. et al." in result.sections[0].footer.paragraphs[0].text


def test_docx_renderer_moves_wide_formula_to_its_own_line(tmp_path: Path) -> None:
    image_path = tmp_path / "formula.png"
    image_path.write_bytes(
        base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVQIHWP4"
            "z8DwHwAFgAI/ScL0RQAAAABJRU5ErkJggg=="
        )
    )
    article = ArticleIR(
        body=[
            ParagraphBlock(
                id="p1",
                runs=[
                    TextRun(text="Before "),
                    TextRun(
                        asset_id="formula",
                        formula_image=True,
                        width_pt=300,
                        height_pt=24,
                    ),
                    TextRun(text=" after"),
                ],
            )
        ],
        assets=[
            {
                "id": "formula",
                "path": "formula.png",
                "media_type": "image/png",
            }
        ],
    )
    output = DocxRenderer().render(
        article,
        tmp_path / "result.docx",
        profile=TemplateProfile(),
        asset_root=tmp_path,
    )
    paragraphs = Document(output).paragraphs

    assert [paragraph.text for paragraph in paragraphs] == ["Before ", "", " after"]
    assert paragraphs[1].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert "w:drawing" in paragraphs[1]._p.xml


def test_docx_renderer_preserves_ole_formula_when_preview_is_wmf(
    tmp_path: Path,
) -> None:
    preview_path = tmp_path / "formula.wmf"
    ole_path = tmp_path / "formula.bin"
    preview_path.write_bytes(b"WMF preview")
    ole_path.write_bytes(b"MathType OLE payload")
    object_xml = """
    <w:object xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
              xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
              xmlns:v="urn:schemas-microsoft-com:vml"
              xmlns:o="urn:schemas-microsoft-com:office:office">
      <v:shape style="width:12pt;height:12pt">
        <v:imagedata r:id="rIdPreview"/>
      </v:shape>
      <o:OLEObject Type="Embed" ProgID="Equation.DSMT4" r:id="rIdOle"/>
    </w:object>
    """
    article = ArticleIR(
        body=[
            ParagraphBlock(
                id="p1",
                runs=[
                    TextRun(text="Before "),
                    TextRun(
                        asset_id="formula-preview",
                        formula_image=True,
                        ole_object_xml=object_xml,
                        ole_object_asset_id="formula-ole",
                        ole_preview_asset_id="formula-preview",
                    ),
                    TextRun(text=" after"),
                ],
            )
        ],
        assets=[
            {
                "id": "formula-preview",
                "path": preview_path.name,
                "media_type": "image/x-wmf",
            },
            {
                "id": "formula-ole",
                "path": ole_path.name,
                "media_type": "application/vnd.openxmlformats-officedocument.oleObject",
            },
        ],
    )

    output = DocxRenderer().render(
        article,
        tmp_path / "ole-result.docx",
        asset_root=tmp_path,
    )
    document = Document(output)

    assert document.paragraphs[0].text == "Before  after"
    assert "[формула]" not in document.paragraphs[0].text
    assert "OLEObject" in document.paragraphs[0]._p.xml
    with ZipFile(output) as archive:
        assert any(name.startswith("word/embeddings/") for name in archive.namelist())


def test_docx_renderer_renames_duplicate_package_parts(tmp_path: Path) -> None:
    document = Document()
    first = Part(
        PackURI("/word/media/image1.png"),
        "image/png",
        b"first",
        document.part.package,
    )
    second = Part(
        PackURI("/word/media/image1.png"),
        "image/png",
        b"second",
        document.part.package,
    )
    document.part.relate_to(first, RT.IMAGE)
    document.part.relate_to(second, RT.IMAGE)

    renamed = DocxRenderer._ensure_unique_package_partnames(document)
    output = tmp_path / "unique-parts.docx"
    document.save(output)

    assert renamed == 1
    with ZipFile(output) as archive:
        names = archive.namelist()
    assert len(names) == len(set(names))


def test_docx_renderer_uses_template_body_font_and_style(tmp_path: Path) -> None:
    template_path = tmp_path / "font-template.docx"
    template = Document()
    body_style = template.styles.add_style(
        "MDPI_3.1_text",
        WD_STYLE_TYPE.PARAGRAPH,
    )
    body_style.font.name = "Palatino Linotype"
    body_style.font.size = Pt(11)
    template.add_paragraph("Template body", style=body_style)
    template.save(template_path)
    profile = DocxTemplateAnalyzer().analyze(template_path)
    article = ArticleIR(
        body=[ParagraphBlock(id="p1", runs=[TextRun(text="Generated body")])]
    )

    output = DocxRenderer().render(
        article,
        tmp_path / "font-result.docx",
        profile=profile,
    )
    paragraph = Document(output).paragraphs[0]

    assert paragraph.style.name == "MDPI_3.1_text"
    assert paragraph.style.font.name == "Palatino Linotype"


def test_docx_renderer_prints_typed_heading_number_without_word_numbering(
    tmp_path: Path,
) -> None:
    article = ArticleIR(
        body=[
            SectionBlock(
                id="section-1",
                title="Методы исследования",
                number="2.1",
                level=2,
            )
        ]
    )

    output = DocxRenderer().render(article, tmp_path / "numbered.docx")
    heading = Document(output).paragraphs[0]

    assert heading.text == "2.1. Методы исследования"
