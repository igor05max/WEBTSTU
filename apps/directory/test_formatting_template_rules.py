import io

from django.test import SimpleTestCase

from apps.directory.formatting_templates import _extract_docx_rules


class DocxFormattingRuleExtractionTests(SimpleTestCase):
    def test_semantic_rules_come_from_custom_style_definitions(self):
        from docx import Document
        from docx.enum.style import WD_STYLE_TYPE
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Cm, Pt

        document = Document()
        normal = document.styles["Normal"]
        normal.font.name = "Palatino Linotype"
        normal.font.size = Pt(10)
        normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        normal.paragraph_format.first_line_indent = Cm(2)
        normal.paragraph_format.line_spacing = 2

        title_style = document.styles.add_style(
            "Publisher_Title",
            WD_STYLE_TYPE.PARAGRAPH,
        )
        title_style.base_style = normal
        title_style.font.size = Pt(18)
        title_style.font.bold = True

        heading_style = document.styles.add_style(
            "Publisher_Heading1",
            WD_STYLE_TYPE.PARAGRAPH,
        )
        heading_style.base_style = normal
        heading_style.font.size = Pt(12)
        heading_style.font.bold = True
        heading_style.paragraph_format.space_before = Pt(12)

        body_style = document.styles.add_style(
            "Publisher_MainText",
            WD_STYLE_TYPE.PARAGRAPH,
        )
        body_style.base_style = normal
        body_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        body_style.paragraph_format.first_line_indent = Cm(0.75)
        body_style.paragraph_format.line_spacing = Pt(14)

        reference_style = document.styles.add_style(
            "Publisher_References",
            WD_STYLE_TYPE.PARAGRAPH,
        )
        reference_style.base_style = normal
        reference_style.font.size = Pt(9)
        reference_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        reference_style.paragraph_format.first_line_indent = Cm(-0.75)
        reference_style.paragraph_format.line_spacing = Pt(14)

        caption_style = document.styles.add_style(
            "Publisher_Figure_Caption",
            WD_STYLE_TYPE.PARAGRAPH,
        )
        caption_style.base_style = normal
        caption_style.font.size = Pt(9)
        caption_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_style.paragraph_format.line_spacing = Pt(14)

        title = document.add_paragraph("Placeholder title", style=title_style)
        title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        title.paragraph_format.first_line_indent = Cm(4)
        title.paragraph_format.line_spacing = 3
        title_run = title.runs[0]
        title_run.font.name = "Arial"
        title_run.font.size = Pt(30)

        heading = document.add_paragraph("1. Placeholder", style=heading_style)
        heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        heading.paragraph_format.space_before = Pt(30)
        heading.runs[0].font.size = Pt(28)

        for index in range(3):
            body = document.add_paragraph(
                f"Scientific body paragraph {index}.",
                style=body_style,
            )
            body.paragraph_format.first_line_indent = Cm(4)
            body.paragraph_format.line_spacing = 3
            body.runs[0].font.name = "Arial"
            body.runs[0].font.size = Pt(30)

        reference = document.add_paragraph(
            "Reference placeholder.",
            style=reference_style,
        )
        reference.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        reference.paragraph_format.first_line_indent = Cm(4)
        reference.runs[0].font.size = Pt(30)

        caption = document.add_paragraph(
            "Figure 1. Placeholder.",
            style=caption_style,
        )
        caption.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        caption.runs[0].font.size = Pt(30)

        output = io.BytesIO()
        document.save(output)
        rules = _extract_docx_rules(output.getvalue())

        blocks = {
            block["role"]: block["style"]
            for block in rules["document"]["blocks"]
        }
        self.assertEqual(blocks["title"]["font_family"], "Palatino Linotype")
        self.assertEqual(blocks["title"]["font_size_pt"], 18.0)
        self.assertEqual(blocks["title"]["alignment"], "left")
        self.assertEqual(blocks["title"]["first_line_indent_cm"], 0)
        self.assertNotIn("line_spacing", blocks["title"])

        self.assertEqual(rules["body"]["font_family"], "Palatino Linotype")
        self.assertEqual(rules["body"]["font_size_pt"], 10.0)
        self.assertEqual(rules["body"]["line_spacing"], 1.4)
        self.assertEqual(rules["body"]["first_line_indent_cm"], 0.75)
        self.assertEqual(rules["headings"]["levels"]["1"]["font_size_pt"], 12.0)
        self.assertEqual(
            rules["headings"]["levels"]["1"]["space_before_pt"],
            12.0,
        )

        self.assertEqual(blocks["references"]["font_size_pt"], 9.0)
        self.assertEqual(blocks["references"]["line_spacing"], 1.56)
        self.assertEqual(
            blocks["references"]["first_line_indent_cm"],
            -0.75,
        )
        figure_caption = rules["document"]["paragraph_styles"]["figure_caption"]
        self.assertEqual(figure_caption["font_size_pt"], 9.0)
        self.assertEqual(figure_caption["alignment"], "center")

    def test_uses_inherited_body_style_instead_of_reference_overrides(self):
        from docx import Document
        from docx.enum.style import WD_STYLE_TYPE
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Cm, Pt

        document = Document()
        normal = document.styles["Normal"]
        normal.font.name = "Palatino Linotype"
        normal.font.size = Pt(10)

        body_style = document.styles.add_style("MDPI_3.1_text", WD_STYLE_TYPE.PARAGRAPH)
        body_style.base_style = normal
        body_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        body_style.paragraph_format.first_line_indent = Cm(0.75)
        body_style.paragraph_format.line_spacing = Pt(14)

        heading_style = document.styles.add_style(
            "MDPI_2.1_heading1",
            WD_STYLE_TYPE.PARAGRAPH,
        )
        heading_style.base_style = normal
        heading_style.font.size = Pt(12)

        title_style = document.styles.add_style(
            "MDPI_1.2_title",
            WD_STYLE_TYPE.PARAGRAPH,
        )
        title_style.base_style = normal
        title_style.font.size = Pt(24)

        reference_style = document.styles.add_style(
            "MDPI_8.1_references",
            WD_STYLE_TYPE.PARAGRAPH,
        )
        reference_style.base_style = normal
        reference_style.font.size = Pt(9)
        reference_style.paragraph_format.first_line_indent = Cm(-0.75)
        reference_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

        document.add_paragraph("Title", style=title_style)
        document.add_paragraph("1. Introduction", style=heading_style)
        for index in range(5):
            document.add_paragraph(f"Main text paragraph {index}.", style=body_style)
        for index in range(8):
            document.add_paragraph(f"Reference {index}.", style=reference_style)

        output = io.BytesIO()
        document.save(output)
        rules = _extract_docx_rules(output.getvalue())

        self.assertEqual(rules["body"]["font_family"], "Palatino Linotype")
        self.assertEqual(rules["body"]["font_size_pt"], 10.0)
        self.assertEqual(rules["body"]["line_spacing"], 1.4)
        self.assertEqual(rules["body"]["first_line_indent_cm"], 0.75)
        self.assertEqual(rules["body"]["alignment"], "justify")
        self.assertEqual(rules["headings"]["font_size_pt"], 12.0)
        self.assertEqual(rules["headings"]["title_font_size_pt"], 24.0)

    def test_reads_complex_script_size_used_by_mdpi_template(self):
        from docx import Document
        from docx.enum.style import WD_STYLE_TYPE
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt

        document = Document()
        normal = document.styles["Normal"]
        normal.font.name = None
        normal.font.size = None

        body_style = document.styles.add_style(
            "MDPI_3.1_text",
            WD_STYLE_TYPE.PARAGRAPH,
        )
        body_style.font.name = "Palatino Linotype"
        run_properties = body_style.element.get_or_add_rPr()
        complex_script_size = OxmlElement("w:szCs")
        complex_script_size.set(qn("w:val"), "22")
        run_properties.append(complex_script_size)
        body_style.paragraph_format.line_spacing = Pt(14)
        body_style.paragraph_format.first_line_indent = Cm(0.75)
        body_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        paragraph = document.add_paragraph(style=body_style)
        paragraph.add_run("A representative MDPI body paragraph.")

        output = io.BytesIO()
        document.save(output)
        rules = _extract_docx_rules(output.getvalue())

        self.assertEqual(rules["body"]["font_size_pt"], 11.0)
        self.assertEqual(rules["body"]["line_spacing"], 1.27)
