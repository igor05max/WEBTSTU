import io
from collections import Counter
from pathlib import Path
import re
from statistics import median
from zipfile import ZIP_DEFLATED, BadZipFile, ZipFile

from django.db import models, transaction

from apps.checks.ai_client import (
    AIProviderError,
    extract_response_text,
    generate_content,
    get_configured_model,
    get_provider,
    is_ai_configured,
)
from apps.directory.models import (
    FormattingTemplate,
    FormattingTemplateStatus,
)
from apps.submissions.document_analysis import (
    TEXT_EXTENSIONS,
    analyze_document_bytes,
    read_file_bytes,
)
from apps.submissions.document_conversion import convert_legacy_doc_to_docx
from document_template_engine import (
    extract_latex_template_rules,
    interpret_template_text,
    normalize_template_rules,
)


TEMPLATE_EXTENSIONS = {
    ".docx",
    ".dotx",
    ".doc",
    ".pdf",
    ".tex",
    ".txt",
    ".md",
    ".rtf",
    ".png",
    ".jpg",
    ".jpeg",
    ".webp",
}

DEFAULT_RULES = {
    "article": {"limits": {"min_words": 2000, "max_words": 12000}},
    "monograph": {"limits": {"min_words": 10000, "max_words": 200000}},
    "theses": {"limits": {"min_words": 500, "max_words": 5000}},
}

FORMATTING_TEMPLATE_ANALYSIS_REVISION = 3


def _round_cm(length):
    if length is None:
        return None
    return round(float(length.cm), 2)


def _round_pt(length):
    if length is None:
        return None
    return round(float(length.pt), 1)


def _dominant(values):
    cleaned = [value for value in values if value not in (None, "")]
    if not cleaned:
        return None
    return Counter(cleaned).most_common(1)[0][0]


_NON_BODY_STYLE_TOKENS = (
    "abstract",
    "affiliation",
    "article_type",
    "author",
    "back_matter",
    "bullet",
    "caption",
    "heading",
    "itemize",
    "keyword",
    "list",
    "reference",
    "title",
    "аннотац",
    "автор",
    "заголов",
    "ключев",
    "литератур",
    "назван",
    "подпис",
    "список",
)


def _style_chain(style):
    while style is not None:
        yield style
        style = style.base_style


def _first_style_value(style, getter):
    for candidate in _style_chain(style):
        value = getter(candidate)
        if value is not None:
            return value
    return None


def _style_font_size(style):
    try:
        from docx.oxml.ns import qn
    except ImportError:
        return None
    for candidate in _style_chain(style):
        direct_size = _round_pt(candidate.font.size)
        if direct_size is not None:
            return direct_size
        run_properties = candidate.element.find(qn("w:rPr"))
        if run_properties is None:
            continue
        for size_name in ("w:sz", "w:szCs"):
            size = run_properties.find(qn(size_name))
            if size is None:
                continue
            try:
                return round(float(size.get(qn("w:val"))) / 2, 1)
            except (TypeError, ValueError):
                continue
    return None


def _style_font_color_hex(style, *, automatic=""):
    for candidate in _style_chain(style):
        try:
            rgb = candidate.font.color.rgb
        except (AttributeError, TypeError, ValueError):
            rgb = None
        value = str(rgb or "").strip().lstrip("#")
        if re.fullmatch(r"[0-9a-fA-F]{6}", value):
            return value.upper()
    return automatic


def _document_default_font_size(document):
    try:
        from docx.oxml.ns import qn
    except ImportError:
        return None
    defaults = document.styles.element.find(qn("w:docDefaults"))
    if defaults is None:
        return None
    run_defaults = defaults.find(qn("w:rPrDefault"))
    run_properties = run_defaults.find(qn("w:rPr")) if run_defaults is not None else None
    if run_properties is None:
        return None
    for size_name in ("w:sz", "w:szCs"):
        size = run_properties.find(qn(size_name))
        if size is None:
            continue
        try:
            return round(float(size.get(qn("w:val"))) / 2, 1)
        except (TypeError, ValueError):
            continue
    return None


def _document_default_font_family(document):
    from docx.oxml.ns import qn

    try:
        styles_root = document.styles.element
    except (AttributeError, TypeError):
        return ""
    defaults = styles_root.find(qn("w:docDefaults"))
    run_defaults = (
        defaults.find(qn("w:rPrDefault"))
        if defaults is not None
        else None
    )
    run_properties = (
        run_defaults.find(qn("w:rPr"))
        if run_defaults is not None
        else None
    )
    fonts = (
        run_properties.find(qn("w:rFonts"))
        if run_properties is not None
        else None
    )
    if fonts is None:
        return ""
    for attribute in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        value = str(fonts.get(qn(attribute)) or "").strip()
        if value and not value.startswith(("+", "major", "minor")):
            return value
    return ""


def _resolved_style_rules(document, paragraphs):
    if not paragraphs:
        return {}
    normal_style = document.styles["Normal"]
    representative_style = paragraphs[0].style or normal_style
    style_font_size = _style_font_size(representative_style)

    font_names = []
    font_sizes = []
    line_spacings = []
    first_line_indents = []
    left_indents = []
    right_indents = []
    spaces_before = []
    spaces_after = []
    alignments = []
    keep_with_next_values = []
    keep_together_values = []
    for paragraph in paragraphs:
        for run in paragraph.runs:
            if not run.text.strip():
                continue
            if run.font.name:
                font_names.append(run.font.name)
            if run.font.size is not None:
                font_sizes.append(_round_pt(run.font.size))

        formatting = paragraph.paragraph_format
        spacing = formatting.line_spacing
        if spacing is None:
            spacing = _first_style_value(
                paragraph.style or representative_style,
                lambda style: style.paragraph_format.line_spacing,
            )
        if spacing is None:
            spacing = normal_style.paragraph_format.line_spacing
        indent = formatting.first_line_indent
        if indent is None:
            indent = _first_style_value(
                paragraph.style or representative_style,
                lambda style: style.paragraph_format.first_line_indent,
            )
        if indent is None:
            indent = normal_style.paragraph_format.first_line_indent
        left_indent = formatting.left_indent
        if left_indent is None:
            left_indent = _first_style_value(
                paragraph.style or representative_style,
                lambda style: style.paragraph_format.left_indent,
            )
        if left_indent is None:
            left_indent = normal_style.paragraph_format.left_indent
        right_indent = formatting.right_indent
        if right_indent is None:
            right_indent = _first_style_value(
                paragraph.style or representative_style,
                lambda style: style.paragraph_format.right_indent,
            )
        if right_indent is None:
            right_indent = normal_style.paragraph_format.right_indent
        space_before = formatting.space_before
        if space_before is None:
            space_before = _first_style_value(
                paragraph.style or representative_style,
                lambda style: style.paragraph_format.space_before,
            )
        if space_before is None:
            space_before = normal_style.paragraph_format.space_before
        space_after = formatting.space_after
        if space_after is None:
            space_after = _first_style_value(
                paragraph.style or representative_style,
                lambda style: style.paragraph_format.space_after,
            )
        if space_after is None:
            space_after = normal_style.paragraph_format.space_after
        keep_with_next = formatting.keep_with_next
        if keep_with_next is None:
            keep_with_next = _first_style_value(
                paragraph.style or representative_style,
                lambda style: style.paragraph_format.keep_with_next,
            )
        keep_together = formatting.keep_together
        if keep_together is None:
            keep_together = _first_style_value(
                paragraph.style or representative_style,
                lambda style: style.paragraph_format.keep_together,
            )
        alignment = paragraph.alignment
        if alignment is None:
            alignment = _first_style_value(
                paragraph.style or representative_style,
                lambda style: style.paragraph_format.alignment,
            )
        if alignment is None:
            alignment = normal_style.paragraph_format.alignment

        first_line_indents.append(_round_cm(indent))
        left_indents.append(_round_cm(left_indent))
        right_indents.append(_round_cm(right_indent))
        spaces_before.append(_round_pt(space_before))
        spaces_after.append(_round_pt(space_after))
        if keep_with_next is not None:
            keep_with_next_values.append(bool(keep_with_next))
        if keep_together is not None:
            keep_together_values.append(bool(keep_together))
        if alignment is not None:
            alignments.append(str(alignment).split()[0].casefold())

        reference_size = (
            _style_font_size(paragraph.style or representative_style)
            or _dominant(font_sizes)
            or _style_font_size(normal_style)
            or _document_default_font_size(document)
        )
        if hasattr(spacing, "pt") and reference_size:
            ratio = round(float(spacing.pt) / reference_size, 2)
            if 0.5 <= ratio <= 4:
                line_spacings.append(ratio)
        elif isinstance(spacing, (int, float)):
            ratio = round(float(spacing), 2)
            if 0.5 <= ratio <= 4:
                line_spacings.append(ratio)

    style_font_name = _first_style_value(
        representative_style,
        lambda style: style.font.name,
    )
    result = {
        "font_family": (
            _dominant(font_names)
            or style_font_name
            or normal_style.font.name
            or _document_default_font_family(document)
            or ""
        ),
        "font_size_pt": (
            style_font_size
            or _dominant(font_sizes)
            or _style_font_size(normal_style)
            or _document_default_font_size(document)
        ),
        "line_spacing": _dominant(line_spacings),
        "first_line_indent_cm": _dominant(first_line_indents),
        "left_indent_cm": _dominant(left_indents),
        "right_indent_cm": _dominant(right_indents),
        "space_before_pt": _dominant(spaces_before),
        "space_after_pt": _dominant(spaces_after),
        "alignment": _dominant(alignments) or "",
    }
    bold = _first_style_value(
        representative_style,
        lambda style: style.font.bold,
    )
    italic = _first_style_value(
        representative_style,
        lambda style: style.font.italic,
    )
    if bold is not None:
        result["bold"] = bool(bold)
    if italic is not None:
        result["italic"] = bool(italic)
    if keep_with_next_values:
        result["keep_with_next"] = _dominant(keep_with_next_values)
    if keep_together_values:
        result["keep_together"] = _dominant(keep_together_values)
    return {
        key: value
        for key, value in result.items()
        if value not in (None, "")
    }


def _style_definition_rules(
    document,
    paragraphs,
    *,
    include_color=False,
    automatic_color_hex="",
):
    """
    Resolve formatting from a representative paragraph style definition.

    Placeholder text in a DOCX template often carries direct formatting which
    is not part of the reusable style.  Font family and size may legitimately
    inherit from Normal/document defaults, but paragraph layout must come only
    from the custom style chain; otherwise Normal or a formatted placeholder
    can silently change captions, headings, references, and front matter.
    """

    if not paragraphs:
        return {}

    normal_style = document.styles["Normal"]
    style_id = _dominant(
        paragraph.style.style_id
        for paragraph in paragraphs
        if paragraph.style is not None
    )
    representative_style = next(
        (
            paragraph.style
            for paragraph in paragraphs
            if paragraph.style is not None
            and paragraph.style.style_id == style_id
        ),
        normal_style,
    )

    def custom_style_value(getter):
        for candidate in _style_chain(representative_style):
            if candidate.style_id == normal_style.style_id:
                break
            value = getter(candidate)
            if value is not None:
                return value
        return None

    font_family = (
        _first_style_value(
            representative_style,
            lambda style: style.font.name,
        )
        or normal_style.font.name
        or _document_default_font_family(document)
        or ""
    )
    font_size = (
        _style_font_size(representative_style)
        or _style_font_size(normal_style)
        or _document_default_font_size(document)
    )
    spacing = custom_style_value(
        lambda style: style.paragraph_format.line_spacing
    )
    line_spacing = None
    if hasattr(spacing, "pt") and font_size:
        ratio = round(float(spacing.pt) / font_size, 2)
        if 0.5 <= ratio <= 4:
            line_spacing = ratio
    elif isinstance(spacing, (int, float)):
        ratio = round(float(spacing), 2)
        if 0.5 <= ratio <= 4:
            line_spacing = ratio

    alignment = custom_style_value(
        lambda style: style.paragraph_format.alignment
    )
    result = {
        "font_family": font_family,
        "font_size_pt": font_size,
        "line_spacing": line_spacing,
        "first_line_indent_cm": _round_cm(
            custom_style_value(
                lambda style: style.paragraph_format.first_line_indent
            )
        ),
        "left_indent_cm": _round_cm(
            custom_style_value(
                lambda style: style.paragraph_format.left_indent
            )
        ),
        "right_indent_cm": _round_cm(
            custom_style_value(
                lambda style: style.paragraph_format.right_indent
            )
        ),
        "space_before_pt": _round_pt(
            custom_style_value(
                lambda style: style.paragraph_format.space_before
            )
        ),
        "space_after_pt": _round_pt(
            custom_style_value(
                lambda style: style.paragraph_format.space_after
            )
        ),
        "alignment": (
            str(alignment).split()[0].casefold()
            if alignment is not None
            else "left"
        ),
    }
    if include_color:
        result["color_hex"] = _style_font_color_hex(
            representative_style,
            automatic=automatic_color_hex,
        )
    for key, getter in (
        ("bold", lambda style: style.font.bold),
        ("italic", lambda style: style.font.italic),
        (
            "keep_with_next",
            lambda style: style.paragraph_format.keep_with_next,
        ),
        (
            "keep_together",
            lambda style: style.paragraph_format.keep_together,
        ),
    ):
        value = custom_style_value(getter)
        if value is not None:
            result[key] = bool(value)
    return {
        key: value
        for key, value in result.items()
        if value not in (None, "")
    }


def _body_paragraphs(document):
    candidates = []
    for paragraph in document.paragraphs[:1000]:
        if not paragraph.text.strip():
            continue
        style_name = (paragraph.style.name or "").casefold() if paragraph.style else ""
        if any(token in style_name for token in _NON_BODY_STYLE_TOKENS):
            continue
        candidates.append(paragraph)
    if not candidates:
        candidates = [
            paragraph
            for paragraph in document.paragraphs[:1000]
            if paragraph.text.strip()
        ]
    style_id = _dominant(
        paragraph.style.style_id
        for paragraph in candidates
        if paragraph.style is not None
    )
    selected = [
        paragraph
        for paragraph in candidates
        if paragraph.style is not None and paragraph.style.style_id == style_id
    ]
    return selected or candidates


_DOCX_ROLE_STYLE_TOKENS = {
    "title": ("title", "назван"),
    "authors": ("authorname", "author_name", "author names", "автор"),
    "institution": ("affiliation", "institution", "организац"),
    "abstract": ("abstract", "аннотац"),
    "keywords": ("keyword", "ключев"),
    "references": ("reference", "литератур"),
}

_DOCX_SPECIAL_STYLE_TOKENS = {
    "figure_caption": (
        "figure_caption",
        "figure caption",
        "рисунок",
        "подпись рисун",
    ),
    "table_caption": (
        "table_caption",
        "table caption",
        "подпись табли",
    ),
    "list_itemize": ("itemize", "numbered list", "нумерован"),
    "list_bullet": ("bullet", "bulleted list", "маркирован"),
    "table_body": ("table_body", "table body", "текст табли"),
}


def _paragraphs_with_style_tokens(document, tokens, *, limit=1000):
    normalized_tokens = tuple(str(token).casefold() for token in tokens)
    return [
        paragraph
        for paragraph in document.paragraphs[:limit]
        if paragraph.text.strip()
        and paragraph.style is not None
        and any(
            token in (paragraph.style.name or "").casefold()
            for token in normalized_tokens
        )
    ]


def _table_paragraphs_with_style_tokens(document, tokens):
    normalized_tokens = tuple(str(token).casefold() for token in tokens)
    paragraphs = []
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if (
                        paragraph.text.strip()
                        and paragraph.style is not None
                        and any(
                            token in (paragraph.style.name or "").casefold()
                            for token in normalized_tokens
                        )
                    ):
                        paragraphs.append(paragraph)
    return paragraphs


def _docx_heading_style_level(paragraph):
    style_name = (
        (paragraph.style.name or "").casefold()
        if paragraph.style is not None
        else ""
    )
    match = re.search(r"(?:heading|заголов\w*)[\s_.-]*([1-6])\b", style_name)
    return int(match.group(1)) if match else None


def _semantic_docx_blocks(document, body_rules):
    blocks = []
    for role in (
        "title",
        "authors",
        "institution",
        "abstract",
        "keywords",
        "body",
        "references",
    ):
        if role == "body":
            style = dict(body_rules or {})
            # Body styles describe the paragraph defaults.  Inline bold and
            # italic emphasis in the scientific text must remain untouched.
            style.pop("bold", None)
            style.pop("italic", None)
        else:
            paragraphs = _paragraphs_with_style_tokens(
                document,
                _DOCX_ROLE_STYLE_TOKENS[role],
            )
            if not paragraphs:
                continue
            style = _style_definition_rules(document, paragraphs)
            # An absent Word alignment means the default left alignment.  It
            # must remain explicit here so generic block defaults do not turn
            # an intentionally left-aligned publisher title into a centered
            # one.
            style.setdefault("alignment", "left")
            if role in {
                "title",
                "authors",
                "institution",
                "abstract",
                "keywords",
            }:
                style.setdefault("first_line_indent_cm", 0)
        if not style:
            continue
        blocks.append(
            {
                "role": role,
                "required": role == "body",
                "style": style,
            }
        )
    return blocks


def _special_docx_paragraph_styles(document):
    result = {}
    for role, tokens in _DOCX_SPECIAL_STYLE_TOKENS.items():
        paragraphs = (
            _table_paragraphs_with_style_tokens(document, tokens)
            if role == "table_body"
            else _paragraphs_with_style_tokens(document, tokens)
        )
        if not paragraphs:
            continue
        style = _style_definition_rules(document, paragraphs)
        style.setdefault("alignment", "left")
        if role in {"figure_caption", "table_caption"}:
            style.setdefault("first_line_indent_cm", 0)
        result[role] = style
    return result


def _pdf_font_family(base_font):
    name = str(base_font or "").split("+")[-1].casefold()
    if "palladio" in name or "palatino" in name:
        return "Palatino Linotype"
    if "times" in name or "nimbusroman" in name:
        return "Times New Roman"
    if "helvetica" in name or "arial" in name:
        return "Arial"
    if "computer-modern" in name or "latinmodern" in name:
        return "Latin Modern Roman"
    return ""


def _percentile(values, fraction):
    if not values:
        return None
    ordered = sorted(values)
    index = min(len(ordered) - 1, max(0, round((len(ordered) - 1) * fraction)))
    return ordered[index]


def _extract_pdf_rules(data):
    """Extract visual page and body rules which plain PDF text cannot retain."""
    try:
        from pypdf import PdfReader
        from pypdf._text_extraction import mult
    except ImportError:
        return {}

    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception:
        return {}
    if not reader.pages:
        return {}

    first_page = reader.pages[0]
    width_pt = float(first_page.mediabox.width)
    height_pt = float(first_page.mediabox.height)
    fragments = []
    # The title page is often deliberately different. Later pages provide a
    # much more reliable sample of the article's main typography.
    page_indexes = range(1, min(len(reader.pages), 16)) if len(reader.pages) > 1 else range(1)
    for page_index in page_indexes:
        page = reader.pages[page_index]

        def collect(text, cm, tm, font_dict, font_size):
            clean = " ".join(str(text or "").split())
            if not clean or font_dict is None or not (7 <= float(font_size or 0) <= 13):
                return
            x, y = mult(tm, cm)[4:]
            if x <= 0 or y <= 35 or y >= height_pt - 25:
                return
            if re.fullmatch(r"\d+", clean) and float(font_size) < 7:
                return
            fragments.append(
                {
                    "page": page_index,
                    "text": clean,
                    "x": float(x),
                    "y": float(y),
                    "size": round(float(font_size), 2),
                    "base_font": str(font_dict.get("/BaseFont") or ""),
                }
            )

        try:
            page.extract_text(visitor_text=collect)
        except Exception:
            continue

    page_width_cm = width_pt / 72 * 2.54
    page_height_cm = height_pt / 72 * 2.54
    is_a4 = (
        abs(min(page_width_cm, page_height_cm) - 21.0) <= 0.35
        and abs(max(page_width_cm, page_height_cm) - 29.7) <= 0.35
    )
    page_rules = {
        "size": "A4" if is_a4 else "",
        "orientation": "landscape" if width_pt > height_pt else "portrait",
    }
    if not fragments:
        return {"page": page_rules}

    font_weights = Counter()
    for item in fragments:
        if 8 <= item["size"] <= 12:
            font_weights[(item["base_font"], item["size"])] += len(item["text"])
    if not font_weights:
        return {"page": page_rules}
    (body_base_font, body_size), _weight = font_weights.most_common(1)[0]
    body_fragments = [
        item
        for item in fragments
        if item["base_font"] == body_base_font and abs(item["size"] - body_size) <= 0.15
    ]

    x_counts = Counter(round(item["x"]) for item in body_fragments)
    body_left_pt = float(x_counts.most_common(1)[0][0])
    indent_candidates = Counter(
        round(item["x"]) - round(body_left_pt)
        for item in body_fragments
        if 10 <= round(item["x"]) - round(body_left_pt) <= 40
    )
    first_line_indent_pt = (
        float(indent_candidates.most_common(1)[0][0])
        if indent_candidates
        else 0.0
    )

    baselines = {}
    for item in body_fragments:
        baselines.setdefault(item["page"], set()).add(round(item["y"], 1))
    line_gaps = []
    for page_values in baselines.values():
        ordered = sorted(page_values, reverse=True)
        for first, second in zip(ordered, ordered[1:]):
            gap = first - second
            if body_size * 1.05 <= gap <= body_size * 1.8:
                line_gaps.append(gap)
    baseline_gap = median(line_gaps) if line_gaps else body_size
    line_spacing = round(baseline_gap / body_size, 2) if body_size else None

    page_tops = []
    page_bottoms = []
    for page_index in set(item["page"] for item in body_fragments):
        ys = [item["y"] for item in body_fragments if item["page"] == page_index]
        if ys:
            page_tops.append(max(ys))
            page_bottoms.append(min(ys))
    top_baseline = _percentile(page_tops, 0.75)
    bottom_baseline = _percentile(page_bottoms, 0.25)
    margins = {
        "left": round(body_left_pt / 72 * 2.54, 2),
        # PDF text coordinates do not expose a dependable right edge for every
        # embedded font. A narrow journal margin is inferred from the available
        # line width after the dominant left boundary.
        "right": round(max(0.8, min(2.5, (width_pt - body_left_pt) / 72 * 2.54 * 0.08)), 2),
        "top": round((height_pt - top_baseline + body_size * 0.2) / 72 * 2.54, 2),
        "bottom": round(max(1.5, bottom_baseline / 72 * 2.54), 2),
    }
    page_rules["margins_cm"] = margins
    bold_sizes = Counter()
    for item in fragments:
        if "bold" in item["base_font"].casefold() and item["size"] >= body_size:
            bold_sizes[item["size"]] += len(item["text"])
    title_size = max(
        max(bold_sizes, default=0),
        round(body_size * 1.8, 1),
    )

    return {
        "page": page_rules,
        "body": {
            "font_family": _pdf_font_family(body_base_font),
            "font_size_pt": round(body_size, 1),
            "line_spacing": line_spacing,
            "first_line_indent_cm": round(first_line_indent_pt / 72 * 2.54, 2),
            "alignment": "justify",
        },
        "headings": {
            "font_family": _pdf_font_family(body_base_font),
            "font_size_pt": round(body_size, 1),
            "title_font_size_pt": round(title_size, 1),
            "color_hex": "000000",
        },
    }


def _classify_template_source(file_name, text):
    suffix = Path(file_name or "").suffix.casefold()
    if suffix != ".pdf":
        return {"kind": "normative_template", "confidence": 1.0, "reasons": []}

    source = str(text or "").casefold().replace("ё", "е")
    name = Path(file_name or "").name.casefold()
    manuscript_signals = []
    if re.search(r"(?:peer[-_ ]?review|proof|manuscript)", name):
        manuscript_signals.append("filename")
    if re.search(r"\bsubmitted\s+to\b|\bversion\b.{0,80}\bsubmitted\s+to\b", source):
        manuscript_signals.append("submission_footer")
    if re.search(r"\babstract\b", source) and re.search(r"\bkeywords?\b", source):
        manuscript_signals.append("article_front_matter")
    if re.search(r"\b(?:references|bibliography)\b", source):
        manuscript_signals.append("references")
    if re.search(r"https?://doi\.org/10\.|doi\s*:", source):
        manuscript_signals.append("doi")
    numbered_sections = len(
        re.findall(
            r"(?:^|\n)\s*\d+(?:\.\d+)*\.?\s+[A-ZА-ЯЁ][^\n]{2,80}",
            str(text or ""),
        )
    )
    if numbered_sections >= 3:
        manuscript_signals.append("numbered_article_sections")

    instruction_signals = []
    for pattern, label in (
        (
            r"\b(?:instructions?\s+for\s+authors?|author\s+guidelines?|"
            r"manuscript\s+preparation|use\s+this\s+template)\b",
            "author_instructions",
        ),
        (
            r"\b(?:требовани\w*\s+к\s+оформлени\w*|инструкци\w*\s+для\s+автор\w*|"
            r"используйте\s+(?:этот\s+)?шаблон)\b",
            "author_instructions_ru",
        ),
    ):
        if re.search(pattern, source):
            instruction_signals.append(label)

    if len(manuscript_signals) >= 3 and not instruction_signals:
        return {
            "kind": "sample_manuscript",
            "confidence": min(0.99, 0.7 + len(manuscript_signals) * 0.05),
            "reasons": manuscript_signals,
        }
    return {
        "kind": "normative_template",
        "confidence": 0.75 if manuscript_signals else 0.9,
        "reasons": instruction_signals,
    }


def _sample_manuscript_rules(deterministic_rules, classification, text):
    page = deterministic_rules.get("page") or {}
    safe_page = {
        key: page.get(key)
        for key in ("size", "orientation")
        if page.get(key) not in (None, "")
    }
    letters = [character for character in str(text or "") if character.isalpha()]
    cyrillic = sum("а" <= character.casefold() <= "я" or character in "Ёё" for character in letters)
    languages = ["ru"] if letters and cyrillic / len(letters) >= 0.35 else ["en"]
    return {
        "page": safe_page,
        "document": {
            "source_kind": "sample_manuscript",
            "rules_reliable": False,
            "latex_generation_allowed": False,
            "source_notice": (
                "Загруженный PDF похож на готовую свёрстанную статью, а не на "
                "нормативный шаблон. Разделы, поля и шрифты конкретной статьи "
                "не применяются как обязательные требования."
            ),
        },
        "structure": {"required_sections": [], "section_order": []},
        "metadata": {"required_fields": []},
        "languages": languages,
        "notes": [
            "Для точного LaTeX используйте официальный TEX/ZIP-шаблон издателя."
        ],
        "source_classification": classification,
    }


def _publisher_sample_profile(file_name, text, deterministic_rules=None):
    """
    Return a conservative formatting profile for a recognised publisher sample.

    A published article is not a normative template: its text column and section
    names must not become mandatory rules.  Some publisher samples can still be
    used safely for typography when the publisher and journal are unambiguous.
    """

    name = Path(file_name or "").name.casefold()
    source = str(text or "").casefold()
    mdpi_sensors_name = bool(
        re.search(r"(?:^|[/\\])?sensors-\d+-peer-review-v\d+\.pdf$", name)
    )
    mdpi_sensors_text = (
        "mdpi" in source
        and (
            re.search(r"\bsensors\b", source)
            or "mdpi.com/journal/sensors" in source
        )
    )
    if not (mdpi_sensors_name or mdpi_sensors_text):
        return {}

    body_rules = {
        "font_family": "Palatino Linotype",
        "font_size_pt": 10,
        "line_spacing": 1,
        "alignment": "justify",
    }
    detected_body = (
        deterministic_rules.get("body") or {}
        if isinstance(deterministic_rules, dict)
        else {}
    )
    for key in (
        "line_spacing",
        "first_line_indent_cm",
        "alignment",
    ):
        value = detected_body.get(key)
        if value not in (None, ""):
            body_rules[key] = value

    return {
        "page": {"size": "A4", "orientation": "portrait"},
        # Text-column margins from a typeset article are intentionally ignored,
        # but paragraph typography is safe to reproduce in an editable DOCX.
        "body": body_rules,
        "headings": {
            "font_family": "Palatino Linotype",
            "font_size_pt": 10,
            "title_font_size_pt": 18,
            "color_hex": "000000",
        },
        "document": {
            "source_kind": "publisher_sample",
            "publisher_profile": "mdpi_sensors",
            "auto_format_mode": "safe_typography",
            "rules_reliable": True,
            # A standalone generic LaTeX file cannot replace the official MDPI
            # class and Definitions directory.  DOCX formatting is available.
            "latex_generation_allowed": False,
            "source_notice": (
                "PDF распознан как образец Sensors/MDPI. Автоматическое "
                "оформление DOCX включено в безопасном режиме: применяются "
                "типографика и формат страницы, а поля колонки и заголовки "
                "конкретной статьи не становятся обязательными."
            ),
        },
        "structure": {"required_sections": [], "section_order": []},
        "metadata": {"required_fields": []},
        "languages": ["en"],
        "notes": [
            "Поля журнальной колонки и боковая издательская панель не копируются.",
            "Для LaTeX требуется официальный MDPI TEX/ZIP с каталогом Definitions.",
        ],
        "source_classification": {
            "kind": "publisher_sample",
            "confidence": 0.99,
            "publisher_profile": "mdpi_sensors",
        },
    }


def _extract_docx_rules(data):
    try:
        from docx import Document
    except ImportError:
        return {}

    source = io.BytesIO(data)
    try:
        document = Document(source)
    except ValueError:
        source = io.BytesIO()
        try:
            with ZipFile(io.BytesIO(data)) as input_archive, ZipFile(
                source,
                "w",
                ZIP_DEFLATED,
            ) as output_archive:
                for info in input_archive.infolist():
                    value = input_archive.read(info)
                    if info.filename == "[Content_Types].xml":
                        value = value.replace(
                            b"application/vnd.openxmlformats-officedocument."
                            b"wordprocessingml.template.main+xml",
                            b"application/vnd.openxmlformats-officedocument."
                            b"wordprocessingml.document.main+xml",
                        )
                    output_archive.writestr(info, value)
        except (OSError, BadZipFile, KeyError):
            return {}
        source.seek(0)
        try:
            document = Document(source)
        except (OSError, ValueError):
            return {}
    page_rules = {}
    if document.sections:
        section = document.sections[0]
        width_cm = _round_cm(section.page_width)
        height_cm = _round_cm(section.page_height)
        page_rules = {
            "size": "A4" if width_cm and height_cm and sorted((round(width_cm), round(height_cm))) == [21, 30] else "",
            "orientation": "landscape" if width_cm and height_cm and width_cm > height_cm else "portrait",
            "margins_cm": {
                "top": _round_cm(section.top_margin),
                "right": _round_cm(section.right_margin),
                "bottom": _round_cm(section.bottom_margin),
                "left": _round_cm(section.left_margin),
            },
        }

    body_paragraphs = _body_paragraphs(document)
    normal_style_id = document.styles["Normal"].style_id
    body_style_id = _dominant(
        paragraph.style.style_id
        for paragraph in body_paragraphs
        if paragraph.style is not None
    )
    body_rules = (
        _style_definition_rules(document, body_paragraphs)
        if body_style_id and body_style_id != normal_style_id
        else _resolved_style_rules(document, body_paragraphs)
    )
    heading_paragraphs = [
        paragraph
        for paragraph in document.paragraphs[:1000]
        if paragraph.text.strip()
        and _docx_heading_style_level(paragraph) is not None
    ]
    heading_levels = {}
    for level in range(1, 7):
        level_paragraphs = [
            paragraph
            for paragraph in heading_paragraphs
            if _docx_heading_style_level(paragraph) == level
        ]
        if not level_paragraphs:
            continue
        level_rules = _style_definition_rules(
            document,
            level_paragraphs,
            include_color=True,
            automatic_color_hex="000000",
        )
        level_rules.setdefault("alignment", "left")
        level_rules.setdefault("first_line_indent_cm", 0)
        level_rules.setdefault("bold", False)
        level_rules.setdefault("italic", False)
        heading_levels[str(level)] = level_rules
    heading_rules = (
        dict(heading_levels.get("1") or {})
        or _style_definition_rules(
            document,
            heading_paragraphs,
            include_color=True,
            automatic_color_hex="000000",
        )
    )
    title_paragraphs = [
        paragraph
        for paragraph in document.paragraphs[:100]
        if paragraph.text.strip()
        and paragraph.style is not None
        and (
            "title" in (paragraph.style.name or "").casefold()
            or "назван" in (paragraph.style.name or "").casefold()
        )
    ]
    title_rules = _style_definition_rules(document, title_paragraphs)
    if title_rules.get("font_size_pt"):
        heading_rules["title_font_size_pt"] = title_rules["font_size_pt"]
    if heading_levels:
        heading_rules["levels"] = heading_levels

    result = {
        "page": page_rules,
        "body": body_rules,
        "document": {
            "analysis_revision": FORMATTING_TEMPLATE_ANALYSIS_REVISION,
            "blocks": _semantic_docx_blocks(document, body_rules),
            "paragraph_styles": _special_docx_paragraph_styles(document),
        },
    }
    if heading_paragraphs:
        result["headings"] = heading_rules
    return result


def _extract_template_content(template):
    with template.file.open("rb") as source:
        data = read_file_bytes(source)
    suffix = Path(template.file.name).suffix.casefold()
    analysis_data = data
    analysis_name = template.file.name
    if suffix == ".doc":
        analysis_data = convert_legacy_doc_to_docx(data)
        analysis_name = f"{Path(template.file.name).stem}.docx"
    snapshot = analyze_document_bytes(analysis_data, analysis_name)
    text = snapshot.get("text") or ""
    parse_warning = snapshot.get("parse_error") or ""
    if suffix in {".doc", ".docx", ".dotx"}:
        deterministic_rules = _extract_docx_rules(analysis_data)
    elif suffix == ".pdf":
        deterministic_rules = _extract_pdf_rules(data)
    elif suffix == ".tex":
        deterministic_rules = extract_latex_template_rules(data)
    else:
        deterministic_rules = {}

    if suffix == ".pdf":
        try:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(data))
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
            if text.strip():
                parse_warning = ""
        except Exception:
            text = ""
    elif suffix == ".tex":
        for encoding in ("utf-8-sig", "utf-8", "cp1251", "utf-16le"):
            try:
                text = data.decode(encoding)
            except UnicodeDecodeError:
                continue
            else:
                break
        else:
            text = data.decode("utf-8", errors="replace")
        parse_warning = ""
    elif suffix in TEXT_EXTENSIONS:
        text = snapshot.get("text") or ""
    return text[:120_000], deterministic_rules, parse_warning


def _merge_dict(base, override):
    result = dict(base or {})
    for key, value in (override or {}).items():
        if (
            key == "blocks"
            and isinstance(value, list)
            and isinstance(result.get(key), list)
        ):
            merged_blocks = []
            positions = {}
            for block in [*result[key], *value]:
                if not isinstance(block, dict):
                    continue
                role = str(block.get("role") or "").strip().casefold()
                if not role:
                    merged_blocks.append(block)
                    continue
                if role in positions:
                    position = positions[role]
                    merged_blocks[position] = _merge_dict(
                        merged_blocks[position],
                        block,
                    )
                else:
                    positions[role] = len(merged_blocks)
                    merged_blocks.append(block)
            result[key] = merged_blocks
        elif isinstance(value, dict) and isinstance(result.get(key), dict):
            result[key] = _merge_dict(result[key], value)
        elif value not in (None, "", [], {}):
            result[key] = value
    return result


def _extract_rules_with_ai(template, text):
    if not text.strip() or not is_ai_configured():
        return {}
    if get_provider() != "openai_compatible":
        raise ValueError(
            "Для интерпретации шаблонов требуется локальная OpenAI-совместимая модель."
        )

    def complete_json(prompt):
        payload = {
            "contents": [{"role": "user", "parts": [{"text": prompt}]}],
            "generationConfig": {
                "temperature": 0.1,
                "maxOutputTokens": 6144,
                "responseMimeType": "application/json",
            },
        }
        response, _model = generate_content(
            payload,
            model=get_configured_model(),
            timeout=120,
        )
        return extract_response_text(response)

    return interpret_template_text(
        document_type=template.article_type.name,
        target_name=template.target_name,
        text=text,
        complete_json=complete_json,
    )


def build_rules_snapshot(*, article_type, template=None, journal=None):
    effective = _merge_dict({}, DEFAULT_RULES.get(article_type.code, {}))
    sources = [{"kind": "material_type", "label": article_type.name, "priority": 10}]
    conflicts = []

    policy = getattr(journal, "editorial_policy", {}) if journal is not None else {}
    if isinstance(policy, dict) and policy:
        journal_rules = {
            "structure": {"required_sections": policy.get("required_sections") or []},
            "limits": {
                "min_words": policy.get("min_words"),
                "max_words": policy.get("max_words"),
            },
        }
        effective = _merge_dict(effective, journal_rules)
        sources.append({"kind": "journal", "label": journal.name, "priority": 20})

    if template is not None and template.extracted_rules:
        old_limits = dict(effective.get("limits") or {})
        effective = _merge_dict(effective, template.extracted_rules)
        new_limits = effective.get("limits") or {}
        for key in ("min_words", "max_words"):
            if old_limits.get(key) and new_limits.get(key) and old_limits[key] != new_limits[key]:
                conflicts.append(
                    {
                        "field": f"limits.{key}",
                        "lower_value": old_limits[key],
                        "selected_value": new_limits[key],
                        "message": "Требование шаблона имеет приоритет над общим правилом типа материала.",
                    }
                )
        sources.append(
            {
                "kind": "uploaded_template",
                "label": f"{template.target_name}, шаблон v{template.version_number}",
                "priority": 30,
                "template_id": template.id,
            }
        )
    return {
        "effective": normalize_template_rules(effective),
        "sources": sources,
        "conflicts": conflicts,
    }


def has_manual_rule_overrides(snapshot):
    snapshot = snapshot or {}
    effective = snapshot.get("effective") or {}
    document_rules = effective.get("document") or {}
    return bool(
        document_rules.get("manual_override_confirmed")
        or any(
            source.get("kind") == "manual"
            for source in (snapshot.get("sources") or [])
            if isinstance(source, dict)
        )
    )


@transaction.atomic
def create_formatting_template(
    *,
    article_type,
    uploaded_by,
    file,
    journal=None,
    publication_topic=None,
):
    if (journal is None) == (publication_topic is None):
        raise ValueError("Шаблон должен относиться либо к журналу, либо к теме/событию.")
    suffix = Path(file.name or "").suffix.casefold()
    if suffix not in TEMPLATE_EXTENSIONS:
        allowed = ", ".join(sorted(value.lstrip(".").upper() for value in TEMPLATE_EXTENSIONS))
        raise ValueError(f"Формат шаблона не поддерживается. Разрешены: {allowed}.")

    filters = {"article_type": article_type}
    if journal is not None:
        filters["journal"] = journal
    else:
        filters["publication_topic"] = publication_topic
    last_version = (
        FormattingTemplate.objects.select_for_update()
        .filter(**filters)
        .aggregate(value=models.Max("version_number"))["value"]
        or 0
    )
    return FormattingTemplate.objects.create(
        article_type=article_type,
        journal=journal,
        publication_topic=publication_topic,
        version_number=last_version + 1,
        file=file,
        uploaded_by=uploaded_by,
    )


def formatting_template_needs_processing(template):
    if (
        template.analysis_status not in {"ready", "partial"}
        or not template.extracted_rules
    ):
        return True
    suffix = Path(template.file.name or "").suffix.casefold()
    if suffix not in {".docx", ".dotx"}:
        return False
    document_rules = (
        template.extracted_rules.get("document")
        if isinstance(template.extracted_rules, dict)
        else {}
    ) or {}
    return (
        document_rules.get("analysis_revision")
        != FORMATTING_TEMPLATE_ANALYSIS_REVISION
    )


def process_formatting_template(template):
    template.analysis_status = FormattingTemplateStatus.PROCESSING
    template.analysis_message = "Извлекаем правила из шаблона."
    template.save(update_fields=["analysis_status", "analysis_message"])
    try:
        text, deterministic_rules, parse_warning = _extract_template_content(template)
        classification = _classify_template_source(template.file.name, text)
        if classification["kind"] == "sample_manuscript":
            publisher_profile = _publisher_sample_profile(
                template.file.name,
                text,
                deterministic_rules,
            )
            template.source_text = text
            template.extracted_rules = normalize_template_rules(
                publisher_profile
                or _sample_manuscript_rules(
                    deterministic_rules,
                    classification,
                    text,
                )
            )
            template.rule_conflicts = []
            template.analysis_status = FormattingTemplateStatus.PARTIAL
            if publisher_profile:
                template.analysis_message = (
                    "PDF распознан как образец Sensors/MDPI. Автоматическое "
                    "оформление DOCX включено в безопасном режиме без ложных "
                    "полей и обязательных заголовков конкретной статьи."
                )
            else:
                template.analysis_message = (
                    "PDF распознан как готовая свёрстанная статья, а не как "
                    "нормативный шаблон. Автоматическое применение полей, шрифтов "
                    "и разделов отключено. Загрузите официальный TEX/ZIP-шаблон "
                    "или текст требований издателя."
                )
            template.save(
                update_fields=[
                    "analysis_status",
                    "analysis_message",
                    "source_text",
                    "extracted_rules",
                    "rule_conflicts",
                ]
            )
            return template
        ai_rules = {}
        ai_warning = ""
        try:
            ai_rules = _extract_rules_with_ai(template, text)
        except (AIProviderError, ValueError) as exc:
            ai_warning = str(exc)
        raw_rules = _merge_dict(ai_rules, deterministic_rules)
        extracted_rules = normalize_template_rules(raw_rules) if raw_rules else {}
        template.source_text = text
        template.extracted_rules = extracted_rules
        template.rule_conflicts = []
        if extracted_rules:
            template.analysis_status = (
                FormattingTemplateStatus.READY
                if text and not parse_warning and not ai_warning
                else FormattingTemplateStatus.PARTIAL
            )
            warnings = [value for value in (parse_warning, ai_warning) if value]
            template.analysis_message = (
                "Правила извлечены."
                if not warnings
                else "Правила извлечены частично. " + " ".join(warnings)
            )
        else:
            template.analysis_status = FormattingTemplateStatus.PARTIAL
            template.analysis_message = (
                "Файл сохранён, но автоматически извлечь правила не удалось. "
                "Правила можно уточнить вручную для конкретной работы."
            )
    except Exception as exc:
        template.analysis_status = FormattingTemplateStatus.FAILED
        template.analysis_message = f"Не удалось обработать шаблон: {type(exc).__name__}."
    template.save(
        update_fields=[
            "analysis_status",
            "analysis_message",
            "source_text",
            "extracted_rules",
            "rule_conflicts",
        ]
    )
    return template


def get_latest_formatting_template(*, article_type, journal=None, publication_topic=None):
    queryset = FormattingTemplate.objects.filter(article_type=article_type)
    if journal is not None:
        queryset = queryset.filter(journal=journal)
    elif publication_topic is not None:
        queryset = queryset.filter(publication_topic=publication_topic)
    else:
        return None
    return queryset.select_related("uploaded_by", "article_type", "journal", "publication_topic").order_by(
        "-version_number",
        "-created_at",
    ).first()
