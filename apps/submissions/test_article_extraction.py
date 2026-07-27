from io import BytesIO
from unittest.mock import patch
from zipfile import ZIP_DEFLATED, ZipFile

from django.test import SimpleTestCase

from apps.submissions.article_extraction import (
    extract_article_structure,
    refine_article_with_model,
)
from apps.directory.formatting_templates import _extract_docx_rules
from apps.submissions.document_conversion import LegacyDocConversionError
from apps.submissions.document_analysis import analyze_document_bytes


def _badly_formatted_article_docx():
    from docx import Document
    from docx.shared import Pt

    document = Document()
    title = document.add_paragraph("Надёжное извлечение данных из плохо оформленных статей")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(18)
    document.add_paragraph("Петров П.П.; Сидорова А.В.")
    document.add_paragraph("Тамбовский государственный технический университет")
    document.add_paragraph("Аннотация")
    document.add_paragraph(
        "Предложен устойчивый конвейер анализа документов, не зависящий от "
        "единственного стиля Word и регистра заголовков."
    )
    document.add_paragraph("Ключевые слова — документы; извлечение; структура")
    introduction = document.add_paragraph("Введение")
    introduction.runs[0].bold = True
    document.add_paragraph(
        "Основной текст статьи расположен после метаданных и отдельного заголовка."
    )
    references = document.add_paragraph("Литература")
    references.runs[0].bold = True
    document.add_paragraph("1. Петров П. П. Анализ документов. Тамбов, 2025.")
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _article_with_uppercase_surnames_docx():
    from docx import Document
    from docx.shared import Pt

    document = Document()
    title = document.add_paragraph(
        "АЛГОРИТМ АДАПТАЦИИ ВИРТУАЛЬНОЙ СРЕДЫ ДЛЯ РЕАБИЛИТАЦИИ"
    )
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(16)
    document.add_paragraph("ВОЛКОВ Андрей Андреевич, аспирант,")
    document.add_paragraph(
        "ассистент кафедры «Системы автоматизированной поддержки принятия решений»,"
    )
    document.add_paragraph("Тамбовский государственный технический университет")
    document.add_paragraph("ОБУХОВ Артем Дмитриевич, доктор технических наук,")
    document.add_paragraph(
        "профессор кафедры «Системы автоматизированной поддержки принятия решений»,"
    )
    document.add_paragraph("Тамбовский государственный технический университет")
    document.add_paragraph(
        "Аннотация. Рассматривается автоматическая адаптация виртуальной среды."
    )
    document.add_paragraph(
        "Ключевые слова: реабилитация, адаптивное управление."
    )
    introduction = document.add_paragraph("Введение")
    introduction.runs[0].bold = True
    document.add_paragraph("Основной текст научной статьи.")
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _docx_with_formula_and_figure():
    document_xml = """<?xml version="1.0" encoding="UTF-8"?>
    <w:document
      xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
      xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"
      xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
      xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
      <w:body>
        <w:p><w:r><w:t>FORMULA AND FIGURE EXTRACTION</w:t></w:r></w:p>
        <w:p><m:oMath><m:r><m:t>x=1</m:t></m:r></m:oMath></w:p>
        <w:p><w:r><w:drawing><a:blip r:embed="rId1"/></w:drawing></w:r></w:p>
      </w:body>
    </w:document>
    """
    relationships = """<?xml version="1.0" encoding="UTF-8"?>
    <Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
      <Relationship Id="rId1"
        Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"
        Target="media/image1.png"/>
    </Relationships>
    """
    output = BytesIO()
    with ZipFile(output, "w", ZIP_DEFLATED) as archive:
        archive.writestr("word/document.xml", document_xml)
        archive.writestr("word/_rels/document.xml.rels", relationships)
        archive.writestr("word/media/image1.png", b"not-decoded-by-parser")
    return output.getvalue()


def _semantic_style_article_docx():
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE

    document = Document()
    for style_name in (
        "MDPI_1.2_title",
        "MDPI_1.3_authornames",
        "MDPI_1.4_affiliation",
        "MDPI_2.1_abstract",
        "MDPI_2.2_keywords",
    ):
        document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    document.add_paragraph(
        "A Resilient Parser for Scientific Articles",
        style="MDPI_1.2_title",
    )
    document.add_paragraph(
        "Alice Researcher; Bob Scientist",
        style="MDPI_1.3_authornames",
    )
    document.add_paragraph(
        "Document Intelligence Laboratory, Example University",
        style="MDPI_1.4_affiliation",
    )
    document.add_paragraph(
        "Abstract: Layout and semantic styles are combined.",
        style="MDPI_2.1_abstract",
    )
    document.add_paragraph(
        "Keywords: document parsing; semantic styles",
        style="MDPI_2.2_keywords",
    )
    document.add_paragraph("1. Introduction", style="Heading 1")
    document.add_paragraph("The article body.")
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _as_dotx(data):
    output = BytesIO()
    with ZipFile(BytesIO(data)) as source, ZipFile(
        output,
        "w",
        ZIP_DEFLATED,
    ) as target:
        for info in source.infolist():
            value = source.read(info)
            if info.filename == "[Content_Types].xml":
                value = value.replace(
                    b"application/vnd.openxmlformats-officedocument."
                    b"wordprocessingml.document.main+xml",
                    b"application/vnd.openxmlformats-officedocument."
                    b"wordprocessingml.template.main+xml",
                )
            target.writestr(info, value)
    return output.getvalue()


def _docx_with_embedded_vector_equation():
    document_xml = """<?xml version="1.0" encoding="UTF-8"?>
    <w:document
      xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
      xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
      xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
      <w:body>
        <w:p><w:r><w:t>VECTOR EQUATION CLASSIFICATION</w:t></w:r></w:p>
        <w:p><w:r><w:drawing><a:blip r:embed="rId1"/></w:drawing></w:r></w:p>
        <w:p><w:r><w:drawing><a:blip r:embed="rId2"/></w:drawing></w:r></w:p>
      </w:body>
    </w:document>
    """
    relationships = """<?xml version="1.0" encoding="UTF-8"?>
    <Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
      <Relationship Id="rId1"
        Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"
        Target="media/equation1.wmf"/>
      <Relationship Id="rId2"
        Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"
        Target="media/figure1.png"/>
    </Relationships>
    """
    output = BytesIO()
    with ZipFile(output, "w", ZIP_DEFLATED) as archive:
        archive.writestr("word/document.xml", document_xml)
        archive.writestr("word/_rels/document.xml.rels", relationships)
        archive.writestr("word/media/equation1.wmf", b"equation-vector")
        archive.writestr("word/media/figure1.png", b"figure-raster")
    return output.getvalue()


class ArticleExtractionTests(SimpleTestCase):
    def test_utf8_russian_text_is_not_misread_as_cp1251(self):
        source = (
            "Проверка быстрой загрузки полей\n"
            "Аннотация. Технический текст для проверки интерфейса."
        ).encode("utf-8")

        snapshot = analyze_document_bytes(source, "article.txt")

        self.assertIn("Проверка быстрой загрузки полей", snapshot["text"])
        self.assertNotIn("РџСЂ", snapshot["text"])

    def test_extracts_structure_when_word_styles_and_labels_are_inconsistent(self):
        snapshot = analyze_document_bytes(
            _badly_formatted_article_docx(),
            "bad-layout.docx",
        )

        metadata = snapshot["metadata"]
        article = snapshot["article"]
        self.assertEqual(
            metadata["title"],
            "Надёжное извлечение данных из плохо оформленных статей",
        )
        self.assertEqual(metadata["authors"], ["Петров П.П.", "Сидорова А.В."])
        self.assertEqual(
            metadata["organizations"],
            "Тамбовский государственный технический университет",
        )
        self.assertIn("устойчивый конвейер", metadata["abstract"])
        self.assertEqual(metadata["keywords"], "документы, извлечение, структура")
        self.assertEqual(
            [section["title"] for section in article["sections"]],
            ["Введение", "Литература"],
        )
        self.assertEqual(len(article["references"]), 1)
        self.assertGreaterEqual(metadata["confidence"]["title"], 0.7)

    def test_preserves_uppercase_surnames_in_full_author_names(self):
        snapshot = analyze_document_bytes(
            _article_with_uppercase_surnames_docx(),
            "uppercase-surnames.docx",
        )

        self.assertEqual(
            snapshot["metadata"]["authors"],
            [
                "ВОЛКОВ Андрей Андреевич",
                "ОБУХОВ Артем Дмитриевич",
            ],
        )
        self.assertEqual(
            snapshot["metadata"]["document_authors"],
            "ВОЛКОВ Андрей Андреевич\nОБУХОВ Артем Дмитриевич",
        )

    def test_extracts_real_omml_formulas_and_drawing_relationships(self):
        snapshot = analyze_document_bytes(
            _docx_with_formula_and_figure(),
            "structured.docx",
        )

        self.assertEqual(snapshot["formulas"][0]["text"], "x=1")
        self.assertEqual(snapshot["image_count"], 1)
        self.assertEqual(snapshot["figures"][0]["relationship_id"], "rId1")

    def test_pdf_with_text_layer_is_parsed_instead_of_marked_unsupported(self):
        from reportlab.pdfgen.canvas import Canvas

        output = BytesIO()
        canvas = Canvas(output)
        canvas.drawString(72, 780, "ROBUST SCIENTIFIC DOCUMENT PARSING")
        canvas.drawString(72, 755, "A. Researcher")
        canvas.drawString(72, 730, "Introduction")
        canvas.drawString(72, 705, "The PDF text layer is available for analysis.")
        canvas.drawString(140, 680, "Figure 1. Extraction architecture.")
        canvas.drawString(160, 655, "Table 1. Evaluation results.")
        canvas.drawString(180, 630, "x = y + 1. (1)")
        canvas.save()

        snapshot = analyze_document_bytes(output.getvalue(), "article.pdf")

        self.assertEqual(snapshot["parse_error"], "")
        self.assertFalse(snapshot["requires_ocr"])
        self.assertIn("ROBUST SCIENTIFIC", snapshot["text"])
        self.assertTrue(snapshot["pages"])
        self.assertEqual(len(snapshot["figures"]), 1)
        self.assertEqual(len(snapshot["tables"]), 1)
        self.assertEqual(len(snapshot["formulas"]), 1)

    def test_local_model_cannot_inject_text_absent_from_source_blocks(self):
        snapshot = analyze_document_bytes(
            _badly_formatted_article_docx(),
            "bad-layout.docx",
        )
        article = snapshot["article"]
        source_id = snapshot["paragraphs"][0]["block_id"]

        refined = refine_article_with_model(
            snapshot,
            article,
            complete_json=lambda _prompt: (
                '{"title":{"block_ids":["'
                + source_id
                + '"],"confidence":1,"value":"ВЫДУМАННОЕ НАЗВАНИЕ"}}'
            ),
        )

        self.assertEqual(
            refined["title"]["value"],
            "Надёжное извлечение данных из плохо оформленных статей",
        )

    def test_local_model_classifies_unlabeled_blocks_and_sections(self):
        texts = [
            "A STUDY WITHOUT STANDARD WORD STYLES",
            "Alice Researcher; Bob Scientist",
            "Example University",
            "This paragraph is the abstract even though it has no label.",
            "document parsing; block classification",
            "Background and Motivation",
            "The main article text follows the heading.",
        ]
        paragraphs = [
            {
                "index": index,
                "block_id": f"p:{index}",
                "text": text,
                "region": "document",
                "container": "body",
                "style": "custom-style",
                "bold": index in {0, 5},
            }
            for index, text in enumerate(texts)
        ]
        snapshot = {
            "suffix": ".docx",
            "paragraphs": paragraphs,
            "figures": [],
            "formulas": [],
            "tables": [],
            "requires_ocr": False,
        }
        article = extract_article_structure(snapshot)

        refined = refine_article_with_model(
            snapshot,
            article,
            complete_json=lambda _prompt: """
            {
              "title": {"block_ids": ["p:0"], "confidence": 0.99},
              "authors": {"block_ids": ["p:1"], "confidence": 0.98},
              "organizations": {"block_ids": ["p:2"], "confidence": 0.95},
              "abstract": {"block_ids": ["p:3"], "confidence": 0.96},
              "keywords": {"block_ids": ["p:4"], "confidence": 0.96},
              "section_headings": [
                {"block_id": "p:5", "level": 1, "confidence": 0.94}
              ]
            }
            """,
        )

        self.assertEqual(
            refined["authors"]["value"],
            ["Alice Researcher", "Bob Scientist"],
        )
        self.assertEqual(
            refined["organizations"]["value"],
            ["Example University"],
        )
        self.assertIn("no label", refined["abstract"]["value"])
        self.assertEqual(
            refined["keywords"]["value"],
            ["document parsing", "block classification"],
        )
        self.assertEqual(
            [section["title"] for section in refined["sections"]],
            ["Background and Motivation"],
        )

    def test_semantic_word_styles_extract_english_front_matter(self):
        snapshot = analyze_document_bytes(
            _semantic_style_article_docx(),
            "mdpi-layout.docx",
        )

        self.assertEqual(
            snapshot["metadata"]["title"],
            "A Resilient Parser for Scientific Articles",
        )
        self.assertEqual(
            snapshot["metadata"]["authors"],
            ["Alice Researcher", "Bob Scientist"],
        )
        self.assertEqual(
            snapshot["metadata"]["organizations"],
            "Document Intelligence Laboratory, Example University",
        )
        self.assertGreaterEqual(snapshot["metadata"]["confidence"]["title"], 0.8)

    def test_dotx_uses_the_same_structural_parser_as_docx(self):
        data = _as_dotx(_semantic_style_article_docx())
        snapshot = analyze_document_bytes(
            data,
            "journal-template.dotx",
        )

        self.assertEqual(snapshot["parse_error"], "")
        self.assertEqual(
            snapshot["metadata"]["title"],
            "A Resilient Parser for Scientific Articles",
        )
        rules = _extract_docx_rules(data)
        self.assertIn("page", rules)
        self.assertTrue(
            rules["body"]["line_spacing"] is None
            or rules["body"]["line_spacing"] < 10
        )

    def test_small_wmf_is_equation_and_not_a_content_figure(self):
        snapshot = analyze_document_bytes(
            _docx_with_embedded_vector_equation(),
            "legacy-equations.docx",
        )

        self.assertEqual(snapshot["embedded_image_count"], 2)
        self.assertEqual(snapshot["image_count"], 1)
        self.assertEqual(len(snapshot["figures"]), 1)
        self.assertEqual(len(snapshot["formulas"]), 1)
        self.assertEqual(
            snapshot["formulas"][0]["source"],
            "embedded_vector_equation",
        )

    def test_latex_bibitem_is_not_split_into_multiple_references(self):
        source = br"""
        \title{Reliable Extraction}
        \author{Alice Researcher}
        \abstract{A sufficiently clear abstract.}
        \keyword{documents; parsing}
        \section{Introduction}
        Body text.
        \section{References}
        \begin{thebibliography}{99}
        \bibitem{one} A. Author. \emph{First title}. Journal, 2024.
        \newblock DOI: \url{https://doi.org/10.1000/one}.
        \bibitem{two} B. Author. Second title. Journal, 2025.
        \end{thebibliography}
        """

        snapshot = analyze_document_bytes(source, "article.tex")

        self.assertEqual(len(snapshot["article"]["references"]), 2)
        self.assertIn(
            "First title",
            snapshot["article"]["references"][0]["text"],
        )
        self.assertIn(
            "10.1000/one",
            snapshot["article"]["references"][0]["text"],
        )

    def test_pdf_header_before_reference_number_is_discarded(self):
        texts = [
            "ROBUST REFERENCE EXTRACTION",
            "Abstract: Reference markers can follow a running header.",
            "Keywords: references; PDF",
            "References",
            (
                "Version July 2026, 25 of 26 12. Patoz, A.; Malatesta, D. "
                "A scientific article. Journal, 2025."
            ),
            "13. Smith, B. Another scientific article. Journal, 2026.",
        ]
        paragraphs = [
            {
                "index": index,
                "block_id": f"p:{index}",
                "text": text,
                "region": "document",
                "container": "page",
            }
            for index, text in enumerate(texts)
        ]

        article = extract_article_structure(
            {
                "suffix": ".pdf",
                "paragraphs": paragraphs,
                "figures": [],
                "formulas": [],
                "tables": [],
                "requires_ocr": False,
            }
        )

        self.assertEqual(
            [reference["number"] for reference in article["references"]],
            [12, 13],
        )

    def test_heading_style_does_not_turn_source_code_into_sections(self):
        from docx import Document

        document = Document()
        document.add_paragraph("PARSER VALIDATION ARTICLE", style="Title")
        document.add_paragraph("Abstract: The article validates heading rules.")
        document.add_paragraph("Keywords: parser; validation")
        document.add_paragraph("for(k : keyCollection){", style="Heading 1")
        document.add_paragraph("if(m.accept(k)){", style="Heading 1")
        document.add_paragraph("Литература", style="Heading 1")
        document.add_paragraph("1. Автор А. А. Название. 2025.")
        output = BytesIO()
        document.save(output)

        snapshot = analyze_document_bytes(output.getvalue(), "source-code.docx")

        self.assertEqual(
            [section["title"] for section in snapshot["article"]["sections"]],
            ["Литература"],
        )

    @patch(
        "apps.submissions.document_conversion.convert_legacy_doc_to_docx",
        return_value=_badly_formatted_article_docx(),
    )
    def test_legacy_doc_is_converted_then_parsed(self, _convert):
        snapshot = analyze_document_bytes(
            bytes.fromhex("d0cf11e0a1b11ae1") + b"legacy",
            "legacy.doc",
        )

        self.assertEqual(snapshot["parse_error"], "")
        self.assertEqual(snapshot["suffix"], ".doc")
        self.assertEqual(snapshot["converted_suffix"], ".docx")
        self.assertEqual(snapshot["article"]["source_format"], "doc")
        self.assertEqual(
            snapshot["metadata"]["title"],
            "Надёжное извлечение данных из плохо оформленных статей",
        )

    @patch(
        "apps.submissions.document_conversion.convert_legacy_doc_to_docx",
        side_effect=LegacyDocConversionError("Конвертер недоступен."),
    )
    def test_legacy_doc_reports_conversion_failure(self, _convert):
        snapshot = analyze_document_bytes(
            bytes.fromhex("d0cf11e0a1b11ae1") + b"legacy",
            "legacy.doc",
        )

        self.assertIn("Конвертер недоступен", snapshot["parse_error"])
        self.assertEqual(snapshot["article"]["title"]["value"], "")
