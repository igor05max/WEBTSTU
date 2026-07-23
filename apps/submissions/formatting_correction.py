import io
from pathlib import Path

from apps.submissions.document_analysis import read_file_bytes


class FormattingCorrectionError(ValueError):
    pass


def _as_float(value):
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def _iter_paragraphs(document):
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def build_corrected_docx(submission):
    version = submission.current_version
    if version is None or not version.file:
        raise FormattingCorrectionError("У заявки нет текущей версии файла.")
    if Path(version.file.name).suffix.casefold() != ".docx":
        raise FormattingCorrectionError("Автоматическое исправление доступно только для DOCX.")

    rules = (submission.formatting_rules_snapshot or {}).get("effective") or {}
    if not rules:
        raise FormattingCorrectionError("Для этой заявки не сохранены правила оформления.")

    try:
        from docx import Document
        from docx.enum.section import WD_ORIENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt
    except ImportError as exc:
        raise FormattingCorrectionError("На сервере не установлен модуль обработки DOCX.") from exc

    with version.file.open("rb") as source:
        original_bytes = read_file_bytes(source)
    try:
        document = Document(io.BytesIO(original_bytes))
    except Exception as exc:
        raise FormattingCorrectionError("Не удалось открыть DOCX для исправления.") from exc

    changes = []
    page_rules = rules.get("page") or {}
    margins = page_rules.get("margins_cm") or {}
    for section in document.sections:
        for field_name, attribute_name, label in (
            ("top", "top_margin", "верхнее поле"),
            ("right", "right_margin", "правое поле"),
            ("bottom", "bottom_margin", "нижнее поле"),
            ("left", "left_margin", "левое поле"),
        ):
            value = _as_float(margins.get(field_name))
            if value is not None:
                setattr(section, attribute_name, Cm(value))
                changes.append(f"{label}: {value:g} см")

        orientation = str(page_rules.get("orientation") or "").casefold()
        if orientation in {"landscape", "альбомная", "album"}:
            if section.page_height > section.page_width:
                section.page_width, section.page_height = section.page_height, section.page_width
            section.orientation = WD_ORIENT.LANDSCAPE
            changes.append("альбомная ориентация")
        elif orientation in {"portrait", "книжная"}:
            if section.page_width > section.page_height:
                section.page_width, section.page_height = section.page_height, section.page_width
            section.orientation = WD_ORIENT.PORTRAIT
            changes.append("книжная ориентация")

    body_rules = rules.get("body") or {}
    font_family = str(body_rules.get("font_family") or "").strip()
    font_size = _as_float(body_rules.get("font_size_pt"))
    line_spacing = _as_float(body_rules.get("line_spacing"))
    first_line_indent = _as_float(body_rules.get("first_line_indent_cm"))
    alignment = str(body_rules.get("alignment") or "").casefold()
    alignment_map = {
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "justified": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "по ширине": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "центр": WD_ALIGN_PARAGRAPH.CENTER,
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }

    normal_style = document.styles["Normal"]
    if font_family:
        normal_style.font.name = font_family
        style_fonts = normal_style._element.get_or_add_rPr().get_or_add_rFonts()
        style_fonts.set(qn("w:ascii"), font_family)
        style_fonts.set(qn("w:hAnsi"), font_family)
        style_fonts.set(qn("w:eastAsia"), font_family)
        changes.append(f"основной шрифт: {font_family}")
    if font_size is not None:
        normal_style.font.size = Pt(font_size)
        changes.append(f"размер основного шрифта: {font_size:g} пт")

    for paragraph in _iter_paragraphs(document):
        if not paragraph.text.strip():
            continue
        style_name = (paragraph.style.name or "").casefold() if paragraph.style else ""
        is_heading = "heading" in style_name or "заголов" in style_name
        if not is_heading:
            if line_spacing is not None:
                paragraph.paragraph_format.line_spacing = line_spacing
            if first_line_indent is not None:
                paragraph.paragraph_format.first_line_indent = Cm(first_line_indent)
            if alignment in alignment_map:
                paragraph.alignment = alignment_map[alignment]
            for run in paragraph.runs:
                if font_family:
                    run.font.name = font_family
                    run_fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
                    run_fonts.set(qn("w:ascii"), font_family)
                    run_fonts.set(qn("w:hAnsi"), font_family)
                    run_fonts.set(qn("w:eastAsia"), font_family)
                if font_size is not None:
                    run.font.size = Pt(font_size)

    if line_spacing is not None:
        changes.append(f"межстрочный интервал: {line_spacing:g}")
    if first_line_indent is not None:
        changes.append(f"абзацный отступ: {first_line_indent:g} см")
    if alignment in alignment_map:
        changes.append(f"выравнивание: {alignment}")

    output = io.BytesIO()
    document.save(output)
    return output.getvalue(), list(dict.fromkeys(changes))
