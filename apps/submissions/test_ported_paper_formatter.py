from io import BytesIO

from django.test import SimpleTestCase
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from apps.submissions.formatting_correction import _build_with_ported_formatter


def _document_bytes(document):
    output = BytesIO()
    document.save(output)
    return output.getvalue()


class PortedPaperFormatterTests(SimpleTestCase):
    def test_rebuilds_source_using_original_docx_template(self):
        source = Document()
        source.add_heading("Исследование шаблонного оформления", level=0)
        source.add_paragraph("Иванов Иван Иванович")
        source.add_heading("1. Введение", level=1)
        source.add_paragraph(
            "Основной текст статьи должен сохраниться после полного "
            "структурного преобразования."
        )
        source.add_table(rows=2, cols=2)
        source.tables[0].cell(0, 0).text = "Показатель"
        source.tables[0].cell(0, 1).text = "Значение"
        source.tables[0].cell(1, 0).text = "Точность"
        source.tables[0].cell(1, 1).text = "95"

        template = Document()
        title_style = template.styles["Title"]
        title_style.font.name = "Arial"
        title_style.font.size = Pt(18)
        title_style.font.bold = True
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        template.styles["Normal"].font.name = "Arial"
        template.styles["Normal"].font.size = Pt(10)
        template.add_paragraph("Пример названия", style="Title")
        template.add_paragraph("Пример основного текста")

        result_bytes, changes = _build_with_ported_formatter(
            _document_bytes(source),
            ("journal-template.docx", _document_bytes(template)),
        )
        result = Document(BytesIO(result_bytes))
        result_text = "\n".join(
            paragraph.text
            for paragraph in result.paragraphs
            if paragraph.text.strip()
        )

        self.assertIn("Исследование шаблонного оформления", result_text)
        self.assertIn("Основной текст статьи должен сохраниться", result_text)
        self.assertEqual(len(result.tables), 1)
        self.assertIn("применён полный профиль файла-шаблона", changes)
