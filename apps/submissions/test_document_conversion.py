import io
from zipfile import ZIP_DEFLATED, ZipFile

from django.test import SimpleTestCase

from apps.submissions.document_conversion import normalize_docx_compatibility


def _docx_with_modern_alignment(value):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    document = Document()
    paragraph = document.add_paragraph("Текст")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    source = io.BytesIO()
    document.save(source)

    output = io.BytesIO()
    with ZipFile(io.BytesIO(source.getvalue())) as input_archive, ZipFile(
        output,
        "w",
        ZIP_DEFLATED,
    ) as output_archive:
        for info in input_archive.infolist():
            payload = input_archive.read(info.filename)
            if info.filename == "word/document.xml":
                payload = payload.replace(
                    b'w:val="left"',
                    f'w:val="{value}"'.encode(),
                )
            output_archive.writestr(info, payload)
    return output.getvalue()


class DocxCompatibilityTests(SimpleTestCase):
    def test_normalizes_libreoffice_start_and_end_alignments(self):
        from docx import Document

        for source_value, expected_value in (("start", "left"), ("end", "right")):
            normalized = normalize_docx_compatibility(
                _docx_with_modern_alignment(source_value)
            )
            with ZipFile(io.BytesIO(normalized)) as archive:
                document_xml = archive.read("word/document.xml")
            self.assertIn(
                f'w:val="{expected_value}"'.encode(),
                document_xml,
            )
            self.assertNotIn(
                f'w:val="{source_value}"'.encode(),
                document_xml,
            )
            self.assertIsNotNone(Document(io.BytesIO(normalized)).paragraphs[0].alignment)
