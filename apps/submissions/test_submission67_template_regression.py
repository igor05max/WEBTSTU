from io import BytesIO

from django.test import SimpleTestCase
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

from apps.directory.formatting_templates import _extract_docx_rules
from document_template_engine import (
    build_docx_from_template,
    normalize_template_rules,
)


def _docx_bytes(document):
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _add_paragraph_style(
    document,
    name,
    *,
    size,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    bold=False,
    italic=False,
    left_indent=0,
    right_indent=0,
    first_line_indent=0,
    space_before=0,
    space_after=0,
    line_spacing=1.0,
    keep_with_next=False,
    keep_together=False,
):
    style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = "Palatino Linotype"
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.italic = italic
    style.paragraph_format.alignment = alignment
    style.paragraph_format.left_indent = Cm(left_indent)
    style.paragraph_format.right_indent = Cm(right_indent)
    style.paragraph_format.first_line_indent = Cm(first_line_indent)
    style.paragraph_format.space_before = Pt(space_before)
    style.paragraph_format.space_after = Pt(space_after)
    style.paragraph_format.line_spacing = line_spacing
    style.paragraph_format.keep_with_next = keep_with_next
    style.paragraph_format.keep_together = keep_together
    return style


class DocxTemplateSemanticExtractionTests(SimpleTestCase):
    def test_extracts_semantic_blocks_extended_properties_and_heading_levels(self):
        template = Document()

        _add_paragraph_style(
            template,
            "MDPI_1.2_title",
            size=20,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            bold=True,
            space_after=8,
            keep_with_next=True,
        )
        _add_paragraph_style(
            template,
            "MDPI_1.3_authornames",
            size=11,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            italic=True,
            space_after=4,
            keep_with_next=True,
        )
        _add_paragraph_style(
            template,
            "MDPI_1.7_abstract",
            size=10,
            left_indent=0.4,
            right_indent=0.3,
            space_before=5,
            space_after=5,
            line_spacing=1.05,
            keep_together=True,
        )
        _add_paragraph_style(
            template,
            "MDPI_1.8_keywords",
            size=10,
            left_indent=0.4,
            right_indent=0.3,
            space_after=6,
        )
        _add_paragraph_style(
            template,
            "MDPI_3.1_text",
            size=11,
            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
            left_indent=0.2,
            right_indent=0.1,
            first_line_indent=0.75,
            space_before=2,
            space_after=3,
            line_spacing=1.15,
            keep_together=True,
        )
        _add_paragraph_style(
            template,
            "MDPI_8.1_references",
            size=9,
            left_indent=0.5,
            first_line_indent=-0.5,
            space_after=2,
        )
        _add_paragraph_style(
            template,
            "MDPI_2.1_heading1",
            size=15,
            bold=True,
            space_before=12,
            space_after=4,
            keep_with_next=True,
        )
        _add_paragraph_style(
            template,
            "MDPI_2.2_heading2",
            size=13,
            italic=True,
            left_indent=0.25,
            space_before=8,
            space_after=3,
            keep_with_next=True,
        )
        _add_paragraph_style(
            template,
            "MDPI_2.3_heading3",
            size=11,
            bold=True,
            italic=True,
            left_indent=0.5,
            space_before=6,
            space_after=2,
            keep_with_next=True,
        )

        for text, style_name in (
            ("A Real Research Article", "MDPI_1.2_title"),
            ("Anna Author and Boris Researcher", "MDPI_1.3_authornames"),
            (
                "Abstract: This study presents a sufficiently concrete abstract.",
                "MDPI_1.7_abstract",
            ),
            (
                "Keywords: document processing; semantic styles",
                "MDPI_1.8_keywords",
            ),
            ("1. Introduction", "MDPI_2.1_heading1"),
            ("1.1. Motivation", "MDPI_2.2_heading2"),
            ("1.1.1. Details", "MDPI_2.3_heading3"),
            (
                "This paragraph represents the ordinary main text of the article.",
                "MDPI_3.1_text",
            ),
            (
                "[1] Author, A. A concrete bibliographic reference. 2026.",
                "MDPI_8.1_references",
            ),
        ):
            template.add_paragraph(text, style=style_name)

        rules = _extract_docx_rules(_docx_bytes(template))

        self.assertEqual(rules["document"]["analysis_revision"], 3)
        block_styles = {
            block["role"]: block["style"]
            for block in rules["document"]["blocks"]
        }
        self.assertTrue(
            {
                "title",
                "authors",
                "abstract",
                "keywords",
                "body",
                "references",
            }.issubset(block_styles)
        )

        self.assertEqual(block_styles["title"]["font_size_pt"], 20)
        self.assertEqual(block_styles["title"]["alignment"], "center")
        self.assertIs(block_styles["title"]["bold"], True)
        self.assertIs(block_styles["authors"]["italic"], True)

        body = block_styles["body"]
        self.assertAlmostEqual(body["left_indent_cm"], 0.2, places=2)
        self.assertAlmostEqual(body["right_indent_cm"], 0.1, places=2)
        self.assertAlmostEqual(body["first_line_indent_cm"], 0.75, places=2)
        self.assertAlmostEqual(body["space_before_pt"], 2, places=1)
        self.assertAlmostEqual(body["space_after_pt"], 3, places=1)
        self.assertAlmostEqual(body["line_spacing"], 1.15, places=2)
        self.assertIs(body["keep_together"], True)

        abstract = block_styles["abstract"]
        self.assertAlmostEqual(abstract["left_indent_cm"], 0.4, places=2)
        self.assertAlmostEqual(abstract["right_indent_cm"], 0.3, places=2)
        self.assertAlmostEqual(abstract["space_before_pt"], 5, places=1)
        self.assertAlmostEqual(abstract["space_after_pt"], 5, places=1)
        self.assertIs(abstract["keep_together"], True)

        levels = rules["headings"]["levels"]
        self.assertEqual(levels["1"]["font_size_pt"], 15)
        self.assertIs(levels["1"]["bold"], True)
        self.assertIs(levels["1"]["keep_with_next"], True)
        self.assertEqual(levels["1"]["color_hex"], "000000")
        self.assertEqual(levels["2"]["font_size_pt"], 13)
        self.assertIs(levels["2"]["italic"], True)
        self.assertAlmostEqual(levels["2"]["left_indent_cm"], 0.25, places=2)
        self.assertEqual(levels["3"]["font_size_pt"], 11)
        self.assertIs(levels["3"]["bold"], True)
        self.assertIs(levels["3"]["italic"], True)


class DocxTemplateHeadingApplicationTests(SimpleTestCase):
    def test_applies_numbered_heading_levels_and_uses_metadata_for_title(self):
        source = Document()
        source.add_paragraph("A Real Research Article")
        source.add_paragraph("1. Introduction")
        source.add_paragraph("First body paragraph.")
        source.add_paragraph("1.1. Scope")
        source.add_paragraph("Second body paragraph.")
        source.add_paragraph("2. Methods", style="Heading 1")
        source.paragraphs[-1].runs[0].font.color.rgb = RGBColor(
            0x00,
            0x66,
            0xCC,
        )

        rules = normalize_template_rules(
            {
                "body": {
                    "font_family": "Arial",
                    "font_size_pt": 10,
                    "alignment": "justify",
                },
                "headings": {
                    "font_family": "Arial",
                    "font_size_pt": 14,
                    "title_font_size_pt": 29,
                    "levels": {
                        "1": {
                            "font_family": "Arial",
                            "font_size_pt": 16,
                            "left_indent_cm": 0.25,
                            "bold": True,
                            "color_hex": "000000",
                        },
                        "2": {
                            "font_family": "Arial",
                            "font_size_pt": 13,
                            "left_indent_cm": 0.75,
                            "italic": True,
                        },
                        "3": {
                            "font_family": "Arial",
                            "font_size_pt": 11,
                        },
                    },
                },
                "document": {
                    "analysis_revision": 3,
                    "blocks": [
                        {
                            "role": "title",
                            "style": {
                                "font_family": "Arial",
                                "font_size_pt": 23,
                                "alignment": "center",
                                "bold": True,
                            },
                        },
                        {
                            "role": "body",
                            "style": {
                                "font_family": "Arial",
                                "font_size_pt": 10,
                                "alignment": "justify",
                            },
                        },
                    ],
                },
            }
        )

        built, _changes, _plan = build_docx_from_template(
            _docx_bytes(source),
            rules,
            metadata={"title": "A Real Research Article"},
        )
        corrected = Document(BytesIO(built))
        paragraphs = {paragraph.text: paragraph for paragraph in corrected.paragraphs}

        title = paragraphs["A Real Research Article"]
        level_one = paragraphs["1. Introduction"]
        level_two = paragraphs["1.1. Scope"]
        styled_level_one = paragraphs["2. Methods"]

        self.assertAlmostEqual(title.runs[0].font.size.pt, 23, places=1)
        self.assertEqual(title.alignment, WD_ALIGN_PARAGRAPH.CENTER)
        self.assertIs(title.runs[0].bold, True)

        self.assertAlmostEqual(level_one.runs[0].font.size.pt, 16, places=1)
        self.assertAlmostEqual(
            level_one.paragraph_format.left_indent.cm,
            0.25,
            places=2,
        )
        self.assertIs(level_one.runs[0].bold, True)

        self.assertAlmostEqual(level_two.runs[0].font.size.pt, 13, places=1)
        self.assertAlmostEqual(
            level_two.paragraph_format.left_indent.cm,
            0.75,
            places=2,
        )
        self.assertIs(level_two.runs[0].italic, True)

        self.assertAlmostEqual(
            styled_level_one.runs[0].font.size.pt,
            16,
            places=1,
        )
        self.assertEqual(
            str(styled_level_one.runs[0].font.color.rgb),
            "000000",
        )
        self.assertNotEqual(
            styled_level_one.runs[0].font.size.pt,
            title.runs[0].font.size.pt,
        )


class DocxTemplateParagraphStyleApplicationTests(SimpleTestCase):
    def test_applies_caption_list_and_table_body_styles(self):
        source = Document()
        source.add_paragraph("Integration Article")
        source.add_paragraph("Рисунок 1. Схема обработки документа")
        source.add_paragraph("Таблица 1. Результаты проверки")
        table = source.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "Значение в таблице"
        source.add_paragraph("RQ 1. Как изменяется качество форматирования?")
        source.add_paragraph("Маркированный пункт", style="List Bullet")

        rules = normalize_template_rules(
            {
                "body": {
                    "font_family": "Arial",
                    "font_size_pt": 10,
                    "line_spacing": 1.4,
                    "alignment": "justify",
                },
                "document": {
                    "analysis_revision": 3,
                    "blocks": [
                        {
                            "role": "body",
                            "style": {
                                "font_family": "Arial",
                                "font_size_pt": 10,
                                "alignment": "justify",
                            },
                        },
                    ],
                    "paragraph_styles": {
                        "figure_caption": {
                            "font_family": "Arial",
                            "font_size_pt": 8,
                            "alignment": "center",
                            "space_before_pt": 5,
                            "italic": True,
                        },
                        "table_caption": {
                            "font_family": "Arial",
                            "font_size_pt": 9,
                            "alignment": "right",
                            "space_after_pt": 4,
                            "bold": True,
                        },
                        "list_itemize": {
                            "font_family": "Arial",
                            "font_size_pt": 11,
                            "left_indent_cm": 0.9,
                            "first_line_indent_cm": -0.25,
                        },
                        "list_bullet": {
                            "font_family": "Arial",
                            "font_size_pt": 12,
                            "left_indent_cm": 1.2,
                            "italic": True,
                        },
                        "table_body": {
                            "font_family": "Arial",
                            "alignment": "center",
                        },
                    },
                },
            }
        )

        built, changes, _plan = build_docx_from_template(
            _docx_bytes(source),
            rules,
            metadata={"title": "Integration Article"},
        )
        corrected = Document(BytesIO(built))
        paragraphs = {paragraph.text: paragraph for paragraph in corrected.paragraphs}

        figure_caption = paragraphs["Рисунок 1. Схема обработки документа"]
        table_caption = paragraphs["Таблица 1. Результаты проверки"]
        itemized = paragraphs["RQ 1. Как изменяется качество форматирования?"]
        bulleted = paragraphs["Маркированный пункт"]
        table_body = corrected.tables[0].cell(0, 0).paragraphs[0]

        self.assertAlmostEqual(
            figure_caption.runs[0].font.size.pt,
            8,
            places=1,
        )
        self.assertEqual(
            figure_caption.alignment,
            WD_ALIGN_PARAGRAPH.CENTER,
        )
        self.assertIs(figure_caption.runs[0].italic, True)

        self.assertAlmostEqual(
            table_caption.runs[0].font.size.pt,
            9,
            places=1,
        )
        self.assertEqual(table_caption.alignment, WD_ALIGN_PARAGRAPH.RIGHT)
        self.assertIs(table_caption.runs[0].bold, True)

        self.assertAlmostEqual(itemized.runs[0].font.size.pt, 11, places=1)
        self.assertAlmostEqual(
            itemized.paragraph_format.left_indent.cm,
            0.9,
            places=2,
        )
        self.assertAlmostEqual(
            itemized.paragraph_format.first_line_indent.cm,
            -0.25,
            places=2,
        )

        self.assertAlmostEqual(bulleted.runs[0].font.size.pt, 12, places=1)
        self.assertAlmostEqual(
            bulleted.paragraph_format.left_indent.cm,
            1.2,
            places=2,
        )
        self.assertIs(bulleted.runs[0].italic, True)

        self.assertAlmostEqual(table_body.runs[0].font.size.pt, 10, places=1)
        self.assertEqual(table_body.paragraph_format.line_spacing, 1.4)
        self.assertEqual(table_body.alignment, WD_ALIGN_PARAGRAPH.CENTER)

        changes_text = " ".join(str(change) for change in changes).casefold()
        self.assertIn("подпис", changes_text)
        self.assertIn("списк", changes_text)
        self.assertIn("текст", changes_text)
        self.assertIn("таблиц", changes_text)

    def test_skips_layout_tables_and_protected_cells(self):
        from docx.oxml import OxmlElement

        source = Document()
        source.add_paragraph("Safe Table Article")
        source.add_paragraph("Layout container")
        layout_table = source.add_table(rows=1, cols=1)
        layout_paragraph = layout_table.cell(0, 0).paragraphs[0]
        layout_paragraph.text = "Layout value"
        layout_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        layout_paragraph.paragraph_format.line_spacing = 2
        layout_paragraph.runs[0].font.name = "Courier New"
        layout_paragraph.runs[0].font.size = Pt(18)

        source.add_paragraph("Таблица 2. Защищённые ячейки")
        data_table = source.add_table(rows=1, cols=3)
        regular = data_table.cell(0, 0).paragraphs[0]
        regular.text = "Regular data"
        regular.runs[0].font.name = "Courier New"
        regular.runs[0].font.size = Pt(18)

        equation = data_table.cell(0, 1).paragraphs[0]
        equation.text = "Formula data"
        equation.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        equation.paragraph_format.line_spacing = 2
        equation.runs[0].font.name = "Courier New"
        equation.runs[0].font.size = Pt(18)
        equation._p.append(OxmlElement("m:oMath"))

        nested_cell = data_table.cell(0, 2)
        nested = nested_cell.paragraphs[0]
        nested.text = "Nested layout"
        nested.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        nested.paragraph_format.line_spacing = 2
        nested.runs[0].font.name = "Courier New"
        nested.runs[0].font.size = Pt(18)
        nested_cell.add_table(rows=1, cols=1).cell(0, 0).text = "Inner"

        rules = normalize_template_rules(
            {
                "body": {
                    "font_family": "Arial",
                    "font_size_pt": 10,
                    "line_spacing": 1.4,
                    "first_line_indent_cm": 1.25,
                    "alignment": "justify",
                },
                "document": {
                    "blocks": [{"role": "body", "style": {}}],
                    "paragraph_styles": {
                        "table_body": {
                            "font_family": "Arial",
                            "alignment": "center",
                        },
                    },
                },
            }
        )

        built, _changes, _plan = build_docx_from_template(
            _docx_bytes(source),
            rules,
            metadata={"title": "Safe Table Article"},
        )
        corrected = Document(BytesIO(built))
        result_layout = corrected.tables[0].cell(0, 0).paragraphs[0]
        result_regular = corrected.tables[1].cell(0, 0).paragraphs[0]
        result_equation = corrected.tables[1].cell(0, 1).paragraphs[0]
        result_nested = corrected.tables[1].cell(0, 2).paragraphs[0]

        self.assertEqual(result_layout.runs[0].font.name, "Courier New")
        self.assertAlmostEqual(result_layout.runs[0].font.size.pt, 18, places=1)
        self.assertEqual(result_layout.alignment, WD_ALIGN_PARAGRAPH.RIGHT)
        self.assertEqual(result_layout.paragraph_format.line_spacing, 2)

        self.assertEqual(result_regular.runs[0].font.name, "Arial")
        self.assertAlmostEqual(result_regular.runs[0].font.size.pt, 10, places=1)
        self.assertEqual(result_regular.paragraph_format.line_spacing, 1.4)
        self.assertEqual(result_regular.alignment, WD_ALIGN_PARAGRAPH.CENTER)
        self.assertIsNone(result_regular.paragraph_format.first_line_indent)

        for protected in (result_equation, result_nested):
            self.assertEqual(protected.runs[0].font.name, "Courier New")
            self.assertAlmostEqual(protected.runs[0].font.size.pt, 18, places=1)
            self.assertEqual(protected.alignment, WD_ALIGN_PARAGRAPH.RIGHT)
            self.assertEqual(protected.paragraph_format.line_spacing, 2)
