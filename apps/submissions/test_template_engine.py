from io import BytesIO

from django.test import SimpleTestCase

from document_template_engine import (
    build_docx_from_template,
    build_docx_plan,
    build_latex_template,
    check_latex_against_template,
    extract_latex_template_rules,
    interpret_template_text,
    normalize_template_rules,
)


LEGACY_RULES = {
    "page": {
        "size": "A4",
        "orientation": "portrait",
        "margins_cm": {"top": 2, "right": 2, "bottom": 2, "left": 2},
    },
    "body": {
        "font_family": "Times New Roman",
        "font_size_pt": 14,
        "line_spacing": 1,
        "first_line_indent_cm": 1,
        "alignment": "justify",
    },
    "structure": {
        "required_sections": [
            "UDC index",
            "Title",
            "Author initials and surname",
            "Scientific supervisor (optional)",
            "Institution name",
            "City, Country",
            "Abstract text",
            "References (optional)",
        ],
    },
    "metadata": {
        "required_fields": [
            "UDC index",
            "Title",
            "Author initials and surname",
            "Institution name",
            "City",
            "Country",
        ]
    },
}


def _sample_docx():
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt

    document = Document()
    for section in document.sections:
        section.top_margin = Cm(2)
        section.right_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)
    normal.paragraph_format.first_line_indent = Cm(1)
    normal.paragraph_format.line_spacing = 1

    values = [
        ("УДК 004.8", WD_ALIGN_PARAGRAPH.LEFT, 0),
        ("ИНТЕЛЛЕКТУАЛЬНАЯ СИСТЕМА АНАЛИЗА ВИДЕО", WD_ALIGN_PARAGRAPH.CENTER, 0),
        ("А.А. АРХИПОВ, И.И. МАКСИМОВ", WD_ALIGN_PARAGRAPH.CENTER, 0),
        ("Научный руководитель А. Д. ОБУХОВ", WD_ALIGN_PARAGRAPH.CENTER, 0),
        ("Тамбовский государственный технический университет", WD_ALIGN_PARAGRAPH.CENTER, 0),
        ("Тамбов, Россия", WD_ALIGN_PARAGRAPH.CENTER, 0),
        (
            "Основной научный текст сохраняется без изменения содержания и фактов.",
            WD_ALIGN_PARAGRAPH.JUSTIFY,
            1,
        ),
        ("СПИСОК ИСПОЛЬЗОВАННОЙ ЛИТЕРАТУРЫ", WD_ALIGN_PARAGRAPH.CENTER, 0),
        ("1. Иванов И. И. Анализ документов. 2026.", WD_ALIGN_PARAGRAPH.LEFT, 0),
    ]
    for text, alignment, indent in values:
        paragraph = document.add_paragraph(text)
        paragraph.alignment = alignment
        paragraph.paragraph_format.first_line_indent = Cm(indent)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


class ReusableTemplateEngineTests(SimpleTestCase):
    def test_latex_rules_are_extracted_without_executing_source(self):
        source = r"""
            \documentclass[14pt,a4paper]{article}
            \usepackage[top=2cm,right=1.5cm,bottom=2cm,left=2.5cm]{geometry}
            \usepackage{fontspec}
            \setmainfont{Times New Roman}
            \usepackage{setspace}
            \setstretch{1.5}
            \setlength{\parindent}{1cm}
            \title{НАЗВАНИЕ}
            \author{И.О. Фамилия}
            \begin{document}
            \section{Введение}
            Научный текст.
            \end{document}
        """

        rules = extract_latex_template_rules(source)

        self.assertEqual(rules["page"]["size"], "A4")
        self.assertEqual(rules["page"]["margins_cm"]["right"], 1.5)
        self.assertEqual(rules["body"]["font_family"], "Times New Roman")
        self.assertEqual(rules["body"]["font_size_pt"], 14)
        self.assertEqual(rules["body"]["line_spacing"], 1.5)
        self.assertEqual(rules["body"]["first_line_indent_cm"], 1)
        self.assertEqual(rules["structure"]["required_sections"], ["Введение"])

    def test_generated_latex_uses_rules_and_passes_layout_check(self):
        rules = {
            "page": {
                "size": "A4",
                "orientation": "portrait",
                "margins_cm": {"top": 2, "right": 2, "bottom": 2, "left": 2},
            },
            "body": {
                "font_family": "Times New Roman",
                "font_size_pt": 14,
                "line_spacing": 1,
                "first_line_indent_cm": 1,
                "alignment": "justify",
            },
            "document": {
                "blocks": [
                    {"role": "udc", "required": True},
                    {"role": "title", "required": True},
                    {"role": "authors", "required": True},
                    {"role": "abstract", "required": True},
                    {"role": "keywords", "required": True},
                    {"role": "body", "required": True},
                ]
            },
        }

        source = build_latex_template(
            rules,
            metadata={
                "title": "НАЗВАНИЕ РАБОТЫ",
                "authors": "И.О. Фамилия",
                "abstract": "Краткое описание работы.",
                "keywords": "анализ; система",
            },
        )
        report = check_latex_against_template(source, rules)

        decoded = source.decode("utf-8")
        self.assertIn(r"\setmainfont{Times New Roman}", decoded)
        self.assertIn(r"\usepackage[a4paper,top=2cm,right=2cm,bottom=2cm,left=2cm]{geometry}", decoded)
        self.assertEqual(report["issues"], [])
        self.assertEqual(report["metrics"]["source_format"], "latex")

    def test_title_centering_does_not_change_detected_body_alignment(self):
        source = r"""
            \documentclass[12pt]{article}
            \begin{document}
            {\centering НАЗВАНИЕ\par}
            Основной текст выровнен по ширине по умолчанию.
            \end{document}
        """

        rules = extract_latex_template_rules(source)

        self.assertEqual(rules["body"]["alignment"], "justify")
        self.assertEqual(rules["body"]["font_size_pt"], 12)

    def test_ai_unspecified_false_values_do_not_erase_author_formatting(self):
        rules = interpret_template_text(
            document_type="Тезисы",
            target_name="Конференция",
            text="Список литературы при необходимости оформляется по СТБ.",
            complete_json=lambda _prompt: """
                {
                  "document": {
                    "blocks": [
                      {
                        "role": "title",
                        "required": true,
                        "style": {
                          "first_line_indent_cm": null,
                          "bold": false,
                          "italic": false
                        }
                      },
                      {
                        "role": "authors",
                        "required": true,
                        "style": {"line_spacing": 2}
                      },
                      {
                        "role": "references",
                        "required": false,
                        "style": {
                          "alignment": "justify",
                          "first_line_indent_cm": 1,
                          "bold": false
                        }
                      }
                    ]
                  },
                  "body": {"line_spacing": 1}
                }
            """,
        )

        blocks = {
            block["role"]: block
            for block in rules["document"]["blocks"]
        }
        self.assertNotIn("bold", blocks["title"]["style"])
        self.assertNotIn("italic", blocks["title"]["style"])
        self.assertEqual(blocks["title"]["style"]["first_line_indent_cm"], 0)
        self.assertEqual(blocks["authors"]["style"]["line_spacing"], 1)
        self.assertNotIn("alignment", blocks["references"]["style"])
        self.assertNotIn("first_line_indent_cm", blocks["references"]["style"])

    def test_legacy_placeholders_become_blocks_not_required_sections(self):
        normalized = normalize_template_rules(LEGACY_RULES)

        self.assertEqual(normalized["schema_version"], "2.0")
        self.assertEqual(normalized["structure"]["required_sections"], [])
        blocks = {
            block["role"]: block
            for block in normalized["document"]["blocks"]
        }
        self.assertTrue(blocks["udc"]["required"])
        self.assertTrue(blocks["title"]["required"])
        self.assertTrue(blocks["body"]["required"])
        self.assertFalse(blocks["supervisor"]["required"])
        self.assertFalse(blocks["references"]["required"])

    def test_english_template_block_labels_are_displayed_in_russian(self):
        normalized = normalize_template_rules(
            {
                "document": {
                    "blocks": [
                        {"role": "title", "label": "Title", "required": True},
                        {
                            "role": "authors",
                            "label": "Firstname Lastname",
                            "required": True,
                        },
                        {
                            "role": "institution",
                            "label": "Affiliation",
                            "required": True,
                        },
                        {"role": "abstract", "label": "Abstract", "required": True},
                        {"role": "keywords", "label": "Keywords", "required": True},
                        {"role": "body", "label": "Main Text", "required": True},
                        {
                            "role": "references",
                            "label": "References",
                            "required": True,
                        },
                    ]
                }
            }
        )

        labels = {
            block["role"]: block["label"]
            for block in normalized["document"]["blocks"]
        }
        self.assertEqual(labels["title"], "Название")
        self.assertEqual(labels["authors"], "Авторы")
        self.assertEqual(labels["institution"], "Организация")
        self.assertEqual(labels["abstract"], "Аннотация")
        self.assertEqual(labels["keywords"], "Ключевые слова")
        self.assertEqual(labels["body"], "Основной текст")
        self.assertEqual(labels["references"], "Список литературы")

    def test_affiliation_is_a_real_required_organization_block(self):
        normalized = normalize_template_rules(
            {"structure": {"required_sections": ["Affiliation"]}}
        )
        blocks = {
            block["role"]: block
            for block in normalized["document"]["blocks"]
        }

        self.assertTrue(blocks["institution"]["required"])
        self.assertEqual(blocks["institution"]["label"], "Организация")
        self.assertEqual(normalized["structure"]["required_sections"], [])

    def test_institutional_review_statement_is_not_an_organization(self):
        normalized = normalize_template_rules(
            {
                "structure": {
                    "required_sections": ["Institutional Review Board Statement"]
                }
            }
        )

        roles = {
            block["role"]
            for block in normalized["document"]["blocks"]
        }
        self.assertNotIn("institution", roles)
        self.assertEqual(
            normalized["structure"]["required_sections"],
            ["Institutional Review Board Statement"],
        )

    def test_saved_false_institution_block_is_migrated_to_required_section(self):
        normalized = normalize_template_rules(
            {
                "document": {
                    "blocks": [
                        {
                            "role": "institution",
                            "label": "Организация",
                            "source_label": "Institutional Review Board Statement",
                            "required": True,
                        }
                    ]
                }
            }
        )

        roles = {
            block["role"]
            for block in normalized["document"]["blocks"]
        }
        self.assertNotIn("institution", roles)
        self.assertEqual(
            normalized["structure"]["required_sections"],
            ["Institutional Review Board Statement"],
        )

    def test_required_organization_from_metadata_is_ready_to_fill(self):
        from docx import Document

        document = Document()
        document.add_paragraph("НАЗВАНИЕ ИССЛЕДОВАНИЯ")
        document.add_paragraph(
            "Основной текст исследования содержит достаточно слов для распознавания."
        )
        source = BytesIO()
        document.save(source)
        plan = build_docx_plan(
            source.getvalue(),
            {
                "document": {
                    "blocks": [
                        {
                            "role": "institution",
                            "label": "Affiliation",
                            "required": True,
                        }
                    ]
                }
            },
            metadata={"organizations": "Тамбовский государственный университет"},
        )
        organization = next(
            block for block in plan["blocks"] if block["role"] == "institution"
        )

        self.assertFalse(organization["found"])
        self.assertTrue(organization["can_fill"])
        self.assertEqual(organization["status"], "ready_to_fill")
        self.assertNotIn(organization, plan["missing_blocks"])

    def test_plan_recognizes_title_block_without_false_missing_sections(self):
        plan = build_docx_plan(_sample_docx(), LEGACY_RULES)

        self.assertEqual(plan["issues"], [])
        found = {
            block["role"]
            for block in plan["blocks"]
            if block["found"]
        }
        self.assertTrue(
            {"udc", "title", "authors", "supervisor", "institution", "city_country", "body", "references"}
            <= found
        )

    def test_builder_formats_roles_separately_and_preserves_text(self):
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        source = _sample_docx()
        built, changes, plan = build_docx_from_template(source, LEGACY_RULES)
        original = Document(BytesIO(source))
        result = Document(BytesIO(built))

        self.assertEqual(
            [paragraph.text for paragraph in result.paragraphs],
            [paragraph.text for paragraph in original.paragraphs],
        )
        self.assertEqual(result.paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.LEFT)
        self.assertAlmostEqual(
            result.paragraphs[0].paragraph_format.first_line_indent.cm,
            0,
            places=1,
        )
        self.assertEqual(result.paragraphs[1].alignment, WD_ALIGN_PARAGRAPH.CENTER)
        self.assertAlmostEqual(
            result.paragraphs[1].paragraph_format.first_line_indent.cm,
            0,
            places=1,
        )
        self.assertEqual(result.paragraphs[6].alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)
        self.assertAlmostEqual(
            result.paragraphs[6].paragraph_format.first_line_indent.cm,
            1,
            places=1,
        )
        self.assertEqual(result.paragraphs[7].alignment, WD_ALIGN_PARAGRAPH.CENTER)
        self.assertEqual(plan["issues"], [])
        self.assertIn("титульный блок оформлен отдельно от основного текста", changes)

    def test_builder_centers_and_scales_figures_to_the_new_text_width(self):
        import base64
        from copy import deepcopy

        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Cm

        pixel = base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR4nGNg"
            "YAAAAAMAASsJTYQAAAAASUVORK5CYII="
        )
        document = Document()
        document.add_paragraph(
            "Основной научный текст перед рисунком сохраняется без изменений."
        )
        figure = document.add_paragraph()
        figure.paragraph_format.first_line_indent = Cm(1.5)
        figure.add_run().add_picture(BytesIO(pixel), width=Cm(18), height=Cm(6))
        caption = document.add_paragraph("Рисунок 1 — Экспериментальная схема.")
        caption.paragraph_format.first_line_indent = Cm(1.5)
        source = BytesIO()
        document.save(source)
        rules = deepcopy(LEGACY_RULES)
        rules["page"]["margins_cm"] = {
            "top": 2,
            "right": 1.2,
            "bottom": 2,
            "left": 5.8,
        }

        built, changes, _plan = build_docx_from_template(
            source.getvalue(),
            rules,
        )
        result = Document(BytesIO(built))
        result_figure = next(
            paragraph
            for paragraph in result.paragraphs
            if paragraph._p.xpath(".//w:drawing")
        )
        result_caption = next(
            paragraph
            for paragraph in result.paragraphs
            if paragraph.text.startswith("Рисунок 1")
        )
        section = result.sections[0]
        maximum_width = (
            section.page_width
            - section.left_margin
            - section.right_margin
            - Cm(0.2)
        )

        self.assertLessEqual(result.inline_shapes[0].width, maximum_width)
        self.assertEqual(result_figure.alignment, WD_ALIGN_PARAGRAPH.CENTER)
        self.assertAlmostEqual(
            result_figure.paragraph_format.first_line_indent.cm,
            0,
            places=2,
        )
        self.assertTrue(result_figure.paragraph_format.keep_with_next)
        self.assertEqual(result_caption.alignment, WD_ALIGN_PARAGRAPH.CENTER)
        self.assertAlmostEqual(
            result_caption.paragraph_format.first_line_indent.cm,
            0,
            places=2,
        )
        self.assertIn("рисунки вписаны в рабочую область страницы: 1", changes)

    def test_builder_merges_adjacent_citation_markers_from_an_older_result(self):
        from docx import Document

        document = Document()
        document.add_paragraph(
            "Сложность оценки качества ответа [1]. [8, 9] "
            "Следующее предложение сохраняется."
        )
        source = BytesIO()
        document.save(source)

        built, changes, _plan = build_docx_from_template(
            source.getvalue(),
            LEGACY_RULES,
        )
        result = Document(BytesIO(built))

        self.assertEqual(
            result.paragraphs[0].text,
            "Сложность оценки качества ответа [1, 8, 9]. "
            "Следующее предложение сохраняется.",
        )
        self.assertIn("объединены соседние маркеры источников: 1", changes)
