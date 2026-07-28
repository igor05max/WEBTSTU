from io import BytesIO
from tempfile import TemporaryDirectory
from unittest.mock import patch

from django.contrib.auth import get_user_model
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TestCase, override_settings

from apps.directory.formatting_templates import process_formatting_template
from apps.directory.models import ArticleType, FormattingTemplate, Journal
from apps.submissions.document_conversion import LEGACY_DOC_SIGNATURE
from apps.submissions.models import SubmissionStatus
from apps.submissions.services import (
    add_submission_version,
    create_submission_with_initial_version,
)


def _editable_docx():
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt

    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.right_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    for index in range(8):
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.first_line_indent = Cm(1.25)
        paragraph.paragraph_format.line_spacing = 1.5
        run = paragraph.add_run(
            f"Обычный абзац научной статьи номер {index} с достаточным объёмом текста."
        )
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


@override_settings(SUBMISSION_CHECKS_ASYNC=False)
class LegacyDocWorkflowTests(TestCase):
    def setUp(self):
        self.media_dir = TemporaryDirectory()
        self.settings_override = override_settings(MEDIA_ROOT=self.media_dir.name)
        self.settings_override.enable()
        self.user = get_user_model().objects.create_user(
            username="legacy_doc_author",
            password="1234",
        )
        self.journal = Journal.objects.create(name="DOC Journal")
        self.article_type = ArticleType.objects.create(
            code="article",
            name="Статья",
        )

    def tearDown(self):
        self.settings_override.disable()
        self.media_dir.cleanup()

    @patch(
        "apps.submissions.document_uploads.convert_legacy_doc_to_docx",
        return_value=_editable_docx(),
    )
    def test_initial_doc_is_preserved_and_current_version_is_editable_docx(
        self,
        mocked_convert,
    ):
        source_bytes = LEGACY_DOC_SIGNATURE + b"legacy-source"
        submission = create_submission_with_initial_version(
            author=self.user,
            title="DOC article",
            abstract="",
            journal=self.journal,
            article_type=self.article_type,
            file=SimpleUploadedFile("Статья.doc", source_bytes),
            defer_checks=True,
            mark_as_checking=False,
        )

        versions = list(submission.versions.order_by("version_number"))
        self.assertEqual(len(versions), 2)
        self.assertTrue(versions[0].file.name.endswith(".doc"))
        self.assertTrue(versions[1].file.name.endswith(".docx"))
        self.assertEqual(submission.current_version_id, versions[1].id)
        with versions[0].file.open("rb") as source:
            self.assertEqual(source.read(), source_bytes)
        self.assertIn("Рабочая DOCX-версия", versions[1].comment)
        mocked_convert.assert_called_once_with(source_bytes)

    @patch("apps.checks.services.queue_submission_checks")
    @patch(
        "apps.submissions.document_uploads.convert_legacy_doc_to_docx",
        return_value=_editable_docx(),
    )
    def test_new_doc_version_creates_source_and_current_docx(
        self,
        mocked_convert,
        mocked_queue,
    ):
        submission = create_submission_with_initial_version(
            author=self.user,
            title="Existing article",
            abstract="",
            journal=self.journal,
            article_type=self.article_type,
            file=SimpleUploadedFile("article.txt", b"initial"),
            defer_checks=True,
            mark_as_checking=False,
        )
        submission.status = SubmissionStatus.REVISION_REQUESTED
        submission.save(update_fields=["status", "updated_at"])

        current = add_submission_version(
            submission,
            self.user,
            SimpleUploadedFile(
                "Исправленная.doc",
                LEGACY_DOC_SIGNATURE + b"revision",
            ),
            comment="После замечаний",
        )

        self.assertEqual(current.version_number, 3)
        self.assertTrue(current.file.name.endswith(".docx"))
        self.assertEqual(submission.versions.count(), 3)
        self.assertTrue(
            submission.versions.get(version_number=2).file.name.endswith(".doc")
        )
        self.assertIn("После замечаний", current.comment)
        mocked_convert.assert_called_once()
        mocked_queue.assert_called_once()

    @patch(
        "apps.directory.formatting_templates.convert_legacy_doc_to_docx",
        return_value=_editable_docx(),
    )
    @patch(
        "apps.directory.formatting_templates.is_ai_configured",
        return_value=False,
    )
    def test_doc_template_uses_converted_docx_for_visual_rules(
        self,
        _mocked_ai_configured,
        mocked_convert,
    ):
        template = FormattingTemplate.objects.create(
            journal=self.journal,
            article_type=self.article_type,
            file=SimpleUploadedFile(
                "ОБРАЗЕЦ.doc",
                LEGACY_DOC_SIGNATURE + b"template",
            ),
            uploaded_by=self.user,
        )

        process_formatting_template(template)
        template.refresh_from_db()

        self.assertIn(template.analysis_status, {"ready", "partial"})
        self.assertAlmostEqual(
            template.extracted_rules["page"]["margins_cm"]["top"],
            2,
            places=1,
        )
        self.assertEqual(
            template.extracted_rules["body"]["font_family"],
            "Times New Roman",
        )
        self.assertAlmostEqual(
            template.extracted_rules["body"]["font_size_pt"],
            12,
            places=1,
        )
        mocked_convert.assert_called_once()
