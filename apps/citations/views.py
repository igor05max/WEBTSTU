import json
from io import BytesIO

from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.core.files.base import ContentFile
from django.http import FileResponse, HttpResponse, JsonResponse
from django.shortcuts import get_object_or_404, redirect, render
from django.views.decorators.http import require_GET, require_POST
from django.views.decorators.clickjacking import xframe_options_sameorigin
from docx import Document

from apps.citations.analysis import analyze_claims, document_snapshot, text_snapshot
from apps.citations.forms import CitationSearchForm
from apps.citations.index import ensure_index, get_index_status, search_claim
from apps.citations.rerank import rerank_claims
from apps.citations.workspaces import (
    apply_to_docx,
    create_workspace,
    load_workspace,
    prepare_docx_result,
    read_prepared_result,
)
from apps.submissions.models import Submission, SubmissionStatus
from apps.submissions.services import add_submission_version
from apps.submissions.template_processing import prepare_submission_template_by_id


def _read_form_source(form):
    submission = form.cleaned_data.get("submission")
    uploaded = form.cleaned_data.get("file")
    text = (form.cleaned_data.get("text") or "").strip()
    if submission is not None:
        version = submission.current_version
        with version.file.open("rb") as source:
            data = source.read()
        return data, version.file.name.rsplit("/", 1)[-1], submission.title, submission
    if uploaded is not None:
        return uploaded.read(), uploaded.name, uploaded.name, None
    return None, "pasted-text.txt", "Вставленный текст", None


def _submission_for_user(request, submission_id):
    submission = get_object_or_404(
        Submission.objects.select_related("author", "current_version"),
        pk=submission_id,
    )
    if submission.author_id != request.user.pk and not request.user.is_superuser:
        raise PermissionError("Этот материал недоступен.")
    return submission


def _workspace_submission(request, payload):
    submission_id = payload.get("submission_id")
    if not submission_id:
        raise ValueError("Рабочий набор не связан с материалом.")
    submission = _submission_for_user(request, submission_id)
    if submission.current_version_id != payload.get("source_version_id"):
        raise ValueError(
            "Версия материала уже изменилась. Запустите подбор источников заново."
        )
    return submission


@login_required
def workspace(request):
    result = None
    initial = {}
    selected_submission = None
    if request.method == "GET" and request.GET.get("submission"):
        initial["submission"] = request.GET.get("submission")
    form = CitationSearchForm(
        request.POST or None,
        request.FILES or None,
        user=request.user,
        initial=initial,
    )
    requested_submission_id = (
        request.POST.get("submission")
        if request.method == "POST"
        else request.GET.get("submission")
    )
    if str(requested_submission_id or "").isdigit():
        try:
            selected_submission = _submission_for_user(
                request,
                int(requested_submission_id),
            )
        except (PermissionError, Submission.DoesNotExist):
            selected_submission = None
    if request.method == "POST" and form.is_valid():
        file_bytes, file_name, source_title, selected_submission = _read_form_source(form)
        if file_bytes is None:
            snapshot = text_snapshot(form.cleaned_data["text"])
        else:
            snapshot = document_snapshot(file_bytes, file_name)
        if not (snapshot.get("text") or "").strip():
            form.add_error(
                None,
                snapshot.get("parse_error")
                or "Из документа не удалось извлечь текст для анализа.",
            )
        else:
            try:
                index_meta = ensure_index()
                analysis = analyze_claims(
                    snapshot,
                    max_claims=form.cleaned_data["max_claims"],
                )
                claims = analysis["claims"]
                for claim in claims:
                    claim["recommendations"] = search_claim(claim)
                rerank_claims(claims)
                workspace_payload = create_workspace(
                    user_id=request.user.pk,
                    file_bytes=file_bytes,
                    file_name=file_name,
                    snapshot=snapshot,
                    claims=claims,
                    index_status=index_meta,
                    submission_id=(
                        selected_submission.pk if selected_submission is not None else None
                    ),
                    source_version_id=(
                        selected_submission.current_version_id
                        if selected_submission is not None
                        else None
                    ),
                )
                result = {
                    "source_title": source_title,
                    "file_name": file_name,
                    "token": workspace_payload["token"],
                    "can_apply_docx": workspace_payload["suffix"] == ".docx",
                    "claims": claims,
                    "analysis": analysis,
                    "index": index_meta,
                    "submission_id": workspace_payload.get("submission_id"),
                    "total_recommendations": sum(
                        len(claim.get("recommendations") or []) for claim in claims
                    ),
                }
                if not claims:
                    messages.warning(
                        request,
                        "Утверждения без ссылок не обнаружены. Попробуйте передать введение или обзор литературы.",
                    )
            except Exception as exc:
                form.add_error(None, f"Поиск источников не завершён: {exc}")

    return render(
        request,
        "citations/workspace.html",
        {
            "form": form,
            "result": result,
            "index_status": get_index_status(),
            "selected_submission": selected_submission,
            "auto_analyze": bool(
                request.method == "GET"
                and request.GET.get("auto") == "1"
                and selected_submission is not None
            ),
        },
    )


@login_required
@require_POST
def apply_citations(request):
    token = str(request.POST.get("workspace_token") or "")
    try:
        selections = json.loads(request.POST.get("selections") or "[]")
        if not isinstance(selections, list):
            raise ValueError("Некорректный список выбранных источников.")
        output, file_name = apply_to_docx(
            user_id=request.user.pk,
            token=token,
            selections=selections,
        )
    except (ValueError, FileNotFoundError, json.JSONDecodeError) as exc:
        messages.error(request, str(exc))
        return render(
            request,
            "citations/workspace.html",
            {
                "form": CitationSearchForm(user=request.user),
                "result": None,
                "index_status": get_index_status(),
            },
            status=400,
        )
    return FileResponse(
        output,
        as_attachment=True,
        filename=file_name,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@login_required
@require_POST
def prepare_submission_result(request):
    token = str(request.POST.get("workspace_token") or "")
    try:
        selections = json.loads(request.POST.get("selections") or "[]")
        if not isinstance(selections, list):
            raise ValueError("Некорректный список выбранных источников.")
        payload = load_workspace(user_id=request.user.pk, token=token)
        _workspace_submission(request, payload)
        prepare_docx_result(
            user_id=request.user.pk,
            token=token,
            selections=selections,
        )
    except (ValueError, PermissionError, FileNotFoundError, json.JSONDecodeError) as exc:
        messages.error(request, str(exc))
        return redirect("citations:workspace")
    return redirect("citations:submission_result_preview", token=token)


@login_required
@require_GET
def submission_result_preview(request, token):
    try:
        payload, _result_bytes = read_prepared_result(
            user_id=request.user.pk,
            token=token,
        )
        submission = _workspace_submission(request, payload)
    except (ValueError, PermissionError, FileNotFoundError) as exc:
        messages.error(request, str(exc))
        return redirect("citations:workspace")
    return render(
        request,
        "citations/submission_result_preview.html",
        {
            "submission": submission,
            "workspace_token": token,
            "display_filename": payload.get("result_file_name") or "material-with-sources.docx",
        },
    )


@login_required
@require_GET
@xframe_options_sameorigin
def submission_result_content(request, token):
    try:
        payload, result_bytes = read_prepared_result(
            user_id=request.user.pk,
            token=token,
        )
        _workspace_submission(request, payload)
        document = Document(BytesIO(result_bytes))
        paragraphs = [
            paragraph.text
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        ]
        tables = [
            [[cell.text for cell in row.cells] for row in table.rows]
            for table in document.tables
        ]
    except (ValueError, PermissionError, FileNotFoundError) as exc:
        return HttpResponse(str(exc), status=404)
    except Exception:
        return HttpResponse(
            "Не удалось показать DOCX. Скачайте подготовленный файл для просмотра.",
            status=422,
        )
    return render(
        request,
        "citations/submission_result_content.html",
        {"paragraphs": paragraphs, "tables": tables},
    )


@login_required
@require_GET
def submission_result_download(request, token):
    try:
        payload, result_bytes = read_prepared_result(
            user_id=request.user.pk,
            token=token,
        )
        _workspace_submission(request, payload)
    except (ValueError, PermissionError, FileNotFoundError) as exc:
        messages.error(request, str(exc))
        return redirect("citations:workspace")
    return FileResponse(
        BytesIO(result_bytes),
        as_attachment=True,
        filename=payload.get("result_file_name") or "material-with-sources.docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@login_required
@require_POST
def use_submission_result(request, token):
    try:
        payload, result_bytes = read_prepared_result(
            user_id=request.user.pk,
            token=token,
        )
        submission = _workspace_submission(request, payload)
        if submission.status != SubmissionStatus.DRAFT:
            raise ValueError("Источники можно добавить только до запуска проверок.")
        if submission.formatting_template_id:
            prepare_submission_template_by_id(
                submission.pk,
                template_id=submission.formatting_template_id,
                expected_version_id=submission.current_version_id,
                start_checks=False,
            )
            submission.refresh_from_db()
        add_submission_version(
            submission,
            request.user,
            ContentFile(
                result_bytes,
                name=payload.get("result_file_name") or "material-with-sources.docx",
            ),
            comment="Добавлены выбранные источники из локальной RAG-системы.",
            expected_current_version_id=payload.get("source_version_id"),
        )
    except (ValueError, PermissionError, FileNotFoundError) as exc:
        messages.error(request, str(exc))
        return redirect("citations:workspace")
    messages.success(
        request,
        "Создана новая версия с выбранными источниками. Автоматические проверки запущены.",
    )
    return redirect("submissions:detail", pk=submission.pk)


@login_required
@require_GET
def index_status(request):
    return JsonResponse(get_index_status())
