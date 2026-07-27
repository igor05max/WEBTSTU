import json
from io import BytesIO
from pathlib import Path
from tempfile import TemporaryDirectory
from types import SimpleNamespace
from unittest.mock import patch

from django.contrib.auth import get_user_model
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TestCase, override_settings
from django.urls import reverse
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from apps.citations.analysis import analyze_claims, text_snapshot
from apps.citations.checks import build_citation_coverage_report
from apps.citations.forms import CitationSearchForm
from apps.citations.index import build_index, search_claim
from apps.citations.matching import (
    build_source_identity,
    claims_with_recommendations,
    remove_source_article,
)
from apps.citations.rerank import _remove_weak_results, rerank_claims
from apps.citations.workspaces import apply_to_docx, create_workspace
from apps.directory.models import ArticleType, Journal
from apps.submissions.models import SubmissionStatus
from apps.submissions.services import create_submission_with_initial_version


class CitationSystemTests(TestCase):
    def setUp(self):
        self.temp_dir = TemporaryDirectory()
        root = Path(self.temp_dir.name)
        self.corpus_root = root / "corpus"
        self.corpus_root.mkdir()
        self.index_path = root / "citation.sqlite3"
        self.workspace_root = root / "workspaces"
        self.settings_override = override_settings(
            CITATION_CORPUS_ROOT=self.corpus_root,
            CITATION_INDEX_PATH=self.index_path,
            CITATION_WORKSPACE_ROOT=self.workspace_root,
            CITATION_INDEX_AUTO_BUILD=True,
            CITATION_LLM_ANALYSIS_ENABLED=False,
            CITATION_LLM_RERANK_ENABLED=False,
            CITATION_EMBEDDING_MODEL="",
            CITATION_CHECK_MIN_TEXT_LENGTH=80,
        )
        self.settings_override.enable()
        self._build_corpus()
        build_index(corpus_root=self.corpus_root, index_path=self.index_path)
        self.user = get_user_model().objects.create_user(
            username="citation_author",
            password="1234",
        )

    def tearDown(self):
        self.settings_override.disable()
        self.temp_dir.cleanup()

    def _build_corpus(self):
        (self.corpus_root / "journal_articles_full_metadata.csv").write_text(
            "\n".join(
                [
                    "year;issue_display_name;section;article_id;title;authors;journal;article_year;pages;doi;edn;article_url;citation_elibrary",
                    "2024;Т. 1 № 1;Информатика;1001;НЕЙРОННЫЕ СЕТИ ДЛЯ АНАЛИЗА ИЗОБРАЖЕНИЙ;Иванов И.И.;Тестовый журнал;2024;1-9;10.1000/test.1;ABCDEF;https://example.test/1;Иванов И. И. Нейронные сети для анализа изображений // Тестовый журнал. 2024. № 1. С. 1-9.",
                    "2023;Т. 2 № 1;Химия;1002;ЭКСТРАКЦИЯ РАСТИТЕЛЬНОГО СЫРЬЯ;Петров П.П.;Тестовый журнал;2023;10-20;;;https://example.test/2;Петров П. П. Экстракция растительного сырья // Тестовый журнал. 2023. № 1. С. 10-20.",
                ]
            ),
            encoding="utf-8",
        )
        issue = self.corpus_root / "2024"
        issue.mkdir()
        (issue / "1001 - article.html").write_text(
            """
            <div id="abstract1">Свёрточные нейронные сети применяются для классификации
            медицинских изображений и повышают точность распознавания.</div>
            <p>Ключевые слова: нейронные сети; классификация; изображения</p>
            """,
            encoding="utf-8",
        )
        (issue / "1002 - article.html").write_text(
            '<div id="abstract1">Описан метод экстракции веществ из растений.</div>',
            encoding="utf-8",
        )

    def test_claim_analysis_and_hybrid_search(self):
        snapshot = text_snapshot(
            "Введение\nСвёрточные нейронные сети широко применяются для классификации "
            "медицинских изображений и обеспечивают высокую точность распознавания."
        )
        analysis = analyze_claims(snapshot, max_claims=3)

        self.assertEqual(len(analysis["claims"]), 1)
        self.assertEqual(analysis["claims"][0]["type"], "method")
        results = search_claim(analysis["claims"][0], limit=2)
        self.assertEqual(results[0]["article_id"], "1001")
        self.assertEqual(results[0]["doi"], "10.1000/test.1")
        self.assertTrue(results[0]["evidence"])

    def test_workspace_page_returns_claims_and_real_metadata(self):
        self.client.force_login(self.user)
        response = self.client.post(
            reverse("citations:workspace"),
            {
                "text": (
                    "Нейронные сети широко используются для классификации медицинских "
                    "изображений и позволяют повысить точность распознавания."
                ),
                "max_claims": 3,
            },
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "НЕЙРОННЫЕ СЕТИ ДЛЯ АНАЛИЗА ИЗОБРАЖЕНИЙ")
        self.assertContains(response, "10.1000/test.1")
        self.assertContains(response, "Фрагмент из вашей статьи")
        self.assertContains(response, "поставим сразу после этого фрагмента")
        self.assertContains(response, "Заключение локальной модели")
        self.assertContains(response, "Вариант 1")
        self.assertContains(response, "Показать подтверждающий фрагмент источника")
        self.assertContains(
            response,
            "Свёрточные нейронные сети применяются для классификации",
        )
        self.assertContains(response, "абзац 1")
        self.assertNotContains(response, "Почему рекомендуется")
        self.assertNotContains(response, "Поисковые запросы")

    def test_uploaded_file_wins_over_stale_pasted_text(self):
        document = Document()
        document.add_paragraph("Документ для проверки загрузки.")
        source = BytesIO()
        document.save(source)
        form = CitationSearchForm(
            data={
                "text": "Текст, оставшийся в форме от предыдущего выбора.",
                "max_claims": 3,
            },
            files={
                "file": SimpleUploadedFile(
                    "article.docx",
                    source.getvalue(),
                    content_type=(
                        "application/vnd.openxmlformats-officedocument."
                        "wordprocessingml.document"
                    ),
                )
            },
            user=self.user,
        )

        self.assertTrue(form.is_valid(), form.errors)
        self.assertEqual(form.cleaned_data["file"].name, "article.docx")
        self.assertEqual(form.cleaned_data["text"], "")
        self.assertIsNone(form.cleaned_data["submission"])

    def test_legacy_doc_is_converted_to_docx_before_search(self):
        document = Document()
        document.add_paragraph("Документ после конвертации.")
        converted = BytesIO()
        document.save(converted)
        with patch(
            "apps.citations.forms.convert_legacy_doc_to_docx",
            return_value=converted.getvalue(),
        ) as mocked_convert:
            form = CitationSearchForm(
                data={"text": "", "max_claims": 3},
                files={
                    "file": SimpleUploadedFile(
                        "Статья.doc",
                        bytes.fromhex("d0cf11e0a1b11ae1") + b"legacy-word",
                        content_type="application/msword",
                    )
                },
                user=self.user,
            )

            self.assertTrue(form.is_valid(), form.errors)

        mocked_convert.assert_called_once()
        self.assertEqual(form.cleaned_data["file"].name, "Статья.docx")
        self.assertGreater(form.cleaned_data["file"].size, 100)

    def test_source_form_explains_missing_input_without_reset_instruction(self):
        form = CitationSearchForm(
            data={"submission": "", "text": "", "max_claims": 3},
            user=self.user,
        )

        self.assertFalse(form.is_valid())
        self.assertIn(
            "Загрузите файл, выберите материал или вставьте текст.",
            form.non_field_errors(),
        )

    def test_russian_article_does_not_offer_fragments_from_english_translation(self):
        snapshot = text_snapshot(
            "Введение\n"
            "Выбор направления подготовки формулируется как задача ранжирования "
            "образовательных программ по многокомпонентному цифровому профилю абитуриента.\n"
            "Для поддержки принятия решений используются методы многокритериального "
            "анализа, позволяющие сопоставлять альтернативные образовательные траектории.\n"
            "The selection of a field of study is formulated as the task of ranking "
            "a fixed set of degree programmes according to a multi-component applicant "
            "digital profile."
        )

        analysis = analyze_claims(snapshot, max_claims=6)

        self.assertEqual(analysis["document_language"], "ru")
        self.assertTrue(analysis["claims"])
        self.assertTrue(
            all("The selection of a field" not in claim["text"] for claim in analysis["claims"])
        )
        self.assertTrue(all(claim["section"] == "Введение" for claim in analysis["claims"]))

    def test_source_article_is_removed_but_same_candidate_can_match_two_fragments(self):
        identity = build_source_identity(
            source_title="Собственная статья автора",
        )
        repeated_candidate = {
            "article_id": "1001",
            "title": "Общий источник для двух фрагментов",
        }
        claims = [
            {
                "id": "claim-1",
                "recommendations": [
                    {"article_id": "self", "title": "Собственная статья автора"},
                    dict(repeated_candidate),
                ],
            },
            {
                "id": "claim-2",
                "recommendations": [dict(repeated_candidate)],
            },
            {
                "id": "claim-3",
                "recommendations": [],
            },
        ]

        removed = remove_source_article(claims, identity)
        visible_claims = claims_with_recommendations(claims)

        self.assertEqual(removed, 1)
        self.assertEqual([claim["id"] for claim in visible_claims], ["claim-1", "claim-2"])
        self.assertEqual(
            [claim["recommendations"][0]["article_id"] for claim in visible_claims],
            ["1001", "1001"],
        )

    def test_uploaded_article_is_not_recommended_to_itself(self):
        document = Document()
        document.add_paragraph("НЕЙРОННЫЕ СЕТИ ДЛЯ АНАЛИЗА ИЗОБРАЖЕНИЙ")
        document.add_paragraph("Иванов И. И.")
        document.add_paragraph(
            "Свёрточные нейронные сети широко применяются для классификации "
            "медицинских изображений и обеспечивают высокую точность распознавания."
        )
        source = BytesIO()
        document.save(source)
        self.client.force_login(self.user)

        response = self.client.post(
            reverse("citations:workspace"),
            {
                "file": SimpleUploadedFile(
                    "same-article.docx",
                    source.getvalue(),
                    content_type=(
                        "application/vnd.openxmlformats-officedocument."
                        "wordprocessingml.document"
                    ),
                ),
                "text": "Этот текст остался в поле до выбора файла и должен быть очищен.",
                "max_claims": 3,
            },
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.context["result"]["claims"], [])
        self.assertEqual(response.context["result"]["analyzed_claim_count"], 1)
        self.assertContains(response, "Анализ выполнен")
        self.assertContains(response, "Файл успешно загружен и обработан")
        self.assertContains(response, "Слабые тематические совпадения")
        self.assertContains(response, "Подходящие новые источники не найдены")
        self.assertNotContains(response, "data-citation-source-form")

    def test_submission_check_contains_exact_citation_locations(self):
        submission = SimpleNamespace(
            title="Анализ медицинских изображений",
            abstract="Классификация изображений нейронными сетями.",
        )
        snapshot = text_snapshot(
            "Введение\nСвёрточные нейронные сети широко применяются для классификации "
            "медицинских изображений и обеспечивают высокую точность распознавания."
        )

        passed, payload = build_citation_coverage_report(
            submission,
            None,
            snapshot=snapshot,
            max_claims=3,
            results_per_claim=2,
        )

        self.assertTrue(passed)
        self.assertEqual(payload["check_code"], "article_recommendations")
        self.assertEqual(payload["metrics"]["claims_needing_citation"], 1)
        self.assertEqual(payload["metrics"]["minimum_score_percent"], 21)
        self.assertEqual(payload["issues"][0]["location"], "Введение, абзац 2")
        self.assertIn("Свёрточные нейронные сети", payload["issues"][0]["context_highlight"])
        self.assertTrue(payload["citation_claims"][0]["recommendations"])

    def test_scores_at_or_below_twenty_and_rejected_sources_are_removed(self):
        claims = [
            {
                "recommendations": [
                    {"title": "Нулевой", "score_percent": 0, "verdict": "partial"},
                    {"title": "Граница", "score_percent": 20, "verdict": "partial"},
                    {"title": "Выше границы", "score_percent": 21, "verdict": "partial"},
                    {"title": "Отклонённый", "score_percent": 87, "verdict": "not_supports"},
                    {"title": "Подходящий", "score_percent": 72, "verdict": "supports"},
                ]
            }
        ]

        filtered = _remove_weak_results(claims)

        self.assertEqual(
            [item["title"] for item in filtered[0]["recommendations"]],
            ["Выше границы", "Подходящий"],
        )

    def test_verified_results_are_topped_up_to_eight_best_candidates(self):
        claims = []
        rejected = []
        for claim_number in range(2):
            claim_id = f"claim-{claim_number}"
            recommendations = []
            for article_number in range(5):
                article_id = f"article-{claim_number}-{article_number}"
                recommendations.append(
                    {
                        "article_id": article_id,
                        "title": f"Статья {claim_number}-{article_number}",
                        "year": 2024,
                        "evidence": "Тематически близкий фрагмент публикации.",
                        "reason": "Тематическое совпадение.",
                        "hybrid_score": 0.3,
                        "semantic_score": 0.2,
                        "matched_terms": ["метод", "модель"],
                    }
                )
                rejected.append(
                    {
                        "id": f"{claim_id}::{article_id}",
                        "verdict": "not_supports",
                        "score": 0,
                        "reason": "Прямого подтверждения нет.",
                        "evidence": "Тематически близкий фрагмент публикации.",
                    }
                )
            claims.append(
                {
                    "id": claim_id,
                    "text": f"Проверяемое утверждение {claim_number}.",
                    "type": "topic",
                    "recommendations": recommendations,
                }
            )

        for accepted in rejected[:4]:
            accepted["verdict"] = "supports"
            accepted["score"] = 70

        with (
            override_settings(CITATION_LLM_RERANK_ENABLED=True),
            patch("apps.citations.rerank.is_ai_configured", return_value=True),
            patch(
                "apps.citations.rerank.generate_content",
                return_value=({}, "local-test-model"),
            ),
            patch(
                "apps.citations.rerank.extract_response_text",
                return_value=json.dumps({"items": rejected}, ensure_ascii=False),
            ),
        ):
            rerank_claims(claims, best_available_limit=8)

        restored = [
            item
            for claim in claims
            for item in (claim.get("recommendations") or [])
        ]
        self.assertEqual(len(restored), 8)
        self.assertTrue(all(item["score_percent"] > 20 for item in restored))
        self.assertEqual(sum(bool(item.get("best_available")) for item in restored), 4)
        self.assertEqual(
            sum(item.get("verdict") == "supports" for item in restored),
            4,
        )

    def test_apply_selected_source_to_docx(self):
        document = Document()
        document.add_paragraph(
            "Нейронные сети широко используются для анализа медицинских изображений. "
            "Следующее предложение должно остаться после маркера."
        )
        source = BytesIO()
        document.save(source)
        claim = {
            "id": "claim-1",
            "text": "Нейронные сети широко используются для анализа медицинских изображений.",
            "recommendations": [
                {
                    "article_id": "1001",
                    "title": "Нейронные сети для анализа изображений",
                    "citation": "Иванов И. И. Нейронные сети для анализа изображений. 2024.",
                }
            ],
        }
        payload = create_workspace(
            user_id=self.user.pk,
            file_bytes=source.getvalue(),
            file_name="article.docx",
            snapshot={"text": claim["text"]},
            claims=[claim],
            index_status={"ready": True},
        )

        output, name = apply_to_docx(
            user_id=self.user.pk,
            token=payload["token"],
            selections=[{"claim_id": "claim-1", "article_id": "1001"}],
        )
        result = Document(output)
        text = "\n".join(paragraph.text for paragraph in result.paragraphs)

        self.assertEqual(name, "article_with_citations.docx")
        self.assertIn(
            "медицинских изображений. [1] Следующее предложение",
            text,
        )
        self.assertIn("Список литературы", text)
        self.assertIn("Иванов И. И.", text)

    def test_one_source_can_support_two_fragments_without_duplicate_reference(self):
        document = Document()
        first_text = "Нейронные сети применяются для анализа изображений."
        second_text = "Такие модели используются для классификации объектов."
        document.add_paragraph(f"{first_text} {second_text}")
        source = BytesIO()
        document.save(source)
        recommendation = {
            "article_id": "1001",
            "title": "Нейронные сети для анализа изображений",
            "citation": "Иванов И. И. Нейронные сети для анализа изображений. 2024.",
        }
        claims = [
            {
                "id": "claim-1",
                "text": first_text,
                "recommendations": [dict(recommendation)],
            },
            {
                "id": "claim-2",
                "text": second_text,
                "recommendations": [dict(recommendation)],
            },
        ]
        payload = create_workspace(
            user_id=self.user.pk,
            file_bytes=source.getvalue(),
            file_name="two-fragments.docx",
            snapshot={"text": f"{first_text} {second_text}"},
            claims=claims,
            index_status={"ready": True},
        )

        output, _name = apply_to_docx(
            user_id=self.user.pk,
            token=payload["token"],
            selections=[
                {"claim_id": "claim-1", "article_id": "1001"},
                {"claim_id": "claim-2", "article_id": "1001"},
            ],
        )
        result = Document(output)
        text = "\n".join(paragraph.text for paragraph in result.paragraphs)

        self.assertIn(f"{first_text} [1]", text)
        self.assertIn(f"{second_text} [1]", text)
        self.assertEqual(text.count("Иванов И. И. Нейронные сети"), 1)

    def test_added_reference_continues_word_automatic_numbering(self):
        document = Document()
        document.add_paragraph(
            "YOLO используется для обнаружения людей [1]. MediaPipe определяет точки тела [2]. "
            "VideoMAE анализирует последовательность кадров [3]. "
            "Нейронные сети применяются для анализа движений человека."
        )
        document.add_paragraph("Список использованной литературы")
        document.add_paragraph("Первый источник.", style="List Number")
        document.add_paragraph("Второй источник.", style="List Number")
        reference_prototype = document.add_paragraph("Третий источник.", style="List Number")
        reference_prototype.alignment = WD_ALIGN_PARAGRAPH.LEFT
        reference_prototype.paragraph_format.left_indent = Cm(1)
        reference_prototype.paragraph_format.first_line_indent = Cm(-1)
        reference_prototype.paragraph_format.line_spacing = 1
        source = BytesIO()
        document.save(source)
        claim = {
            "id": "claim-1",
            "text": "Нейронные сети применяются для анализа движений человека.",
            "recommendations": [
                {
                    "article_id": "1001",
                    "title": "Система компьютерного зрения",
                    "citation": "Обухов А. Д. Система компьютерного зрения. 2023.",
                }
            ],
        }
        payload = create_workspace(
            user_id=self.user.pk,
            file_bytes=source.getvalue(),
            file_name="numbered.docx",
            snapshot={"text": claim["text"]},
            claims=[claim],
            index_status={"ready": True},
        )

        output, _name = apply_to_docx(
            user_id=self.user.pk,
            token=payload["token"],
            selections=[{"claim_id": "claim-1", "article_id": "1001"}],
        )
        result = Document(output)
        reference = result.paragraphs[-1]

        self.assertEqual(reference.text, "Обухов А. Д. Система компьютерного зрения. 2023.")
        self.assertEqual(reference.style.name, "List Number")
        self.assertNotIn("[4]", reference.text)
        self.assertIn("движений человека. [4]", result.paragraphs[0].text)
        self.assertEqual(reference.alignment, WD_ALIGN_PARAGRAPH.LEFT)
        self.assertAlmostEqual(reference.paragraph_format.left_indent.cm, 1, places=2)
        self.assertAlmostEqual(reference.paragraph_format.first_line_indent.cm, -1, places=2)
        self.assertEqual(reference.paragraph_format.line_spacing, 1)

    def test_added_reference_continues_visible_decimal_bibliography_numbering(self):
        document = Document()
        claim_text = "Нейронные сети применяются для анализа движений человека."
        document.add_paragraph(claim_text)
        document.add_paragraph("Список литературы")
        document.add_paragraph("1. Первый источник.")
        document.add_paragraph("2. Второй источник.")
        document.add_paragraph("3. Третий источник.")
        source = BytesIO()
        document.save(source)
        claim = {
            "id": "claim-decimal",
            "text": claim_text,
            "recommendations": [
                {
                    "article_id": "1001",
                    "title": "Система компьютерного зрения",
                    "citation": "Обухов А. Д. Система компьютерного зрения. 2023.",
                }
            ],
        }
        payload = create_workspace(
            user_id=self.user.pk,
            file_bytes=source.getvalue(),
            file_name="decimal.docx",
            snapshot={"text": claim_text},
            claims=[claim],
            index_status={"ready": True},
        )

        output, _name = apply_to_docx(
            user_id=self.user.pk,
            token=payload["token"],
            selections=[{"claim_id": "claim-decimal", "article_id": "1001"}],
        )
        result = Document(output)

        self.assertIn(f"{claim_text} [4]", result.paragraphs[0].text)
        self.assertEqual(
            result.paragraphs[-1].text,
            "4. Обухов А. Д. Система компьютерного зрения. 2023.",
        )

    def test_added_reference_stays_inside_bibliography_and_copies_run_format(self):
        document = Document()
        claim_text = "Нейронные сети применяются для анализа движений человека."
        document.add_paragraph(claim_text)
        document.add_paragraph("Список литературы", style="Heading 1")
        reference = document.add_paragraph()
        reference_run = reference.add_run("[1] Первый источник.")
        reference_run.font.name = "Times New Roman"
        reference_run.font.size = Pt(9)
        reference_run.italic = True
        reference.paragraph_format.first_line_indent = Cm(-0.5)
        document.add_paragraph("Приложение А", style="Heading 1")
        document.add_paragraph("Материалы приложения.")
        source = BytesIO()
        document.save(source)
        claim = {
            "id": "claim-before-appendix",
            "text": claim_text,
            "recommendations": [
                {
                    "article_id": "1001",
                    "title": "Система компьютерного зрения",
                    "citation": "Обухов А. Д. Система компьютерного зрения. 2023.",
                }
            ],
        }
        payload = create_workspace(
            user_id=self.user.pk,
            file_bytes=source.getvalue(),
            file_name="bibliography-before-appendix.docx",
            snapshot={"text": claim_text},
            claims=[claim],
            index_status={"ready": True},
        )

        output, _name = apply_to_docx(
            user_id=self.user.pk,
            token=payload["token"],
            selections=[
                {
                    "claim_id": claim["id"],
                    "article_id": "1001",
                }
            ],
        )
        result = Document(output)
        texts = [paragraph.text for paragraph in result.paragraphs]
        added_index = texts.index(
            "[2] Обухов А. Д. Система компьютерного зрения. 2023."
        )

        self.assertLess(added_index, texts.index("Приложение А"))
        added = result.paragraphs[added_index]
        self.assertEqual(added.runs[0].font.name, "Times New Roman")
        self.assertEqual(added.runs[0].font.size.pt, 9)
        self.assertTrue(added.runs[0].italic)
        self.assertAlmostEqual(
            added.paragraph_format.first_line_indent.cm,
            -0.5,
            places=2,
        )

    def test_marker_can_be_inserted_into_table_without_losing_media(self):
        import base64

        document = Document()
        claim_text = "Метод применяется для анализа экспериментальных данных."
        table = document.add_table(rows=1, cols=1)
        table.cell(0, 0).text = claim_text
        pixel = BytesIO(
            base64.b64decode(
                "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lE"
                "QVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
            )
        )
        document.add_picture(pixel)
        source = BytesIO()
        document.save(source)
        claim = {
            "id": "claim-table",
            "text": claim_text,
            "recommendations": [
                {
                    "article_id": "1002",
                    "title": "Анализ экспериментальных данных",
                    "citation": "Петров П. П. Анализ экспериментальных данных. 2025.",
                }
            ],
        }
        payload = create_workspace(
            user_id=self.user.pk,
            file_bytes=source.getvalue(),
            file_name="table-and-image.docx",
            snapshot={"text": claim_text},
            claims=[claim],
            index_status={"ready": True},
        )

        output, _name = apply_to_docx(
            user_id=self.user.pk,
            token=payload["token"],
            selections=[{"claim_id": "claim-table", "article_id": "1002"}],
        )
        result = Document(output)

        self.assertEqual(result.tables[0].cell(0, 0).text, f"{claim_text} [1]")
        self.assertEqual(len(result.inline_shapes), 1)

    def test_apply_source_when_docx_has_no_heading_style(self):
        document = Document()
        heading_style = document.styles["Heading 1"]
        heading_style._element.getparent().remove(heading_style._element)
        document.add_paragraph(
            "Нейронные сети используются для анализа медицинских изображений."
        )
        source = BytesIO()
        document.save(source)
        claim = {
            "id": "claim-without-heading-style",
            "text": "Нейронные сети используются для анализа медицинских изображений.",
            "recommendations": [
                {
                    "article_id": "1001",
                    "title": "Нейронные сети для анализа изображений",
                    "citation": "Иванов И. И. Нейронные сети для анализа изображений. 2024.",
                }
            ],
        }
        payload = create_workspace(
            user_id=self.user.pk,
            file_bytes=source.getvalue(),
            file_name="custom-styles.docx",
            snapshot={"text": claim["text"]},
            claims=[claim],
            index_status={"ready": True},
        )

        output, _name = apply_to_docx(
            user_id=self.user.pk,
            token=payload["token"],
            selections=[
                {
                    "claim_id": claim["id"],
                    "article_id": "1001",
                }
            ],
        )
        result = Document(output)
        headings = [
            paragraph.text
            for paragraph in result.paragraphs
            if paragraph.text == "Список литературы"
        ]
        self.assertEqual(headings, ["Список литературы"])

    def test_submission_source_stage_prepares_preview_and_new_version(self):
        document = Document()
        document.add_heading("Анализ медицинских изображений", level=1)
        document.add_paragraph("Иванов И. И.")
        document.add_paragraph(
            "Нейронные сети широко используются для классификации медицинских "
            "изображений и позволяют повысить точность распознавания."
        )
        source = BytesIO()
        document.save(source)
        journal = Journal.objects.create(name="Журнал RAG")
        article_type = ArticleType.objects.create(code="rag-article", name="Статья RAG")
        submission = create_submission_with_initial_version(
            author=self.user,
            title="Анализ медицинских изображений",
            abstract="Нейронные сети для классификации изображений.",
            document_authors="Иванов И. И.",
            keywords="нейронные сети; изображения",
            journal=journal,
            article_type=article_type,
            file=SimpleUploadedFile("article.docx", source.getvalue()),
            defer_checks=True,
            mark_as_checking=False,
        )
        self.client.force_login(self.user)

        page = self.client.get(
            f"{reverse('citations:workspace')}?submission={submission.pk}"
        )
        self.assertContains(page, "Данные материала")
        self.assertContains(page, str(self.user))
        self.assertContains(page, "Авторы не выбраны; показан отправитель материала.")
        self.assertNotContains(page, "Иванов И. И.")
        self.assertContains(page, 'name="max_claims" value="6"')
        self.assertContains(page, "Подбираем источники по тексту материала")
        self.assertNotContains(page, "Проанализировать и найти")

        search = self.client.post(
            reverse("citations:workspace"),
            {"submission": submission.pk, "max_claims": 3},
        )
        self.assertEqual(search.status_code, 200)
        result = search.context["result"]
        self.assertEqual(result["submission_id"], submission.pk)
        self.assertContains(search, 'data-site-loading-title="Добавляем источники"')
        claim = next(
            item for item in result["claims"] if item.get("recommendations")
        )
        article = claim["recommendations"][0]

        prepared = self.client.post(
            reverse("citations:prepare_submission_result"),
            {
                "workspace_token": result["token"],
                "selections": (
                    f'[{{"claim_id":"{claim["id"]}",'
                    f'"article_id":"{article["article_id"]}"}}]'
                ),
            },
        )
        self.assertRedirects(
            prepared,
            reverse("citations:submission_result_preview", args=[result["token"]]),
            fetch_redirect_response=False,
        )
        preview_page = self.client.get(
            reverse("citations:submission_result_preview", args=[result["token"]])
        )
        self.assertContains(
            preview_page,
            'data-site-loading-title="Сохраняем документ"',
        )
        with patch(
            "apps.citations.views.build_docx_bytes_pdf",
            return_value=b"%PDF-1.4\n" + (b"0" * 120),
        ):
            preview_content = self.client.get(
                reverse(
                    "citations:submission_result_content",
                    args=[result["token"]],
                )
            )
        self.assertEqual(preview_content.status_code, 200)
        self.assertEqual(preview_content["Content-Type"], "application/pdf")
        self.assertIn("inline", preview_content["Content-Disposition"])
        self.assertEqual(preview_content.content[:5], b"%PDF-")

        with patch("apps.checks.services.queue_submission_checks") as mocked_queue:
            applied = self.client.post(
                reverse("citations:use_submission_result", args=[result["token"]])
            )
        self.assertRedirects(
            applied,
            reverse("submissions:detail", args=[submission.pk]),
            fetch_redirect_response=False,
        )
        submission.refresh_from_db()
        self.assertEqual(submission.status, SubmissionStatus.DRAFT)
        self.assertEqual(submission.versions.count(), 2)
        mocked_queue.assert_called_once()

    def test_superuser_sees_selected_author_and_can_search_submission(self):
        coauthor = get_user_model().objects.create_user(
            username="manual_coauthor",
            first_name="Артём",
            last_name="Обухов",
            password="1234",
        )
        admin = get_user_model().objects.create_superuser(
            username="citation_admin",
            password="1234",
        )
        journal = Journal.objects.create(name="Журнал ручных авторов")
        article_type = ArticleType.objects.create(
            code="manual-authors-article",
            name="Статья с ручными авторами",
        )
        submission = create_submission_with_initial_version(
            author=self.user,
            title="Материал с выбранными авторами",
            abstract="Нейронные сети для классификации изображений.",
            document_authors="",
            keywords="нейронные сети; изображения",
            journal=journal,
            article_type=article_type,
            authors=[coauthor],
            file=SimpleUploadedFile(
                "article.txt",
                (
                    "Нейронные сети широко используются для классификации медицинских "
                    "изображений и позволяют повысить точность распознавания."
                ).encode(),
            ),
            defer_checks=True,
            mark_as_checking=False,
        )
        self.client.force_login(admin)

        page = self.client.get(
            f"{reverse('citations:workspace')}?submission={submission.pk}"
        )

        self.assertEqual(page.status_code, 200)
        self.assertEqual(
            page.context["selected_submission_author_data"],
            {"display": str(coauthor), "is_selected": True},
        )
        self.assertContains(page, str(coauthor))
        self.assertContains(page, "Указаны в поле «Авторы» материала.")
        self.assertNotContains(page, "показан отправитель материала")
        self.assertNotContains(page, "Не удалось распознать")

        search = self.client.post(
            reverse("citations:workspace"),
            {"submission": submission.pk, "max_claims": 3},
        )

        self.assertEqual(search.status_code, 200)
        self.assertIsNotNone(search.context["result"])
        self.assertFalse(search.context["form"].errors)
