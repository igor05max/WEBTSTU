import json
import re
import uuid
from copy import deepcopy
from io import BytesIO
from pathlib import Path

from django.conf import settings
from docx import Document
from document_template_engine import (
    DocxPreservationError,
    assert_docx_payload_preserved,
)


TOKEN_RE = re.compile(r"^[0-9a-f]{32}$")
CITATION_BLOCK_RE = re.compile(r"\[([\d\s,;]+)\]")
BRACKETED_REFERENCE_RE = re.compile(r"^\s*\[(\d+)\]\s*")
DECIMAL_REFERENCE_RE = re.compile(r"^\s*(\d+)([.)])\s+")


def _workspace_dir(user_id, token):
    if not TOKEN_RE.fullmatch(str(token or "")):
        raise ValueError("Некорректный идентификатор рабочего набора.")
    return Path(settings.CITATION_WORKSPACE_ROOT) / str(int(user_id)) / token


def create_workspace(
    *,
    user_id,
    file_bytes,
    file_name,
    snapshot,
    claims,
    index_status,
    submission_id=None,
    source_version_id=None,
):
    token = uuid.uuid4().hex
    directory = _workspace_dir(user_id, token)
    directory.mkdir(parents=True, exist_ok=False)
    suffix = Path(file_name or "").suffix.casefold()
    source_name = f"source{suffix}" if file_bytes is not None and suffix else ""
    if source_name:
        (directory / source_name).write_bytes(file_bytes)
    payload = {
        "token": token,
        "file_name": Path(file_name or "document").name,
        "suffix": suffix,
        "source_name": source_name,
        "claims": claims,
        "index_status": index_status,
        "text_length": len(snapshot.get("text") or ""),
        "submission_id": submission_id,
        "source_version_id": source_version_id,
    }
    (directory / "workspace.json").write_text(
        json.dumps(payload, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    return payload


def load_workspace(*, user_id, token):
    directory = _workspace_dir(user_id, token)
    payload_path = directory / "workspace.json"
    if not payload_path.exists():
        raise FileNotFoundError("Рабочий набор не найден или уже удалён.")
    payload = json.loads(payload_path.read_text(encoding="utf-8"))
    payload["_directory"] = directory
    return payload


def _selected_sources(payload, selections):
    claims = {claim["id"]: claim for claim in payload.get("claims") or []}
    selected = []
    article_numbers = {}
    for item in selections:
        claim = claims.get(str(item.get("claim_id") or ""))
        if claim is None:
            continue
        article_id = str(item.get("article_id") or "")
        result = next(
            (
                candidate
                for candidate in claim.get("recommendations") or []
                if str(candidate.get("article_id")) == article_id
            ),
            None,
        )
        if result is None:
            continue
        if article_id not in article_numbers:
            article_numbers[article_id] = len(article_numbers)
        selected.append((claim, result, article_id))
    return selected, article_numbers


def _insert_marker_after_claim(paragraph, claim_text, marker):
    normalized_paragraph = " ".join(paragraph.text.casefold().replace("ё", "е").split())
    normalized_claim_text = " ".join(claim_text.casefold().replace("ё", "е").split())
    if f"{normalized_claim_text} {marker}" in normalized_paragraph:
        return True
    full_text = "".join(run.text for run in paragraph.runs)
    normalized_chars = []
    source_positions = []
    previous_was_space = True
    for source_index, character in enumerate(full_text):
        if character.isspace():
            if not previous_was_space:
                normalized_chars.append(" ")
                source_positions.append(source_index)
            previous_was_space = True
            continue
        folded = character.casefold().replace("ё", "е")
        for folded_character in folded:
            normalized_chars.append(folded_character)
            source_positions.append(source_index)
        previous_was_space = False
    if normalized_chars and normalized_chars[-1] == " ":
        normalized_chars.pop()
        source_positions.pop()
    normalized_text = "".join(normalized_chars)
    normalized_claim = " ".join(claim_text.casefold().replace("ё", "е").split())
    position = normalized_text.find(normalized_claim)
    if position < 0 or not normalized_claim:
        return False
    normalized_end = position + len(normalized_claim) - 1
    insertion_at = source_positions[normalized_end] + 1
    offset = 0
    for run in paragraph.runs:
        run_end = offset + len(run.text)
        if insertion_at <= run_end:
            local_offset = max(0, insertion_at - offset)
            run.text = (
                run.text[:local_offset]
                + f" {marker}"
                + run.text[local_offset:]
            )
            return True
        offset = run_end
    return False


def _has_automatic_numbering(paragraph):
    paragraph_properties = paragraph._p.pPr
    if paragraph_properties is not None and paragraph_properties.numPr is not None:
        return True
    style = paragraph.style
    style_properties = style.element.pPr if style is not None else None
    return style_properties is not None and style_properties.numPr is not None


def _copy_paragraph_format(source, target):
    source_properties = source._p.pPr
    if source_properties is None:
        return
    target_properties = target._p.pPr
    if target_properties is not None:
        target._p.remove(target_properties)
    target._p.insert(0, deepcopy(source_properties))


def _copy_run_format(source, target):
    if source is None or source._r.rPr is None:
        return
    if target._r.rPr is not None:
        target._r.remove(target._r.rPr)
    target._r.insert(0, deepcopy(source._r.rPr))


def _is_section_heading(paragraph):
    style_name = (
        (paragraph.style.name or "").casefold()
        if paragraph.style is not None
        else ""
    )
    return "heading" in style_name or "заголов" in style_name


def _bibliography_references(document, heading):
    paragraphs = document.paragraphs
    heading_index = next(
        (
            index
            for index, paragraph in enumerate(paragraphs)
            if paragraph._p is heading._p
        ),
        len(paragraphs) - 1,
    )
    references = []
    for paragraph in paragraphs[heading_index + 1 :]:
        if not paragraph.text.strip():
            continue
        if _is_section_heading(paragraph):
            break
        references.append(paragraph)
    return references


def _prefix_paragraph(paragraph, prefix):
    if paragraph.runs:
        paragraph.runs[0].text = prefix + paragraph.runs[0].text
    else:
        paragraph.add_run(prefix)


def _bibliography_layout(document, heading):
    if heading is None:
        return {
            "prototype": None,
            "anchor": None,
            "mode": "bracket",
            "delimiter": "]",
        }
    references = _bibliography_references(document, heading)
    if not references:
        return {
            "prototype": None,
            "anchor": heading,
            "mode": "bracket",
            "delimiter": "]",
        }

    prototype = references[-1]
    if any(_has_automatic_numbering(paragraph) for paragraph in references):
        numbered_prototype = next(
            paragraph for paragraph in reversed(references) if _has_automatic_numbering(paragraph)
        )
        return {
            "prototype": numbered_prototype,
            "anchor": prototype,
            "mode": "automatic",
            "delimiter": "",
        }

    bracketed = [BRACKETED_REFERENCE_RE.match(paragraph.text) for paragraph in references]
    if all(match is not None for match in bracketed):
        return {
            "prototype": prototype,
            "anchor": prototype,
            "mode": "bracket",
            "delimiter": "]",
        }

    decimal = [DECIMAL_REFERENCE_RE.match(paragraph.text) for paragraph in references]
    if all(match is not None for match in decimal):
        return {
            "prototype": prototype,
            "anchor": prototype,
            "mode": "decimal",
            "delimiter": decimal[-1].group(2),
        }

    # Plain paragraphs after a bibliography heading have no visible mapping to
    # in-text [n] references. Normalize the whole list once so the newly added
    # source does not become the only visibly numbered item.
    if all(match is None for match in bracketed) and all(match is None for match in decimal):
        for number, paragraph in enumerate(references, start=1):
            _prefix_paragraph(paragraph, f"[{number}] ")
        return {
            "prototype": prototype,
            "anchor": prototype,
            "mode": "bracket",
            "delimiter": "]",
        }

    return {
        "prototype": prototype,
        "anchor": prototype,
        "mode": "bracket",
        "delimiter": "]",
    }


def _append_bibliography_entry(document, layout, number, citation):
    prototype = layout.get("prototype")
    paragraph = document.add_paragraph()
    if prototype is not None:
        _copy_paragraph_format(prototype, paragraph)

    mode = layout.get("mode")
    if mode == "automatic":
        text = citation
    elif mode == "decimal":
        text = f"{number}{layout.get('delimiter') or '.'} {citation}"
    else:
        text = f"[{number}] {citation}"
    run = paragraph.add_run(text)
    if prototype is not None and prototype.runs:
        _copy_run_format(prototype.runs[0], run)

    anchor = layout.get("anchor")
    if anchor is not None:
        anchor._p.addnext(paragraph._p)
    layout["anchor"] = paragraph
    layout["prototype"] = paragraph
    return paragraph


def _iter_document_paragraphs(document):
    """Yield body and table-cell paragraphs once, in document order where possible."""

    seen = set()
    for paragraph in document.paragraphs:
        seen.add(id(paragraph._p))
        yield paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    identity = id(paragraph._p)
                    if identity in seen:
                        continue
                    seen.add(identity)
                    yield paragraph


def _existing_reference_numbers(document, bibliography_heading):
    numbers = [
        int(value)
        for paragraph in _iter_document_paragraphs(document)
        for block in CITATION_BLOCK_RE.findall(paragraph.text)
        for value in re.findall(r"\d+", block)
    ]
    if bibliography_heading is None:
        return numbers
    paragraphs = document.paragraphs
    heading_index = next(
        (
            index
            for index, paragraph in enumerate(paragraphs)
            if paragraph._p is bibliography_heading._p
        ),
        len(paragraphs),
    )
    for paragraph in paragraphs[heading_index + 1 :]:
        bracketed = BRACKETED_REFERENCE_RE.match(paragraph.text)
        decimal = DECIMAL_REFERENCE_RE.match(paragraph.text)
        if bracketed is not None:
            numbers.append(int(bracketed.group(1)))
        elif decimal is not None:
            numbers.append(int(decimal.group(1)))
    return numbers


def apply_to_docx(*, user_id, token, selections):
    payload = load_workspace(user_id=user_id, token=token)
    if payload.get("suffix") != ".docx" or not payload.get("source_name"):
        raise ValueError("Автоматическая вставка доступна только для исходного DOCX.")
    selected, article_numbers = _selected_sources(payload, selections)
    if not selected:
        raise ValueError("Не выбрано ни одного источника.")

    source_path = payload["_directory"] / payload["source_name"]
    document = Document(source_path)
    bibliography_heading = next(
        (
            paragraph
            for paragraph in document.paragraphs
            if paragraph.text.strip().casefold()
            in {
                "список литературы",
                "список использованной литературы",
                "список использованных источников",
                "библиографический список",
                "литература",
                "references",
            }
        ),
        None,
    )
    existing_numbers = _existing_reference_numbers(document, bibliography_heading)
    bibliography_reference_count = (
        len(_bibliography_references(document, bibliography_heading))
        if bibliography_heading is not None
        else 0
    )
    start_number = max(
        [*existing_numbers, bibliography_reference_count],
        default=0,
    ) + 1
    number_by_article = {
        article_id: start_number + offset
        for article_id, offset in article_numbers.items()
    }

    selected_by_claim = {}
    for claim, _result, article_id in selected:
        claim_id = str(claim.get("id") or "")
        group = selected_by_claim.setdefault(
            claim_id,
            {"claim": claim, "numbers": []},
        )
        number = number_by_article[article_id]
        if number not in group["numbers"]:
            group["numbers"].append(number)

    document_paragraphs = list(_iter_document_paragraphs(document))
    missing_markers = []
    for group in selected_by_claim.values():
        claim = group["claim"]
        marker = f"[{', '.join(str(value) for value in sorted(group['numbers']))}]"
        target_text = " ".join(str(claim.get("text") or "").split())
        inserted = False
        for paragraph in document_paragraphs:
            normalized = " ".join(paragraph.text.split())
            if target_text and target_text in normalized:
                if not _insert_marker_after_claim(paragraph, target_text, marker):
                    paragraph.add_run(f" {marker}")
                inserted = True
                break
        if not inserted:
            missing_markers.append(target_text[:160])

    if missing_markers:
        raise ValueError(
            "Не удалось поставить ссылку рядом с выбранным фрагментом: "
            f"«{missing_markers[0]}». Документ не был сформирован."
        )

    if bibliography_heading is None:
        try:
            heading_style = document.styles["Heading 1"]
        except KeyError:
            bibliography_heading = document.add_paragraph()
            bibliography_heading.add_run("Список литературы").bold = True
        else:
            bibliography_heading = document.add_paragraph(
                "Список литературы",
                style=heading_style,
            )
    bibliography_layout = _bibliography_layout(document, bibliography_heading)

    unique_results = {}
    for _claim, result, article_id in selected:
        unique_results.setdefault(article_id, result)
    for article_id, result in sorted(
        unique_results.items(),
        key=lambda item: number_by_article[item[0]],
    ):
        number = number_by_article[article_id]
        citation = str(result.get("citation") or result.get("title") or "").strip()
        _append_bibliography_entry(
            document,
            bibliography_layout,
            number,
            citation,
        )

    output = BytesIO()
    document.save(output)
    output.seek(0)
    try:
        assert_docx_payload_preserved(source_path.read_bytes(), output.getvalue())
    except DocxPreservationError as exc:
        raise ValueError(str(exc)) from exc
    original_stem = Path(payload.get("file_name") or "article").stem
    return output, f"{original_stem}_with_citations.docx"


def prepare_docx_result(*, user_id, token, selections):
    output, file_name = apply_to_docx(
        user_id=user_id,
        token=token,
        selections=selections,
    )
    payload = load_workspace(user_id=user_id, token=token)
    selected, _article_numbers = _selected_sources(payload, selections)
    result_name = "result.docx"
    (payload["_directory"] / result_name).write_bytes(output.getvalue())
    payload["result_name"] = result_name
    payload["result_file_name"] = file_name
    payload["selections"] = [
        {
            "claim_id": str(claim.get("id") or ""),
            "article_id": article_id,
        }
        for claim, _result, article_id in selected
    ]
    directory = payload.pop("_directory")
    (directory / "workspace.json").write_text(
        json.dumps(payload, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    return payload


def read_prepared_result(*, user_id, token):
    payload = load_workspace(user_id=user_id, token=token)
    result_name = str(payload.get("result_name") or "")
    if not result_name:
        raise FileNotFoundError("Сначала подготовьте документ с выбранными источниками.")
    result_path = payload["_directory"] / result_name
    if not result_path.exists():
        raise FileNotFoundError("Подготовленный документ больше недоступен.")
    return payload, result_path.read_bytes()
