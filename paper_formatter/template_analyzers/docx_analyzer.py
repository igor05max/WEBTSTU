from __future__ import annotations

from collections import Counter
from pathlib import Path
import re
from statistics import median

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from paper_formatter.exceptions import TemplateAnalysisError
from paper_formatter.models import (
    HeadingStyleProfile,
    LatexTemplateProfile,
    PageLayout,
    TemplateProfile,
    TypographyProfile,
)
from paper_formatter.template_analyzers.base import TemplateAnalyzer
from paper_formatter.renderers.docx_template_styles import DocxTemplateStyleMap


_ALIGNMENTS = {
    WD_ALIGN_PARAGRAPH.LEFT: "left",
    WD_ALIGN_PARAGRAPH.CENTER: "center",
    WD_ALIGN_PARAGRAPH.RIGHT: "right",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
}


class DocxTemplateAnalyzer(TemplateAnalyzer):
    def analyze(self, source: Path) -> TemplateProfile:
        source = Path(source).resolve()
        try:
            document = Document(source)
        except Exception as exc:
            raise TemplateAnalysisError(f"Не удалось открыть DOCX-образец: {exc}") from exc

        page = self._page_layout(document)
        typography = self._typography(document)
        headings = self._headings(document)
        style_map = DocxTemplateStyleMap.from_document(document, source_path=source)
        evidence: dict[str, str | float | int | bool] = {
            "sections": len(document.sections),
            "paragraphs": len(document.paragraphs),
            "tables": len(document.tables),
            "maximum_columns": max(
                self._section_columns(section)[0]
                for section in document.sections
            ),
            **style_map.as_evidence(),
        }
        return TemplateProfile(
            name=source.stem,
            source_path=str(source),
            source_type="docx",
            confidence=0.82,
            page=page,
            typography=typography,
            headings=headings,
            figure_width_fraction=self._figure_width_fraction(document, page),
            latex=LatexTemplateProfile(
                document_class="article",
                class_options=[f"{max(8, round(typography.main_size_pt))}pt"],
            ),
            evidence=evidence,
        )

    def _page_layout(self, document: Document) -> PageLayout:
        first_section = document.sections[0]
        section_data = [
            (section, *self._section_columns(section))
            for section in document.sections
        ]
        maximum_columns = max(columns for _, columns, _ in section_data)
        # DOCX journal templates commonly alternate one-column full-width
        # figures/tables with the normal multi-column body.  The first section
        # is often a one-column title zone, so it is not representative.
        section, columns, column_gap_mm = next(
            item for item in section_data if item[1] == maximum_columns
        )
        width_mm = self._emu_to_mm(section.page_width)
        height_mm = self._emu_to_mm(section.page_height)
        paper_size = "a4paper"
        if width_mm and height_mm and abs(width_mm - 215.9) < 3 and abs(height_mm - 279.4) < 3:
            paper_size = "letterpaper"
        body_left = self._emu_to_mm(section.left_margin) or 20.0
        body_right = self._emu_to_mm(section.right_margin) or 20.0
        title_left = self._emu_to_mm(first_section.left_margin)
        title_right = self._emu_to_mm(first_section.right_margin)
        return PageLayout(
            paper_size=paper_size,
            width_mm=width_mm,
            height_mm=height_mm,
            margin_top_mm=self._emu_to_mm(section.top_margin) or 20.0,
            margin_right_mm=body_right,
            margin_bottom_mm=self._emu_to_mm(section.bottom_margin) or 20.0,
            margin_left_mm=body_left,
            title_margin_left_mm=(
                title_left
                if columns > 1 and title_left is not None and abs(title_left - body_left) > 0.1
                else None
            ),
            title_margin_right_mm=(
                title_right
                if columns > 1 and title_right is not None and abs(title_right - body_right) > 0.1
                else None
            ),
            columns=columns,
            column_gap_mm=column_gap_mm,
            header_distance_mm=self._emu_to_mm(section.header_distance),
            footer_distance_mm=self._emu_to_mm(section.footer_distance),
        )

    def _typography(self, document: Document) -> TypographyProfile:
        normal = self._body_style(document) or document.styles["Normal"]
        font_name = normal.font.name
        size_pt = normal.font.size.pt if normal.font.size else None
        observed_fonts: Counter[str] = Counter()
        observed_sizes: Counter[float] = Counter()
        for paragraph in document.paragraphs[:400]:
            for run in paragraph.runs:
                if run.font.name:
                    observed_fonts[run.font.name] += max(1, len(run.text))
                if run.font.size:
                    observed_sizes[round(run.font.size.pt, 1)] += max(1, len(run.text))
        if not font_name and observed_fonts:
            font_name = observed_fonts.most_common(1)[0][0]
        if size_pt is None and observed_sizes:
            size_pt = observed_sizes.most_common(1)[0][0]

        title_paragraph = self._title_paragraph(document)
        title_style = title_paragraph.style if title_paragraph is not None else None
        title_size = self._paragraph_size(title_paragraph)
        title_bold = self._paragraph_bold(title_paragraph)
        author_paragraph = self._front_matter_paragraph(document, "author")
        affiliation_paragraph = self._front_matter_paragraph(document, "affiliation")
        is_mdpi_front_matter = bool(
            title_style is not None and title_style.name.lower().startswith("mdpi_")
        )
        front_alignment_default = "left" if is_mdpi_front_matter else "center"
        caption = document.styles["Caption"] if "Caption" in document.styles else None
        normal_format = normal.paragraph_format
        line_spacing, line_spacing_pt, line_spacing_rule = self._line_spacing(normal)
        return TypographyProfile(
            main_font=font_name or "Times New Roman",
            main_size_pt=size_pt or 12.0,
            line_spacing=line_spacing,
            line_spacing_pt=line_spacing_pt,
            line_spacing_rule=line_spacing_rule,
            first_line_indent_mm=(
                round(normal_format.first_line_indent.mm, 2)
                if normal_format.first_line_indent is not None
                else 12.5
            ),
            paragraph_space_before_pt=(
                normal_format.space_before.pt
                if normal_format.space_before is not None
                else 0.0
            ),
            paragraph_space_after_pt=(
                normal_format.space_after.pt
                if normal_format.space_after is not None
                else 0.0
            ),
            paragraph_alignment=_ALIGNMENTS.get(
                normal_format.alignment,
                "justify",
            ),
            title_font=title_style.font.name if title_style is not None else None,
            title_size_pt=title_size,
            title_bold=title_bold,
            title_alignment=self._paragraph_alignment(
                title_paragraph,
                front_alignment_default,
            ),
            author_size_pt=self._paragraph_size(author_paragraph),
            author_alignment=self._paragraph_alignment(
                author_paragraph,
                front_alignment_default,
            ),
            affiliation_alignment=self._paragraph_alignment(
                affiliation_paragraph,
                front_alignment_default,
            ),
            abstract_size_pt=self._front_matter_size(document, 6),
            caption_size_pt=(
                caption.font.size.pt
                if caption is not None and caption.font.size is not None
                else None
            ),
        )

    def _headings(self, document: Document) -> list[HeadingStyleProfile]:
        result: list[HeadingStyleProfile] = []
        for level in range(1, 7):
            style = self._heading_style(document, level)
            if style is None:
                continue
            paragraph_format = style.paragraph_format
            result.append(
                HeadingStyleProfile(
                    level=level,
                    font=style.font.name,
                    size_pt=style.font.size.pt if style.font.size else None,
                    bold=bool(style.font.bold) if style.font.bold is not None else True,
                    italic=bool(style.font.italic) if style.font.italic is not None else False,
                    alignment=_ALIGNMENTS.get(paragraph_format.alignment, "left"),
                    space_before_pt=(
                        paragraph_format.space_before.pt
                        if paragraph_format.space_before
                        else None
                    ),
                    space_after_pt=(
                        paragraph_format.space_after.pt
                        if paragraph_format.space_after
                        else None
                    ),
                )
            )
        return result


    @staticmethod
    def _body_style(document: Document):
        style_map = DocxTemplateStyleMap.from_document(document)
        style_name = style_map.paragraph("body")
        if style_name and style_name in document.styles:
            return document.styles[style_name]
        return None

    @staticmethod
    def _line_spacing(style) -> tuple[float, float | None, str | None]:
        """Return a line multiplier or an absolute point value from Word XML.

        python-docx exposes ``atLeast`` and ``exact`` spacing as large
        EMU/Twips integers. Reusing that integer as a multiplier makes output
        DOCX files invalidly spaced and can prevent rendering altogether.
        """

        paragraph_properties = style.element.pPr
        spacing = (
            paragraph_properties.find(qn("w:spacing"))
            if paragraph_properties is not None
            else None
        )
        if spacing is None:
            return 1.15, None, None
        raw_line = spacing.get(qn("w:line"))
        line_rule = (spacing.get(qn("w:lineRule")) or "auto").lower()
        if raw_line and raw_line.isdigit():
            value = int(raw_line)
            if line_rule in {"atleast", "exact"}:
                return 1.0, value / 20, "atLeast" if line_rule == "atleast" else "exact"
            return value / 240, None, "auto"
        value = style.paragraph_format.line_spacing
        return (float(value), None, None) if isinstance(value, float) else (1.15, None, None)

    @staticmethod
    def _figure_width_fraction(
        document: Document,
        page: PageLayout,
    ) -> float:
        widths_mm: list[float] = []
        for blip in document.element.body.xpath(".//a:blip"):
            extents = blip.xpath(
                "ancestor::*[local-name()='inline' or local-name()='anchor'][1]"
                "/*[local-name()='extent']"
            )
            if not extents:
                continue
            raw = extents[0].get("cx")
            if raw and raw.isdigit():
                widths_mm.append(int(raw) / 36000)
        usable_width = (
            (page.width_mm or 210.0)
            - page.margin_left_mm
            - page.margin_right_mm
        )
        if not widths_mm or usable_width <= 0:
            return 0.9
        return max(0.2, min(0.98, median(widths_mm) / usable_width))

    @staticmethod
    def _section_columns(section) -> tuple[int, float | None]:
        columns = 1
        column_gap_mm: float | None = None
        cols = section._sectPr.xpath("./w:cols")
        if cols:
            count = cols[-1].get(qn("w:num"))
            gap = cols[-1].get(qn("w:space"))
            if count and count.isdigit():
                columns = max(1, min(int(count), 4))
            if gap and gap.isdigit():
                column_gap_mm = int(gap) * 25.4 / 1440
        return columns, column_gap_mm

    @staticmethod
    def _title_paragraph(document: Document):
        for paragraph in document.paragraphs[:20]:
            style_name = (paragraph.style.name if paragraph.style else "").lower()
            if re.search(r"(?:^|[_-])title$", style_name) and paragraph.text.strip():
                return paragraph
        for paragraph in document.paragraphs[:20]:
            text = paragraph.text.strip()
            if (
                paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
                and 15 <= len(text) <= 500
                and any(character.isalpha() for character in text)
                and text.upper() == text
            ):
                return paragraph
        return None

    @staticmethod
    def _front_matter_paragraph(document: Document, role: str):
        patterns = {
            "author": r"(?:^|[_-])author(?:names)?$",
            "affiliation": r"(?:^|[_-])affiliation$",
        }
        pattern = patterns[role]
        for paragraph in document.paragraphs[:20]:
            style_name = (paragraph.style.name if paragraph.style else "").lower()
            if re.search(pattern, style_name) and paragraph.text.strip():
                return paragraph
        return None

    @staticmethod
    def _paragraph_alignment(paragraph, default: str) -> str:
        if paragraph is None:
            return default
        alignment = paragraph.alignment
        if alignment is None and paragraph.style is not None:
            alignment = paragraph.style.paragraph_format.alignment
        return _ALIGNMENTS.get(alignment, default)

    @staticmethod
    def _paragraph_size(paragraph) -> float | None:
        if paragraph is None:
            return None
        weighted: Counter[float] = Counter()
        for run in paragraph.runs:
            if run.font.size is not None:
                weighted[round(run.font.size.pt, 1)] += max(1, len(run.text))
        if weighted:
            return weighted.most_common(1)[0][0]
        if paragraph.style is not None and paragraph.style.font.size is not None:
            return paragraph.style.font.size.pt
        return None

    @staticmethod
    def _paragraph_bold(paragraph) -> bool:
        if paragraph is None:
            return True
        weighted = Counter()
        for run in paragraph.runs:
            value = run.bold
            if value is None and paragraph.style is not None:
                value = paragraph.style.font.bold
            weighted[bool(value)] += max(1, len(run.text))
        return weighted.most_common(1)[0][0] if weighted else True

    @classmethod
    def _front_matter_size(cls, document: Document, index: int) -> float | None:
        if index >= len(document.paragraphs):
            return None
        paragraph = document.paragraphs[index]
        return cls._paragraph_size(paragraph) or (
            document.styles["Normal"].font.size.pt
            if document.styles["Normal"].font.size is not None
            else None
        )

    @staticmethod
    def _heading_style(document: Document, level: int):
        custom_pattern = re.compile(
            rf"(?:article|paper|journal)[ _-]*heading[ _-]*{level}$",
            re.IGNORECASE,
        )
        for style in document.styles:
            normalized = style.name.replace(" ", "")
            if custom_pattern.search(normalized):
                return style
        for style in document.styles:
            normalized = style.name.replace(" ", "")
            if re.search(
                rf"(?:^|[_-])heading[ _-]*{level}$",
                normalized,
                re.IGNORECASE,
            ) and normalized.lower() != f"heading{level}":
                return style
        canonical = f"Heading {level}"
        return document.styles[canonical] if canonical in document.styles else None

    @staticmethod
    def _emu_to_mm(value: int | None) -> float | None:
        return round(value / 36000, 2) if value is not None else None
