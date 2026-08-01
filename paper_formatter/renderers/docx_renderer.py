from __future__ import annotations

from copy import deepcopy
import posixpath
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.opc.part import Part
from docx.shared import Mm, Pt, RGBColor

from paper_formatter.models import (
    ArticleIR,
    EquationBlock,
    FigureBlock,
    ListItemBlock,
    ParagraphBlock,
    SectionBlock,
    TableBlock,
    TemplateProfile,
    TextRun,
)
from paper_formatter.renderers.docx_template_styles import DocxTemplateStyleMap
from paper_formatter.renderers.omml_renderer import (
    LatexToOmmlConverter,
    OmmlConversionError,
)


class DocxRenderer:
    """Базовый независимый ArticleIR → DOCX-рендерер."""

    def __init__(self) -> None:
        self._omml = LatexToOmmlConverter()
        self.warnings: list[str] = []
        self._styles = DocxTemplateStyleMap()
        self._template_document_used = False

    def render(
        self,
        article: ArticleIR,
        output_path: Path,
        *,
        profile: TemplateProfile | None = None,
        asset_root: Path | None = None,
    ) -> Path:
        self.warnings = []
        profile = profile or TemplateProfile()
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        primary_language = (
            article.metadata.titles[0].language
            if article.metadata.titles
            else (
                article.metadata.abstracts[0].language
                if article.metadata.abstracts
                else None
            )
        )
        front_abstracts, secondary_abstracts = self._split_localized_front_matter(
            article.metadata.abstracts,
            primary_language,
        )
        front_keywords, secondary_keywords = self._split_keywords(
            article.metadata.keywords,
            primary_language,
        )
        has_front_matter = any(
            (
                article.metadata.titles,
                article.metadata.subtitles,
                article.metadata.authors,
                article.metadata.affiliations,
                article.metadata.udc,
                article.metadata.doi,
                front_abstracts,
                front_keywords,
            )
        )
        document = self._create_document(profile)
        self._configure_document(
            document,
            profile,
            front_matter=has_front_matter,
        )
        # Clearing the template body can remove the section that owns the
        # effective first-page stories. Reattach them explicitly in both the
        # template-backed and generic paths.
        self._copy_template_header_footer(document, profile)
        usable_width_mm = (
            (profile.page.width_mm or 210.0)
            - profile.page.margin_left_mm
            - profile.page.margin_right_mm
        )
        column_width_mm = self._column_width_mm(profile, usable_width_mm)
        full_figure_width_mm = max(
            20.0,
            min(170.0, usable_width_mm * profile.figure_width_fraction),
        )

        if article.metadata.titles:
            paragraph = self._add_paragraph(document, "title", "Title")
            paragraph.alignment = self._alignment(profile.typography.title_alignment)
            self._clear_first_line_indent(paragraph)
            self._apply_front_matter_geometry(paragraph, profile)
            paragraph.add_run(article.metadata.titles[0].text)
        if article.metadata.subtitles:
            paragraph = self._add_paragraph(document, "subtitle", "Subtitle")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._clear_first_line_indent(paragraph)
            self._apply_front_matter_geometry(paragraph, profile)
            paragraph.add_run(article.metadata.subtitles[0].text)
        if article.metadata.authors:
            paragraph = self._add_paragraph(document, "authors")
            paragraph.alignment = self._alignment(profile.typography.author_alignment)
            self._clear_first_line_indent(paragraph)
            self._apply_front_matter_geometry(paragraph, profile)
            paragraph.add_run(", ".join(author.name for author in article.metadata.authors))
        for affiliation in article.metadata.affiliations:
            paragraph = self._add_paragraph(document, "affiliation")
            paragraph.alignment = self._alignment(
                profile.typography.affiliation_alignment
            )
            self._clear_first_line_indent(paragraph)
            self._apply_front_matter_geometry(paragraph, profile)
            run = paragraph.add_run(affiliation.name)
            if not self._styles.is_template_role("affiliation"):
                run.italic = True
        if article.metadata.udc:
            paragraph = self._add_paragraph(document, "body_no_indent")
            self._clear_first_line_indent(paragraph)
            paragraph.add_run("УДК: ").bold = True
            paragraph.add_run(article.metadata.udc)
        if article.metadata.doi:
            paragraph = self._add_paragraph(document, "body_no_indent")
            self._clear_first_line_indent(paragraph)
            paragraph.add_run("DOI: ").bold = True
            paragraph.add_run(article.metadata.doi)
        for abstract in front_abstracts:
            paragraph = self._add_paragraph(document, "abstract")
            self._clear_first_line_indent(paragraph)
            paragraph.add_run("Аннотация. ").bold = True
            paragraph.add_run(abstract.text)
        if front_keywords:
            paragraph = self._add_paragraph(document, "keywords")
            self._clear_first_line_indent(paragraph)
            paragraph.add_run("Ключевые слова: ").bold = True
            paragraph.add_run(", ".join(front_keywords))

        if profile.page.columns > 1 and has_front_matter:
            self._add_layout_section(
                document,
                profile,
                columns=profile.page.columns,
            )

        assets = {asset.id: asset for asset in article.assets}
        figure_index = 0
        current_figure_group: str | None = None
        table_index = 0
        for block in article.body:
            if isinstance(block, SectionBlock):
                level = min(block.level, 6)
                paragraph = self._add_paragraph(
                    document, f"heading{level}", f"Heading {level}"
                )
                heading_text = block.title
                if (
                    block.number
                    and not self._paragraph_has_automatic_numbering(paragraph)
                    and not re.match(
                        rf"^\s*{re.escape(block.number)}(?:\.|\s)",
                        heading_text,
                    )
                ):
                    heading_text = f"{block.number}. {heading_text}"
                paragraph.add_run(heading_text)
            elif isinstance(block, ParagraphBlock):
                self._append_paragraph_block(
                    document,
                    block,
                    assets,
                    asset_root,
                    column_width_mm=column_width_mm,
                )
            elif isinstance(block, ListItemBlock):
                role = "list_number" if block.ordered else "list_bullet"
                fallback = "List Number" if block.ordered else "List Bullet"
                paragraph = self._add_paragraph(document, role, fallback)
                if block.level > 0:
                    base_indent = paragraph.style.paragraph_format.left_indent
                    base_mm = base_indent.mm if base_indent is not None else 0.0
                    paragraph.paragraph_format.left_indent = Mm(base_mm + block.level * 6)
                if not self._styles.is_template_role(role):
                    self._clear_first_line_indent(paragraph)
                self._append_runs(paragraph, block.runs, assets, asset_root)
            elif isinstance(block, EquationBlock):
                full_width = profile.page.columns > 1
                if full_width:
                    self._add_layout_section(document, profile, columns=1)
                paragraph = self._add_paragraph(document, "equation")
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._clear_first_line_indent(paragraph)
                converted = self._append_math(
                    paragraph,
                    block.latex,
                    display=block.display,
                )
                if not converted:
                    run = paragraph.add_run(block.latex)
                    run.font.name = "Cambria Math"
                if block.number:
                    paragraph.add_run(f"    ({block.number})")
                if full_width:
                    self._add_layout_section(
                        document,
                        profile,
                        columns=profile.page.columns,
                    )
            elif isinstance(block, FigureBlock):
                figure_group = block.group_id or block.id
                if figure_group != current_figure_group:
                    figure_index += 1
                    current_figure_group = figure_group
                full_width = profile.page.columns > 1
                if full_width:
                    self._add_layout_section(document, profile, columns=1)
                paragraph = self._add_paragraph(document, "figure")
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._clear_first_line_indent(paragraph)
                asset = assets.get(block.asset_id)
                path = self._asset_path(asset.path, asset_root) if asset else None
                if path and path.exists() and path.suffix.lower() in {".png", ".jpg", ".jpeg"}:
                    paragraph.add_run().add_picture(
                        str(path),
                        width=Mm(
                            full_figure_width_mm
                            if full_width
                            else max(
                                20.0,
                                column_width_mm * profile.figure_width_fraction,
                            )
                        ),
                    )
                else:
                    paragraph.add_run("[рисунок]")
                if block.caption:
                    caption = self._add_paragraph(
                        document, "figure_caption", "Caption"
                    )
                    if not self._styles.is_template_role("figure_caption"):
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        self._clear_first_line_indent(caption)
                    if re.match(
                        r"^(?:Рис(?:унок)?|Fig(?:ure)?)\.?\s*\d+",
                        block.caption,
                        flags=re.IGNORECASE,
                    ):
                        caption.add_run(block.caption)
                    else:
                        label = (
                            "Рис."
                            if re.search(r"[А-Яа-яЁё]", block.caption)
                            else "Fig."
                        )
                        caption.add_run(f"{label} {figure_index}. ")
                        caption.add_run(block.caption)
                if full_width:
                    self._add_layout_section(
                        document,
                        profile,
                        columns=profile.page.columns,
                    )
            elif isinstance(block, TableBlock):
                table_index += 1
                table_columns = max((len(row) for row in block.rows), default=0)
                full_width = profile.page.columns > 1 and table_columns >= 4
                if full_width:
                    self._add_layout_section(document, profile, columns=1)
                self._append_table(
                    document,
                    block,
                    profile=profile,
                    usable_width_mm=(
                        usable_width_mm if full_width else column_width_mm
                    ),
                    number=table_index,
                )
                if full_width:
                    self._add_layout_section(
                        document,
                        profile,
                        columns=profile.page.columns,
                    )

        if article.references:
            heading = self._add_paragraph(document, "heading1", "Heading 1")
            heading.add_run("Литература")
            if not self._styles.is_template_role("heading1"):
                self._apply_reference_layout(heading, profile, heading=True)
            reference_size = max(8.0, profile.typography.main_size_pt - 0.8)
            for index, reference in enumerate(article.references, start=1):
                paragraph = self._add_paragraph(document, "references")
                custom_reference_style = self._styles.is_template_role("references")
                if not custom_reference_style:
                    self._clear_first_line_indent(paragraph)
                    self._apply_reference_layout(paragraph, profile, heading=False)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    paragraph.paragraph_format.line_spacing = 1.0
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(1)
                paragraph_properties = paragraph._p.pPr
                has_automatic_numbering = bool(
                    paragraph_properties is not None
                    and paragraph_properties.find(qn("w:numPr")) is not None
                )
                if not has_automatic_numbering:
                    self._append_reference_sequence(
                        paragraph,
                        index,
                        profile.typography.main_font,
                        reference_size,
                    )
                reference_text = reference.text
                reference_text = re.sub(
                    r"^\s*(?:\d+\s*[.)]?\s*)+",
                    "",
                    reference_text,
                )
                reference_text = re.sub(
                    r"^\s*[.)]+\s*",
                    "",
                    reference_text,
                )
                run = paragraph.add_run(reference_text)
                if not custom_reference_style:
                    run.font.name = profile.typography.main_font
                    run.font.size = Pt(reference_size)
        for note in article.notes:
            paragraph = self._add_paragraph(document, "notes")
            self._clear_first_line_indent(paragraph)
            paragraph.add_run(f"[{note.kind}] ").bold = True
            paragraph.add_run(note.text)

        secondary_titles = article.metadata.titles[1:]
        if secondary_titles or secondary_abstracts or secondary_keywords:
            self._add_layout_section(
                document,
                profile,
                columns=1,
                start_type=WD_SECTION.NEW_PAGE,
            )
            secondary_language = (
                secondary_titles[0].language
                if secondary_titles
                else (
                    secondary_abstracts[0].language
                    if secondary_abstracts
                    else None
                )
            )
            for title in secondary_titles:
                paragraph = self._add_paragraph(document, "title", "Title")
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._clear_first_line_indent(paragraph)
                paragraph.add_run(title.text)
            for variant in article.metadata.author_variants:
                if (
                    secondary_language is None
                    or variant.language is None
                    or variant.language == secondary_language
                ):
                    paragraph = self._add_paragraph(document, "authors")
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    self._clear_first_line_indent(paragraph)
                    paragraph.add_run(variant.text).bold = True
            for abstract in secondary_abstracts:
                paragraph = self._add_paragraph(document, "abstract")
                self._clear_first_line_indent(paragraph)
                paragraph.add_run(
                    "Abstract. " if abstract.language == "en" else "Аннотация. "
                ).bold = True
                paragraph.add_run(abstract.text)
            if secondary_keywords:
                paragraph = self._add_paragraph(document, "keywords")
                self._clear_first_line_indent(paragraph)
                paragraph.add_run(
                    "Keywords: " if secondary_language == "en" else "Ключевые слова: "
                ).bold = True
                paragraph.add_run(", ".join(secondary_keywords))

        renamed_parts = self._ensure_unique_package_partnames(document)
        if renamed_parts:
            self.warnings.append(
                "DOCX: устранены коллизии внутренних имён ресурсов: "
                f"{renamed_parts}."
            )
        document.save(output_path)
        return output_path

    @staticmethod
    def _ensure_unique_package_partnames(document: Document) -> int:
        """Give every OPC part a unique name before ZIP serialization.

        A template can already contain generic ``imageNN`` parts.  python-docx
        may later allocate the same names for newly inserted formula previews,
        producing duplicate ZIP members that LibreOffice refuses to open.
        Relationships point to part objects, so renaming the colliding object
        here keeps all references intact.
        """

        used: set[str] = set()
        renamed = 0
        for part in document.part.package.iter_parts():
            current = str(part.partname)
            folded = current.casefold()
            if folded not in used:
                used.add(folded)
                continue

            stem, suffix = posixpath.splitext(current)
            index = 2
            candidate = f"{stem}-pf-{index}{suffix}"
            while candidate.casefold() in used:
                index += 1
                candidate = f"{stem}-pf-{index}{suffix}"
            part.partname = PackURI(candidate)
            used.add(candidate.casefold())
            renamed += 1
        return renamed

    def _create_document(self, profile: TemplateProfile) -> Document:
        source_path = Path(profile.source_path) if profile.source_path else None
        if (
            profile.source_type == "docx"
            and source_path is not None
            and source_path.exists()
        ):
            try:
                document = Document(source_path)
                self._styles = DocxTemplateStyleMap.from_document(
                    document, source_path=source_path
                )
                self._clear_document_body(document)
                self._template_document_used = True
                return document
            except Exception as exc:
                self.warnings.append(
                    f"DOCX: не удалось использовать DOCX-образец как основу ({exc}); "
                    "использован новый документ."
                )
        document = Document()
        self._styles = DocxTemplateStyleMap.from_document(document)
        self._template_document_used = False
        return document

    @staticmethod
    def _clear_document_body(document: Document) -> None:
        body = document._element.body
        for child in list(body):
            if child.tag != qn("w:sectPr"):
                body.remove(child)

    def _add_paragraph(
        self,
        document: Document,
        role: str,
        fallback: str | None = None,
    ):
        style_name = self._styles.paragraph(role, fallback)
        if style_name and style_name in document.styles:
            paragraph = document.add_paragraph(style=style_name)
        else:
            paragraph = document.add_paragraph()
        self._styles.apply_paragraph_properties(paragraph, role)
        return paragraph

    def _configure_document(
        self,
        document: Document,
        profile: TemplateProfile,
        *,
        front_matter: bool = False,
    ) -> None:
        section = document.sections[0]
        self._set_section_layout(
            section,
            profile,
            columns=(
                1
                if front_matter and profile.page.columns > 1
                else profile.page.columns
            ),
            title_zone=front_matter and profile.page.columns > 1,
        )

        normal = document.styles["Normal"]
        self._set_style_font(normal, profile.typography.main_font)
        self._set_style_color(normal, RGBColor(0, 0, 0))
        normal.font.size = Pt(profile.typography.main_size_pt)
        normal.paragraph_format.line_spacing = (
            Pt(profile.typography.line_spacing_pt)
            if profile.typography.line_spacing_pt is not None
            else profile.typography.line_spacing
        )
        if profile.typography.line_spacing_rule in {"atLeast", "exact"}:
            spacing = normal.element.get_or_add_pPr().get_or_add_spacing()
            spacing.set(qn("w:lineRule"), profile.typography.line_spacing_rule)
        normal.paragraph_format.first_line_indent = Mm(
            profile.typography.first_line_indent_mm
        )
        normal.paragraph_format.space_before = Pt(
            profile.typography.paragraph_space_before_pt
        )
        normal.paragraph_format.space_after = Pt(
            profile.typography.paragraph_space_after_pt
        )
        normal.paragraph_format.alignment = self._alignment(
            profile.typography.paragraph_alignment
        )
        self._set_style_language(normal, "ru-RU")
        self._enable_hyphenation(document)

        if "Title" in document.styles:
            title = document.styles["Title"]
            self._set_style_font(
                title,
                profile.typography.title_font or profile.typography.main_font,
            )
            self._set_style_color(title, RGBColor(0, 0, 0))
            self._remove_style_borders(title)
            if profile.typography.title_size_pt:
                title.font.size = Pt(profile.typography.title_size_pt)
            title.font.bold = profile.typography.title_bold
            title.paragraph_format.alignment = self._alignment(
                profile.typography.title_alignment
            )
            title.paragraph_format.first_line_indent = Mm(0)

        if "Subtitle" in document.styles:
            subtitle = document.styles["Subtitle"]
            self._set_style_font(subtitle, profile.typography.main_font)
            self._set_style_color(subtitle, RGBColor(0, 0, 0))
            self._remove_style_borders(subtitle)
            subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle.paragraph_format.first_line_indent = Mm(0)

        if "Caption" in document.styles:
            caption = document.styles["Caption"]
            self._set_style_font(caption, profile.typography.main_font)
            self._set_style_color(caption, RGBColor(0, 0, 0))
            self._remove_style_borders(caption)
            if profile.typography.caption_size_pt:
                caption.font.size = Pt(profile.typography.caption_size_pt)
            caption.paragraph_format.first_line_indent = Mm(0)

        for heading in profile.headings:
            style_name = f"Heading {heading.level}"
            if style_name not in document.styles:
                continue
            style = document.styles[style_name]
            if heading.font:
                self._set_style_font(style, heading.font)
            self._set_style_color(style, RGBColor(0, 0, 0))
            self._remove_style_borders(style)
            if heading.size_pt:
                style.font.size = Pt(heading.size_pt)
            style.font.bold = heading.bold
            style.font.italic = heading.italic
            style.paragraph_format.alignment = self._alignment(heading.alignment)
            style.paragraph_format.first_line_indent = Mm(0)
            if heading.space_before_pt is not None:
                style.paragraph_format.space_before = Pt(heading.space_before_pt)
            if heading.space_after_pt is not None:
                style.paragraph_format.space_after = Pt(heading.space_after_pt)

    def _set_section_layout(
        self,
        section,
        profile: TemplateProfile,
        *,
        columns: int,
        title_zone: bool = False,
    ) -> None:
        section.top_margin = Mm(profile.page.margin_top_mm)
        section.right_margin = Mm(
            profile.page.title_margin_right_mm
            if title_zone and profile.page.title_margin_right_mm is not None
            else profile.page.margin_right_mm
        )
        section.bottom_margin = Mm(profile.page.margin_bottom_mm)
        section.left_margin = Mm(
            profile.page.title_margin_left_mm
            if title_zone and profile.page.title_margin_left_mm is not None
            else profile.page.margin_left_mm
        )
        if profile.page.width_mm and profile.page.height_mm:
            section.page_width = Mm(profile.page.width_mm)
            section.page_height = Mm(profile.page.height_mm)
        if profile.page.header_distance_mm is not None:
            section.header_distance = Mm(profile.page.header_distance_mm)
        if profile.page.footer_distance_mm is not None:
            section.footer_distance = Mm(profile.page.footer_distance_mm)

        page_numbering = section._sectPr.find(qn("w:pgNumType"))
        if page_numbering is not None:
            page_numbering.attrib.pop(qn("w:start"), None)

        existing = section._sectPr.xpath("./w:cols")
        for item in existing:
            section._sectPr.remove(item)
        cols = OxmlElement("w:cols")
        cols.set(qn("w:num"), str(max(1, columns)))
        if profile.page.column_gap_mm is not None:
            cols.set(
                qn("w:space"),
                str(round(profile.page.column_gap_mm * 1440 / 25.4)),
            )
        section._sectPr.append(cols)

    def _add_layout_section(
        self,
        document: Document,
        profile: TemplateProfile,
        *,
        columns: int,
        start_type=WD_SECTION.CONTINUOUS,
    ):
        section = document.add_section(start_type)
        self._set_section_layout(
            section,
            profile,
            columns=columns,
        )
        return section

    @staticmethod
    def _paragraph_has_automatic_numbering(paragraph) -> bool:
        paragraph_properties = paragraph._p.pPr
        if (
            paragraph_properties is not None
            and paragraph_properties.find(qn("w:numPr")) is not None
        ):
            return True
        style = paragraph.style
        style_properties = style.element.pPr if style is not None else None
        return bool(
            style_properties is not None
            and style_properties.find(qn("w:numPr")) is not None
        )

    def _copy_template_header_footer(
        self,
        document: Document,
        profile: TemplateProfile,
    ) -> None:
        source_path = Path(profile.source_path) if profile.source_path else None
        if (
            profile.source_type != "docx"
            or source_path is None
            or not source_path.exists()
        ):
            return
        try:
            template = Document(source_path)
        except Exception as exc:
            self.warnings.append(
                f"DOCX: не удалось прочитать колонтитулы шаблона ({exc})."
            )
            return

        target_section = document.sections[0]
        source_section = template.sections[0]
        self._copy_simple_story(
            source_section.header,
            target_section.header,
            "верхний колонтитул",
        )
        self._copy_simple_story(
            source_section.footer,
            target_section.footer,
            "нижний колонтитул",
        )
        target_section.different_first_page_header_footer = (
            source_section.different_first_page_header_footer
        )
        if source_section.different_first_page_header_footer:
            self._copy_simple_story(
                source_section.first_page_header,
                target_section.first_page_header,
                "верхний колонтитул первой страницы",
            )
            self._copy_simple_story(
                source_section.first_page_footer,
                target_section.first_page_footer,
                "нижний колонтитул первой страницы",
            )

    def _copy_simple_story(self, source, target, label: str) -> None:
        if not source._element.xpath(".//w:t | .//w:instrText | .//w:drawing"):
            return
        target.is_linked_to_previous = False
        rel_map: dict[str, str] = {}
        for rel_id, relationship in source.part.rels.items():
            if relationship.is_external:
                new_rel_id = target.part.relate_to(
                    relationship.target_ref,
                    relationship.reltype,
                    is_external=True,
                )
            else:
                new_rel_id = target.part.relate_to(
                    relationship.target_part,
                    relationship.reltype,
                )
            rel_map[rel_id] = new_rel_id
        for child in list(target._element):
            target._element.remove(child)
        for child in source._element:
            copied = deepcopy(child)
            for element in copied.iter():
                for attr_name in (qn("r:id"), qn("r:embed"), qn("r:link")):
                    rel_id = element.get(attr_name)
                    if rel_id in rel_map:
                        element.set(attr_name, rel_map[rel_id])
            target._element.append(copied)

    @staticmethod
    def _set_style_language(style, language: str) -> None:
        properties = style.element.get_or_add_rPr()
        lang = properties.find(qn("w:lang"))
        if lang is None:
            lang = OxmlElement("w:lang")
            properties.append(lang)
        lang.set(qn("w:val"), language)
        lang.set(qn("w:eastAsia"), language)

    @staticmethod
    def _enable_hyphenation(document: Document) -> None:
        settings = document.settings.element
        auto = settings.find(qn("w:autoHyphenation"))
        if auto is None:
            auto = OxmlElement("w:autoHyphenation")
            settings.append(auto)
        auto.set(qn("w:val"), "true")
        limit = settings.find(qn("w:consecutiveHyphenLimit"))
        if limit is None:
            limit = OxmlElement("w:consecutiveHyphenLimit")
            settings.append(limit)
        limit.set(qn("w:val"), "2")

    def _append_runs(
        self,
        paragraph,
        runs: list[TextRun],
        assets: dict,
        asset_root: Path | None,
        *,
        max_formula_width_pt: float | None = None,
    ) -> None:
        for item in runs:
            if item.asset_id:
                asset = assets.get(item.asset_id)
                path = self._asset_path(asset.path, asset_root) if asset else None
                if path and path.exists():
                    if self._append_preserved_math_ole(
                        paragraph, path, item, assets, asset_root
                    ):
                        continue
                    if path.suffix.lower() in {".png", ".jpg", ".jpeg"}:
                        self._append_formula_picture(
                            paragraph,
                            path,
                            item,
                            max_width_pt=max_formula_width_pt,
                        )
                        continue
                paragraph.add_run("[формула]")
                continue
            value = item.math_latex if item.math_latex is not None else item.text
            if item.math_latex is not None and self._append_math(
                paragraph,
                item.math_latex,
                display=False,
            ):
                continue
            run = paragraph.add_run(value)
            run.bold = item.bold
            run.italic = item.italic
            run.underline = item.underline
            run.font.superscript = item.superscript
            run.font.subscript = item.subscript
            if item.math_latex is not None:
                run.font.name = "Cambria Math"

    def _append_paragraph_block(
        self,
        document: Document,
        block: ParagraphBlock,
        assets: dict,
        asset_root: Path | None,
        *,
        column_width_mm: float,
    ) -> None:
        """Keep oversized legacy MathType previews out of surrounding prose."""

        column_width_pt = column_width_mm * 72 / 25.4
        inline_limit_pt = max(150.0, column_width_pt * 0.55)
        display_limit_pt = max(150.0, column_width_pt * 0.95)
        pending: list[TextRun] = []

        def flush_pending() -> None:
            if not pending:
                return
            text_value = "".join(
                (item.math_latex if item.math_latex is not None else item.text)
                for item in pending
            ).lstrip()
            role = "body_no_indent" if text_value.startswith(("{", "[", "<")) else "body"
            paragraph = self._add_paragraph(document, role)
            formula_height = max(
                (item.height_pt or 14.0)
                for item in pending
                if item.asset_id and item.formula_image
            ) if any(item.asset_id and item.formula_image for item in pending) else None
            if formula_height is not None:
                self._set_formula_line_spacing(paragraph, formula_height)
            self._append_runs(
                paragraph,
                pending,
                assets,
                asset_root,
                max_formula_width_pt=inline_limit_pt,
            )
            pending.clear()

        for item in block.runs:
            is_formula = bool(item.asset_id and item.formula_image)
            needs_display = is_formula and (
                item.display or (item.width_pt or 0) > inline_limit_pt
            )
            if not needs_display:
                pending.append(item)
                continue
            flush_pending()
            paragraph = self._add_paragraph(document, "equation")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._clear_first_line_indent(paragraph)
            self._set_formula_line_spacing(paragraph, item.height_pt or 14.0)
            self._append_runs(
                paragraph,
                [item],
                assets,
                asset_root,
                max_formula_width_pt=display_limit_pt,
            )
        flush_pending()

    @staticmethod
    def _set_formula_line_spacing(paragraph, formula_height_pt: float) -> None:
        """Reserve vertical room for legacy OLE objects so they cannot overlap prose."""
        paragraph.paragraph_format.line_spacing = Pt(max(14.0, formula_height_pt + 4.0))
        spacing = paragraph._p.get_or_add_pPr().get_or_add_spacing()
        spacing.set(qn("w:lineRule"), "atLeast")

    @staticmethod
    def _append_formula_picture(
        paragraph,
        path: Path,
        item: TextRun,
        *,
        max_width_pt: float | None,
    ) -> None:
        width_pt = item.width_pt
        if width_pt is not None and max_width_pt is not None:
            width_pt = min(width_pt, max_width_pt)
        if width_pt is not None:
            paragraph.add_run().add_picture(str(path), width=Pt(width_pt))
        else:
            paragraph.add_run().add_picture(
                str(path),
                height=Pt(item.height_pt or 13.0),
            )

    @staticmethod
    def _append_preserved_math_ole(
        paragraph,
        preview_path: Path,
        item: TextRun,
        assets: dict,
        asset_root: Path | None,
    ) -> bool:
        """Use the original MathType payload; PNG remains a compatibility fallback."""
        if not item.ole_object_xml or not item.ole_object_asset_id:
            return False
        ole_asset = assets.get(item.ole_object_asset_id)
        preview_asset = assets.get(item.ole_preview_asset_id)
        ole_path = (
            DocxRenderer._asset_path(ole_asset.path, asset_root)
            if ole_asset is not None
            else None
        )
        preview_raw_path = (
            DocxRenderer._asset_path(preview_asset.path, asset_root)
            if preview_asset is not None
            else None
        )
        if (
            ole_path is None
            or not ole_path.exists()
            or preview_raw_path is None
            or not preview_raw_path.exists()
        ):
            return False
        try:
            document_part = paragraph.part
            preview_part = Part(
                PackURI(f"/word/media/{preview_asset.id}-{preview_raw_path.name}"),
                preview_asset.media_type or "image/x-wmf",
                preview_raw_path.read_bytes(),
                document_part.package,
            )
            preview_rel_id = document_part.relate_to(preview_part, RT.IMAGE)
            ole_part = Part(
                PackURI(f"/word/embeddings/{ole_asset.id}-{ole_path.name}"),
                ole_asset.media_type
                or "application/vnd.openxmlformats-officedocument.oleObject",
                ole_path.read_bytes(),
                document_part.package,
            )
            ole_rel_id = document_part.relate_to(ole_part, RT.OLE_OBJECT)
            object_element = parse_xml(item.ole_object_xml)
            for element in object_element.iter():
                local_name = element.tag.rsplit("}", 1)[-1]
                if local_name == "imagedata":
                    element.set(qn("r:id"), preview_rel_id)
                elif local_name == "OLEObject":
                    element.set(qn("r:id"), ole_rel_id)
            paragraph.add_run()._r.append(object_element)
            return True
        except Exception:
            return False

    def _append_math(self, paragraph, latex: str, *, display: bool) -> bool:
        try:
            paragraph._p.append(self._omml.convert(latex, display=display))
        except OmmlConversionError as exc:
            warning = (
                "DOCX: часть формул оставлена как LaTeX-текст, потому что "
                f"преобразование в OMML недоступно ({exc})."
            )
            if warning not in self.warnings:
                self.warnings.append(warning)
            return False
        return True

    def _append_table(
        self,
        document: Document,
        block: TableBlock,
        *,
        profile: TemplateProfile,
        usable_width_mm: float,
        number: int,
    ) -> None:
        if not block.rows:
            return
        columns = max(len(row) for row in block.rows)
        if block.caption:
            caption = self._add_paragraph(
                document, "table_caption", "Caption"
            )
            if not self._styles.is_template_role("table_caption"):
                self._clear_first_line_indent(caption)
                caption.paragraph_format.space_before = Pt(6)
                caption.paragraph_format.space_after = Pt(3)
            match = re.match(
                r"^((?:Table|Таблица)\s+\d+[.:]?)\s*(.*)$",
                block.caption,
                flags=re.IGNORECASE,
            )
            if match:
                caption.add_run(match.group(1) + " ").bold = True
                caption.add_run(match.group(2))
            else:
                label = (
                    "Таблица"
                    if re.search(r"[А-Яа-яЁё]", block.caption)
                    else "Table"
                )
                caption.add_run(f"{label} {number}. ").bold = True
                caption.add_run(block.caption)
        table = document.add_table(rows=len(block.rows), cols=columns)
        if self._styles.table_style and self._styles.table_style in document.styles:
            table.style = self._styles.table_style
        else:
            table.style = None
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        widths_mm = self._column_widths_mm(block, columns, usable_width_mm)
        if self._styles.table_style:
            self._configure_template_table(table, widths_mm)
        else:
            self._configure_academic_table(table, widths_mm)
        table_font_size = profile.typography.caption_size_pt or max(
            8.0,
            profile.typography.main_size_pt - 1.0,
        )
        for row_index, row in enumerate(block.rows):
            table.rows[row_index].height_rule = None
            self._prevent_row_split(table.rows[row_index])
            if row_index < block.header_rows:
                self._repeat_table_header(table.rows[row_index])
            for column_index in range(columns):
                cell = table.cell(row_index, column_index)
                cell.width = Mm(widths_mm[column_index])
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                cell.text = (
                    row[column_index] if column_index < len(row) else ""
                )
                paragraph = cell.paragraphs[0]
                table_body_style = self._styles.paragraph("table_body")
                if table_body_style and table_body_style in document.styles:
                    paragraph.style = table_body_style
                paragraph.paragraph_format.first_line_indent = Mm(0)
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                value = paragraph.text.strip()
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER
                    if row_index < block.header_rows
                    or self._looks_numeric(value)
                    else WD_ALIGN_PARAGRAPH.LEFT
                )
                if row_index < block.header_rows:
                    self._set_cell_bottom_border(cell)
                    for run in paragraph.runs:
                        run.bold = True
                for run in paragraph.runs:
                    run.font.name = profile.typography.main_font
                    run.font.size = Pt(table_font_size)

        for merged in block.merged_cells:
            last_row = min(len(block.rows) - 1, merged.row + merged.row_span - 1)
            last_col = min(columns - 1, merged.column + merged.column_span - 1)
            if (
                merged.row < len(block.rows)
                and merged.column < columns
                and (last_row > merged.row or last_col > merged.column)
            ):
                table.cell(merged.row, merged.column).merge(
                    table.cell(last_row, last_col)
                )

    @staticmethod
    def _column_widths_mm(
        block: TableBlock,
        columns: int,
        usable_width_mm: float,
    ) -> list[float]:
        available = max(20.0, usable_width_mm)
        if block.column_widths_pt and any(
            width is not None and width > 0 for width in block.column_widths_pt
        ):
            raw = [
                max(1.0, float(width or 24.0) * 25.4 / 72)
                for width in block.column_widths_pt[:columns]
            ]
            raw.extend([24.0] * (columns - len(raw)))
        else:
            raw = []
            for column in range(columns):
                values = [
                    row[column].strip()
                    for row in block.rows
                    if column < len(row) and row[column].strip()
                ]
                body_values = values[block.header_rows :] or values
                numeric_share = (
                    sum(DocxRenderer._looks_numeric(value) for value in body_values)
                    / len(body_values)
                    if body_values
                    else 0.0
                )
                longest = max((len(value) for value in values), default=6)
                if numeric_share >= 0.6:
                    raw.append(max(7.0, min(longest, 14) * 0.65))
                else:
                    raw.append(max(8.0, min(longest, 40) * 0.75))

        minimum = min(12.0, available / max(columns, 1))
        widths = [available * value / sum(raw) for value in raw]
        fixed: set[int] = set()
        while True:
            new_fixed = {
                index
                for index, width in enumerate(widths)
                if width < minimum and index not in fixed
            }
            if not new_fixed:
                break
            fixed.update(new_fixed)
            remaining = available - minimum * len(fixed)
            flexible = [index for index in range(columns) if index not in fixed]
            flexible_weight = sum(raw[index] for index in flexible)
            for index in fixed:
                widths[index] = minimum
            for index in flexible:
                widths[index] = (
                    remaining * raw[index] / flexible_weight
                    if flexible_weight
                    else remaining / max(1, len(flexible))
                )
        return [round(value, 2) for value in widths]

    @staticmethod
    def _configure_template_table(table, widths_mm: list[float]) -> None:
        """Keep the template table style and only set deterministic geometry."""
        for index, width_mm in enumerate(widths_mm):
            table.columns[index].width = Mm(width_mm)
        tbl_pr = table._tbl.tblPr
        layout = tbl_pr.find(qn("w:tblLayout"))
        if layout is None:
            layout = OxmlElement("w:tblLayout")
            tbl_pr.append(layout)
        layout.set(qn("w:type"), "fixed")
        width = tbl_pr.find(qn("w:tblW"))
        if width is None:
            width = OxmlElement("w:tblW")
            tbl_pr.append(width)
        width.set(qn("w:type"), "dxa")
        width.set(qn("w:w"), str(round(sum(widths_mm) * 1440 / 25.4)))

    @staticmethod
    def _configure_academic_table(table, widths_mm: list[float]) -> None:
        for index, width_mm in enumerate(widths_mm):
            table.columns[index].width = Mm(width_mm)
        tbl_pr = table._tbl.tblPr
        layout = tbl_pr.find(qn("w:tblLayout"))
        if layout is None:
            layout = OxmlElement("w:tblLayout")
            tbl_pr.append(layout)
        layout.set(qn("w:type"), "fixed")

        width = tbl_pr.find(qn("w:tblW"))
        if width is None:
            width = OxmlElement("w:tblW")
            tbl_pr.append(width)
        width.set(qn("w:type"), "dxa")
        width.set(
            qn("w:w"),
            str(round(sum(widths_mm) * 1440 / 25.4)),
        )

        borders = tbl_pr.find(qn("w:tblBorders"))
        if borders is None:
            borders = OxmlElement("w:tblBorders")
            tbl_pr.append(borders)
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            element = borders.find(qn(f"w:{edge}"))
            if element is None:
                element = OxmlElement(f"w:{edge}")
                borders.append(element)
            element.set(
                qn("w:val"),
                "single" if edge in {"top", "bottom"} else "nil",
            )
            if edge in {"top", "bottom"}:
                element.set(qn("w:sz"), "4")
                element.set(qn("w:color"), "000000")

        cell_margin = tbl_pr.find(qn("w:tblCellMar"))
        if cell_margin is None:
            cell_margin = OxmlElement("w:tblCellMar")
            tbl_pr.append(cell_margin)
        for side, twips in (("top", 45), ("left", 70), ("bottom", 45), ("right", 70)):
            margin = cell_margin.find(qn(f"w:{side}"))
            if margin is None:
                margin = OxmlElement(f"w:{side}")
                cell_margin.append(margin)
            margin.set(qn("w:w"), str(twips))
            margin.set(qn("w:type"), "dxa")

    @staticmethod
    def _set_cell_bottom_border(cell) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = tc_pr.find(qn("w:tcBorders"))
        if borders is None:
            borders = OxmlElement("w:tcBorders")
            tc_pr.append(borders)
        bottom = borders.find(qn("w:bottom"))
        if bottom is None:
            bottom = OxmlElement("w:bottom")
            borders.append(bottom)
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:color"), "000000")

    @staticmethod
    def _repeat_table_header(row) -> None:
        tr_pr = row._tr.get_or_add_trPr()
        header = tr_pr.find(qn("w:tblHeader"))
        if header is None:
            header = OxmlElement("w:tblHeader")
            tr_pr.append(header)
        header.set(qn("w:val"), "true")

    @staticmethod
    def _prevent_row_split(row) -> None:
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))

    @staticmethod
    def _looks_numeric(value: str) -> bool:
        return bool(
            value
            and re.fullmatch(
                r"[\s+\-−±]?(?:\d+(?:[.,]\d+)?|[–—-])(?:\s*[%°])?",
                value,
            )
        )

    @staticmethod
    def _set_style_font(style, font_name: str) -> None:
        style.font.name = font_name
        r_pr = style.element.get_or_add_rPr()
        r_fonts = r_pr.rFonts
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.insert(0, r_fonts)
        for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
            r_fonts.set(qn(f"w:{attribute}"), font_name)

    @staticmethod
    def _set_style_color(style, color: RGBColor) -> None:
        style.font.color.rgb = color

    @staticmethod
    def _remove_style_borders(style) -> None:
        p_pr = style.element.pPr
        if p_pr is None:
            return
        borders = p_pr.find(qn("w:pBdr"))
        if borders is not None:
            p_pr.remove(borders)

    @staticmethod
    def _alignment(value: str):
        return {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }[value]

    @staticmethod
    def _clear_first_line_indent(paragraph) -> None:
        paragraph.paragraph_format.first_line_indent = Mm(0)

    @staticmethod
    def _split_localized_front_matter(items, primary_language):
        if not items:
            return [], []
        if primary_language is None:
            return [items[0]], list(items[1:])
        front = [
            item
            for item in items
            if item.language is None or item.language == primary_language
        ]
        secondary = [item for item in items if item not in front]
        if not front:
            return [items[0]], list(items[1:])
        return front, secondary

    @staticmethod
    def _split_keywords(
        keywords: list[str],
        primary_language: str | None,
    ) -> tuple[list[str], list[str]]:
        if not keywords or primary_language is None:
            return list(keywords), []

        def is_cyrillic(value: str) -> bool:
            letters = [character for character in value if character.isalpha()]
            return bool(letters) and (
                sum("\u0400" <= character <= "\u04ff" for character in letters)
                / len(letters)
                >= 0.35
            )

        if primary_language == "ru":
            front = [value for value in keywords if is_cyrillic(value)]
            secondary = [value for value in keywords if value not in front]
        else:
            front = [value for value in keywords if not is_cyrillic(value)]
            secondary = [value for value in keywords if value not in front]
        if not front:
            return list(keywords), []
        return front, secondary

    @staticmethod
    def _column_width_mm(
        profile: TemplateProfile,
        usable_width_mm: float,
    ) -> float:
        columns = max(1, profile.page.columns)
        gap = profile.page.column_gap_mm or 0.0
        return max(20.0, (usable_width_mm - gap * (columns - 1)) / columns)

    @staticmethod
    def _apply_front_matter_geometry(paragraph, profile: TemplateProfile) -> None:
        if profile.page.columns > 1:
            # Multi-column templates get a real one-column title section.
            return
        left = profile.page.title_margin_left_mm
        right = profile.page.title_margin_right_mm
        if left is not None:
            paragraph.paragraph_format.left_indent = Mm(
                left - profile.page.margin_left_mm
            )
        if right is not None:
            paragraph.paragraph_format.right_indent = Mm(
                right - profile.page.margin_right_mm
            )

    @staticmethod
    def _apply_reference_layout(
        paragraph,
        profile: TemplateProfile,
        *,
        heading: bool,
    ) -> None:
        if profile.page.columns > 1:
            paragraph.paragraph_format.left_indent = Mm(0 if heading else 5.0)
            paragraph.paragraph_format.right_indent = Mm(0)
            paragraph.paragraph_format.first_line_indent = Mm(0 if heading else -5.0)
            return
        DocxRenderer._apply_reference_zone(
            paragraph,
            profile,
            text_offset_mm=0.0 if heading else 8.0,
            hanging_mm=0.0 if heading else 8.0,
        )

    @staticmethod
    def _apply_reference_zone(
        paragraph,
        profile: TemplateProfile,
        *,
        text_offset_mm: float,
        hanging_mm: float,
    ) -> None:
        zone_left = (
            profile.page.title_margin_left_mm
            if profile.page.title_margin_left_mm is not None
            else profile.page.margin_left_mm
        )
        zone_right = (
            profile.page.title_margin_right_mm
            if profile.page.title_margin_right_mm is not None
            else profile.page.margin_right_mm
        )
        text_left = zone_left + text_offset_mm
        paragraph.paragraph_format.left_indent = Mm(
            text_left - profile.page.margin_left_mm
        )
        paragraph.paragraph_format.right_indent = Mm(
            zone_right - profile.page.margin_right_mm
        )
        paragraph.paragraph_format.first_line_indent = Mm(-hanging_mm)

    @staticmethod
    def _append_reference_sequence(
        paragraph,
        index: int,
        font_name: str,
        font_size_pt: float,
    ) -> None:
        field = OxmlElement("w:fldSimple")
        instruction = r" SEQ PFBibliography \* ARABIC "
        if index == 1:
            instruction += r"\r 1 "
        field.set(qn("w:instr"), instruction)
        run = OxmlElement("w:r")
        properties = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
            fonts.set(qn(f"w:{attribute}"), font_name)
        properties.append(fonts)
        size = OxmlElement("w:sz")
        size.set(qn("w:val"), str(round(font_size_pt * 2)))
        properties.append(size)
        run.append(properties)
        text = OxmlElement("w:t")
        text.text = str(index)
        run.append(text)
        field.append(run)
        paragraph._p.append(field)
        separator = paragraph.add_run(".      ")
        separator.font.name = font_name
        separator.font.size = Pt(font_size_pt)

    @staticmethod
    def _asset_path(path: str, asset_root: Path | None) -> Path:
        value = Path(path)
        if value.is_absolute():
            return value
        return (asset_root / value).resolve() if asset_root else value.resolve()
