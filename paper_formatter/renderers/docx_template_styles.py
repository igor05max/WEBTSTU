from __future__ import annotations

from collections import Counter
from copy import deepcopy
from dataclasses import dataclass, field
from pathlib import Path
import re

from docx import Document
from docx.enum.style import WD_STYLE_TYPE


def _normalized(value: str | None) -> str:
    """Return a comparable style token for valid and malformed DOCX styles.

    Real publisher templates occasionally contain ``w:style`` records without
    a ``w:name`` child.  python-docx exposes their name as ``None``; treating
    that value as text used to abort the complete formatting pipeline before
    the template profile could be analyzed.
    """

    if not value:
        return ""
    return re.sub(r"[^a-zа-я0-9]+", "_", value.lower()).strip("_")


@dataclass
class DocxTemplateStyleMap:
    """Maps semantic article roles to actual styles present in a DOCX template.

    The mapper intentionally does not invent formatting.  It first looks for
    journal-specific style names, then uses paragraphs observed in the template,
    and finally falls back to the built-in Word styles.
    """

    paragraph_styles: dict[str, str] = field(default_factory=dict)
    table_style: str | None = None
    source_path: str | None = None
    paragraph_properties: dict[str, object] = field(default_factory=dict, repr=False)

    @classmethod
    def from_document(
        cls,
        document: Document,
        *,
        source_path: Path | None = None,
    ) -> "DocxTemplateStyleMap":
        paragraph_names = [
            style.name
            for style in document.styles
            if style.type == WD_STYLE_TYPE.PARAGRAPH and style.name
        ]
        table_names = [
            style.name
            for style in document.styles
            if style.type == WD_STYLE_TYPE.TABLE and style.name
        ]
        normalized = {name: _normalized(name) for name in paragraph_names}
        usage = Counter(
            paragraph.style.name
            for paragraph in document.paragraphs
            if (
                paragraph.text.strip()
                and paragraph.style is not None
                and paragraph.style.name
            )
        )
        table_usage = Counter(
            table.style.name
            for table in document.tables
            if table.style is not None and table.style.name
        )

        aliases: dict[str, list[str]] = {
            "article_type": ["article_type", "articletype", "document_type"],
            "title": ["1_2_title", "article_title", "paper_title", "title"],
            "subtitle": ["subtitle", "sub_title"],
            "authors": ["authornames", "author_names", "authors", "author"],
            "affiliation": ["affiliation", "affiliations", "institution"],
            "abstract": ["abstract", "annotation", "аннотация"],
            "keywords": ["keywords", "key_words", "ключевые_слова"],
            "body": ["3_1_text", "body_text", "main_text", "article_text"],
            "body_no_indent": ["text_no_indent", "body_no_indent", "no_indent"],
            "list_number": ["3_7_itemize", "itemize", "list_number", "numbered_list"],
            "list_bullet": ["3_8_bullet", "bullet", "list_bullet", "bulleted_list"],
            "equation": ["3_9_equation", "equation", "formula"],
            "equation_number": ["equation_number", "formula_number"],
            "table_caption": ["4_1_table_caption", "table_caption", "caption_table"],
            "table_body": ["4_2_table_body", "table_body", "table_text"],
            "table_footer": ["4_3_table_footer", "table_footer", "table_note"],
            "figure_caption": ["5_1_figure_caption", "figure_caption", "caption_figure"],
            "figure": ["5_2_figure", "figure", "image"],
            "references": ["8_1_references", "references", "bibliography"],
            "back_matter": ["6_2_back_matter", "back_matter"],
            "notes": ["6_3_notes", "notes", "note"],
        }
        for level in range(1, 7):
            aliases[f"heading{level}"] = [
                f"2_{level}_heading{level}",
                f"heading_{level}",
                f"heading{level}",
                f"заголовок_{level}",
            ]

        result: dict[str, str] = {}

        def exact_or_contains(role: str) -> str | None:
            candidates = aliases[role]
            # Alias order is meaningful: journal-specific identifiers such as
            # ``3.1_text`` must beat generic built-ins such as ``Body Text``.
            for alias in candidates:
                alias_n = _normalized(alias)
                exact = [
                    name
                    for name, normalized_name in normalized.items()
                    if normalized_name == alias_n
                ]
                if exact:
                    return max(exact, key=lambda name: usage.get(name, 0))
                contained = []
                for name, normalized_name in normalized.items():
                    if not alias_n or alias_n not in normalized_name:
                        continue
                    if role == "title" and "subtitle" in normalized_name:
                        continue
                    if role == "equation" and "number" in normalized_name:
                        continue
                    contained.append(name)
                if contained:
                    return max(
                        contained,
                        key=lambda name: (
                            int(_normalized(name).startswith(("mdpi_", "journal_", "paper_", "article_"))),
                            usage.get(name, 0),
                        ),
                    )
            return None

        for role in aliases:
            chosen = exact_or_contains(role)
            if (
                role == "body"
                and chosen
                and _normalized(chosen) in {"body_text", "body_text_indent"}
                and usage.get(chosen, 0) == 0
            ):
                chosen = None
            if chosen:
                result[role] = chosen

        # Infer roles from actual template paragraphs when style names are opaque.
        paragraphs = [p for p in document.paragraphs if p.text.strip()]
        for paragraph in paragraphs[:40]:
            text = paragraph.text.strip()
            style_name = paragraph.style.name if paragraph.style is not None else None
            if not style_name:
                continue
            lower = text.lower()
            if (
                lower.strip(" .:–—-") in {"title", "заглавие", "название"}
                and (
                    "title" not in result
                    or usage.get(result["title"], 0) == 0
                )
            ):
                result["title"] = style_name
            if "title" not in result and 20 <= len(text) <= 500:
                if paragraph is paragraphs[0] or (
                    paragraph.style.font.size is not None
                    and paragraph.style.font.size.pt >= 15
                ):
                    result["title"] = style_name
            if "abstract" not in result and lower in {"abstract", "аннотация"}:
                result["abstract"] = style_name
            if "keywords" not in result and lower.startswith(("keywords", "ключевые слова")):
                result["keywords"] = style_name
            if "table_caption" not in result and re.match(
                r"^(?:table|таблица)\s*\d+", text, re.IGNORECASE
            ):
                result["table_caption"] = style_name
            if "figure_caption" not in result and re.match(
                r"^(?:figure|fig\.?|рис(?:унок)?\.?)\s*\d+", text, re.IGNORECASE
            ):
                result["figure_caption"] = style_name

        # Body text is normally the most frequently used long-paragraph style.
        if "body" not in result:
            long_usage = Counter()
            for paragraph in paragraphs:
                if len(paragraph.text.strip()) >= 120 and paragraph.style is not None:
                    name = paragraph.style.name
                    if not name:
                        continue
                    norm = _normalized(name)
                    if not re.search(
                        r"heading|title|abstract|keyword|caption|reference|bibli|bullet|itemize|author|affiliation",
                        norm,
                    ):
                        long_usage[name] += 1
            if long_usage:
                result["body"] = long_usage.most_common(1)[0][0]
            elif usage:
                result["body"] = usage.most_common(1)[0][0]

        builtins = {
            "title": "Title",
            "subtitle": "Subtitle",
            "body": "Normal",
            "body_no_indent": "Normal",
            "authors": "Normal",
            "affiliation": "Normal",
            "abstract": "Normal",
            "keywords": "Normal",
            "list_number": "List Number",
            "list_bullet": "List Bullet",
            "equation": "Normal",
            "table_caption": "Caption",
            "table_body": "Normal",
            "figure_caption": "Caption",
            "figure": "Normal",
            "references": "Normal",
            "back_matter": "Normal",
            "notes": "Normal",
        }
        for level in range(1, 7):
            builtins[f"heading{level}"] = f"Heading {level}"
        available = set(paragraph_names)
        for role, fallback in builtins.items():
            if role not in result and fallback in available:
                if (
                    role in {"table_caption", "figure_caption"}
                    and fallback == "Caption"
                    and usage.get(fallback, 0) == 0
                    and "Normal" in available
                ):
                    result[role] = "Normal"
                else:
                    result[role] = fallback

        prototypes: dict[str, object] = {}
        longest_roles = {"abstract", "keywords", "body", "body_no_indent", "references", "back_matter", "notes"}
        for role, style_name in result.items():
            matches = [
                item
                for item in paragraphs
                if item.style is not None and item.style.name == style_name
            ]
            if role == "body" and style_name == "Normal":
                matches = [
                    item
                    for item in matches
                    if len(item.text.strip()) >= 240
                    and (
                        item.alignment is None
                        or str(item.alignment).startswith("JUSTIFY")
                    )
                ]
            if not matches:
                continue
            paragraph = (
                max(matches, key=lambda item: len(item.text.strip()))
                if role in longest_roles
                else matches[0]
            )
            if paragraph._p.pPr is not None:
                prototypes[role] = deepcopy(paragraph._p.pPr)

        table_style = cls._choose_table_style(table_names, table_usage)
        return cls(
            paragraph_styles=result,
            table_style=table_style,
            source_path=str(source_path) if source_path else None,
            paragraph_properties=prototypes,
        )

    @staticmethod
    def _choose_table_style(
        names: list[str],
        usage: Counter[str],
    ) -> str | None:
        priorities = [
            "three_line_table",
            "three_line",
            "academic_table",
            "journal_table",
            "mdpi_table",
            "table_grid",
        ]
        normalized = {name: _normalized(name) for name in names}
        for token in priorities:
            token_n = _normalized(token)
            for name, norm in normalized.items():
                if token_n in norm:
                    if token == "table_grid" and usage.get(name, 0) == 0:
                        continue
                    return name
        return None

    def paragraph(self, role: str, fallback: str | None = None) -> str | None:
        return self.paragraph_styles.get(role, fallback)


    def apply_paragraph_properties(self, paragraph, role: str) -> None:
        properties = self.paragraph_properties.get(role)
        if properties is None:
            return
        existing = paragraph._p.pPr
        if existing is not None:
            paragraph._p.remove(existing)
        paragraph._p.insert(0, deepcopy(properties))

    def is_template_role(self, role: str) -> bool:
        style_name = self.paragraph_styles.get(role)
        return bool(
            style_name
            and style_name not in {
                "Normal",
                "Title",
                "Subtitle",
                "Caption",
                "List Number",
                "List Bullet",
                "Heading 1",
                "Heading 2",
                "Heading 3",
                "Heading 4",
                "Heading 5",
                "Heading 6",
            }
        )

    def as_evidence(self) -> dict[str, str]:
        payload = {
            f"docx_style_{role}": name
            for role, name in sorted(self.paragraph_styles.items())
        }
        if self.table_style:
            payload["docx_style_table"] = self.table_style
        return payload
