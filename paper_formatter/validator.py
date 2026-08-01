from __future__ import annotations

import re
import zipfile
from collections import Counter
from pathlib import Path
from typing import Any

from lxml import etree

from paper_formatter.models import (
    ArticleIR,
    EquationBlock,
    FigureBlock,
    ListItemBlock,
    ParagraphBlock,
    SectionBlock,
    TableBlock,
    TemplateProfile,
)
from paper_formatter.semantic.models import SemanticAnalysis, SemanticBlock


class ConversionValidator:
    """Проверяет структуру, текст, объекты, ссылки, ресурсы и выходные файлы."""

    def validate(
        self,
        *,
        source_path: Path,
        article: ArticleIR,
        main_tex: Path,
        pdf_path: Path | None,
        semantic_blocks: list[SemanticBlock] | None = None,
        semantic_analysis: SemanticAnalysis | None = None,
        docx_path: Path | None = None,
        compile_log: Path | None = None,
        template_profile: TemplateProfile | None = None,
        package_analysis: dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        semantic_blocks = semantic_blocks or []
        source_counts = self._source_counts(source_path)
        article_counts = self._article_counts(article)
        warnings: list[str] = []
        errors: list[str] = []

        if not article.metadata.titles:
            warnings.append("VALIDATION: не найдено название документа.")
        if not article.metadata.authors:
            warnings.append("VALIDATION: не найдены авторы документа.")
        self._compare_count(
            warnings, "таблиц", source_counts.get("tables"), article_counts["tables"]
        )
        self._compare_count(
            warnings,
            "формул",
            source_counts.get("formulas"),
            article_counts["equations"] + article_counts["formula_images"],
        )
        self._compare_count(
            warnings,
            "рисунков",
            source_counts.get("drawings"),
            article_counts["figures"] + article_counts["formula_images"],
            tolerance=1,
        )

        all_ids = [block.id for block in article.body]
        duplicate_ids = [
            value for value, count in Counter(all_ids).items() if count > 1
        ]
        if duplicate_ids:
            errors.append(
                f"VALIDATION: повторяются идентификаторы блоков: {duplicate_ids[:10]}."
            )

        unknown: list[str] = []
        low_confidence: list[str] = []
        if semantic_analysis:
            for decision in semantic_analysis.decisions:
                if decision.role == "unknown":
                    unknown.append(decision.block_id)
                if decision.confidence < 0.60:
                    low_confidence.append(decision.block_id)
        if unknown:
            warnings.append(f"VALIDATION: не определена роль {len(unknown)} блоков.")
        if low_confidence:
            warnings.append(
                f"VALIDATION: низкая уверенность семантики у {len(low_confidence)} блоков."
            )

        tex_exists = main_tex.exists() and main_tex.stat().st_size > 0
        tex_text = (
            main_tex.read_text(encoding="utf-8", errors="replace") if tex_exists else ""
        )
        if not tex_exists:
            errors.append("VALIDATION: main.tex отсутствует или пуст.")
        tex_integrity = self._tex_integrity(main_tex, article)
        warnings.extend(tex_integrity["warnings"])
        errors.extend(tex_integrity["errors"])

        asset_checks = self._asset_checks(main_tex.parent, article)
        warnings.extend(asset_checks["warnings"])
        errors.extend(asset_checks["errors"])

        text_coverage = self._text_coverage(source_path, article)
        if text_coverage is not None and text_coverage < 0.88:
            warnings.append(
                f"VALIDATION: покрытие исходной лексики только {text_coverage:.1%}."
            )
        warnings.extend(
            self._pandoc_disagreements(package_analysis, article_counts)
        )

        compile_warnings = self._compile_log_warnings(compile_log)
        warnings.extend(compile_warnings)
        pdf_info = self._pdf_info(pdf_path)
        if pdf_path is not None and not pdf_info["exists"]:
            errors.append("VALIDATION: заявленный PDF отсутствует.")
        docx_exists = bool(docx_path and docx_path.exists() and docx_path.stat().st_size > 0)
        if docx_path is not None and not docx_exists:
            errors.append("VALIDATION: заявленный DOCX отсутствует.")
        docx_style_audit = self._docx_style_audit(
            docx_path if docx_exists else None,
            template_profile,
            article,
        )
        warnings.extend(docx_style_audit.get("warnings", []))

        structure_score = self._structure_score(source_counts, article_counts)
        asset_score = asset_checks["score"]
        reference_score = tex_integrity["reference_score"]
        warnings = list(dict.fromkeys(warnings))
        errors = list(dict.fromkeys(errors))
        manual_review = bool(
            errors
            or warnings
            or source_path.suffix.lower() == ".pdf"
            or (template_profile and template_profile.confidence < 0.75)
        )
        return {
            "source": str(source_path),
            "semantic_provider": semantic_analysis.provider if semantic_analysis else article.semantic_provider,
            "source_counts": source_counts,
            "article_ir_counts": article_counts,
            "semantic": {
                "candidate_blocks": len(semantic_blocks),
                "unknown_blocks": unknown,
                "low_confidence_blocks": low_confidence,
            },
            "integrity": {
                "text_coverage": text_coverage,
                "duplicate_block_ids": duplicate_ids,
                "assets": asset_checks,
                "latex": tex_integrity,
                "docx_styles": docx_style_audit,
            },
            "outputs": {
                "main_tex_exists": tex_exists,
                "pdf_exists": pdf_info["exists"],
                "pdf_pages": pdf_info.get("pages"),
                "docx_exists": docx_exists,
            },
            "scores": {
                "structure_preservation": structure_score,
                "asset_preservation": asset_score,
                "reference_integrity": reference_score,
                "template_confidence": template_profile.confidence if template_profile else None,
            },
            "package": package_analysis,
            "manual_review_required": manual_review,
            "warnings": warnings,
            "errors": errors,
        }

    @staticmethod
    def _docx_style_audit(
        docx_path: Path | None,
        template_profile: TemplateProfile | None,
        article: ArticleIR,
    ) -> dict[str, Any]:
        result: dict[str, Any] = {
            "enabled": False,
            "expected": {},
            "used_counts": {},
            "matched_roles": [],
            "missing_roles": [],
            "table_style": None,
            "warnings": [],
        }
        if (
            docx_path is None
            or template_profile is None
            or template_profile.source_type != "docx"
        ):
            return result
        try:
            from docx import Document

            document = Document(docx_path)
        except Exception as exc:
            result["warnings"].append(
                f"VALIDATION: не удалось проверить стили итогового DOCX ({exc})."
            )
            return result

        result["enabled"] = True
        style_counts = Counter(
            paragraph.style.name
            for paragraph in document.paragraphs
            if paragraph.text.strip() and paragraph.style is not None
        )
        result["used_counts"] = dict(style_counts)
        evidence = template_profile.evidence or {}
        expected = {
            key.removeprefix("docx_style_"): value
            for key, value in evidence.items()
            if key.startswith("docx_style_") and key != "docx_style_table"
            and isinstance(value, str)
        }
        result["expected"] = expected

        required: set[str] = set()
        if article.metadata.titles:
            required.add("title")
        if article.metadata.authors:
            required.add("authors")
        if article.metadata.abstracts:
            required.add("abstract")
        if article.metadata.keywords:
            required.add("keywords")
        if any(isinstance(block, ParagraphBlock) for block in article.body):
            required.add("body")
        for level in range(1, 7):
            if any(
                isinstance(block, SectionBlock) and block.level == level
                for block in article.body
            ):
                required.add(f"heading{level}")
        if any(isinstance(block, ListItemBlock) and block.ordered for block in article.body):
            required.add("list_number")
        if any(isinstance(block, ListItemBlock) and not block.ordered for block in article.body):
            required.add("list_bullet")
        if any(isinstance(block, EquationBlock) for block in article.body):
            required.add("equation")
        if any(isinstance(block, TableBlock) and block.caption for block in article.body):
            required.add("table_caption")
        if any(isinstance(block, FigureBlock) and block.caption for block in article.body):
            required.add("figure_caption")
        if article.references:
            required.add("references")

        matched: list[str] = []
        missing: list[str] = []
        for role in sorted(required):
            style_name = expected.get(role)
            if not style_name:
                continue
            if style_counts.get(style_name, 0) > 0:
                matched.append(role)
            else:
                missing.append(role)
        result["matched_roles"] = matched
        result["missing_roles"] = missing
        if missing:
            result["warnings"].append(
                "VALIDATION: итоговый DOCX не использовал стили шаблона для ролей: "
                + ", ".join(missing)
                + "."
            )

        expected_table_style = evidence.get("docx_style_table")
        if isinstance(expected_table_style, str):
            table_styles = [
                table.style.name if table.style is not None else None
                for table in document.tables
            ]
            result["table_style"] = {
                "expected": expected_table_style,
                "used": table_styles,
                "matched": (
                    not document.tables
                    or all(name == expected_table_style for name in table_styles)
                ),
            }
            if document.tables and not result["table_style"]["matched"]:
                result["warnings"].append(
                    "VALIDATION: часть таблиц не получила табличный стиль DOCX-образца."
                )
        return result

    def _source_counts(self, source_path: Path) -> dict[str, int | None]:
        suffix = source_path.suffix.lower()
        if suffix == ".docx":
            return self._docx_counts(source_path)
        if suffix == ".tex":
            text = self._read_tex_tree(source_path)
            return {
                "paragraphs": None,
                "tables": len(re.findall(r"\\begin\s*\{table\*?\}", text)),
                "formulas": len(
                    re.findall(
                        r"\\begin\s*\{(?:equation|align|gather|multline)\*?\}|\$\$",
                        text,
                    )
                ),
                "drawings": len(re.findall(r"\\includegraphics", text)),
                "media_files": None,
            }
        if suffix == ".pdf":
            try:
                import pymupdf

                document = pymupdf.open(source_path)
                images = sum(len(page.get_images(full=True)) for page in document)
                pages = document.page_count
                document.close()
                return {
                    "paragraphs": None,
                    "tables": None,
                    "formulas": None,
                    "drawings": images,
                    "media_files": images,
                    "pages": pages,
                }
            except Exception:
                pass
        return {
            "paragraphs": None,
            "tables": None,
            "formulas": None,
            "drawings": None,
            "media_files": None,
        }

    @staticmethod
    def _docx_counts(source_path: Path) -> dict[str, int]:
        result = {
            "paragraphs": 0,
            "tables": 0,
            "formulas": 0,
            "omml": 0,
            "ole": 0,
            "drawings": 0,
            "media_files": 0,
        }
        try:
            with zipfile.ZipFile(source_path) as archive:
                root = etree.fromstring(archive.read("word/document.xml"))
                result["paragraphs"] = len(root.xpath(".//*[local-name()='p']"))
                tables = root.xpath(".//*[local-name()='tbl']")
                layout_tables = 0
                for table in tables:
                    rows = table.xpath("./*[local-name()='tr']")
                    cells = table.xpath("./*[local-name()='tr']/*[local-name()='tc']")
                    has_math = bool(table.xpath(".//*[local-name()='oMath']"))
                    has_drawing = bool(table.xpath(".//*[local-name()='drawing']"))
                    if has_drawing or (
                        has_math and len(rows) == 1 and len(cells) <= 2
                    ):
                        layout_tables += 1
                result["tables"] = len(tables) - layout_tables
                math_paragraphs = root.xpath(".//*[local-name()='oMathPara']")
                inline_math = root.xpath(
                    ".//*[local-name()='oMath' and not(ancestor::*[local-name()='oMathPara'])]"
                )
                result["omml"] = len(math_paragraphs) + len(inline_math)
                result["ole"] = len(root.xpath(".//*[local-name()='OLEObject']"))
                result["formulas"] = result["omml"] + result["ole"]
                result["drawings"] = len(root.xpath(".//*[local-name()='drawing']"))
                result["media_files"] = len(
                    [
                        name
                        for name in archive.namelist()
                        if name.startswith("word/media/") and not name.endswith("/")
                    ]
                )
        except Exception:
            pass
        return result

    @staticmethod
    def _article_counts(article: ArticleIR) -> dict[str, int]:
        return {
            "paragraphs": sum(isinstance(block, ParagraphBlock) for block in article.body),
            "sections": sum(isinstance(block, SectionBlock) for block in article.body),
            "list_items": sum(isinstance(block, ListItemBlock) for block in article.body),
            "equations": (
                sum(isinstance(block, EquationBlock) for block in article.body)
                + sum(
                    1
                    for block in article.body
                    if isinstance(block, (ParagraphBlock, ListItemBlock))
                    for run in block.runs
                    if run.math_latex is not None
                )
            ),
            "formula_images": sum(
                1
                for block in article.body
                if isinstance(block, (ParagraphBlock, ListItemBlock))
                for run in block.runs
                if run.asset_id and run.formula_image
            ),
            "figures": sum(isinstance(block, FigureBlock) for block in article.body),
            "tables": sum(isinstance(block, TableBlock) for block in article.body),
            "references": len(article.references),
            "citations": len(article.citations),
            "cross_references": len(article.cross_references),
            "notes": len(article.notes),
            "authors": len(article.metadata.authors),
            "titles": len(article.metadata.titles),
            "assets": len(article.assets),
        }

    def _tex_integrity(self, main_tex: Path, article: ArticleIR) -> dict[str, Any]:
        warnings: list[str] = []
        errors: list[str] = []
        texts: list[str] = []
        for path in (main_tex, main_tex.parent / "metadata.tex", main_tex.parent / "body.tex"):
            if path.exists():
                texts.append(path.read_text(encoding="utf-8", errors="replace"))
        text = "\n".join(texts)
        labels = re.findall(r"\\label\s*\{([^}]+)\}", text)
        refs = re.findall(r"\\(?:ref|eqref|autoref)\s*\{([^}]+)\}", text)
        citations = [
            key.strip()
            for group in re.findall(r"\\cite[a-zA-Z*]*(?:\[[^\]]*\])*\{([^}]+)\}", text)
            for key in group.split(",")
        ]
        bibitems = re.findall(r"\\bibitem(?:\[[^\]]*\])?\{([^}]+)\}", text)
        duplicate_labels = [
            value for value, count in Counter(labels).items() if count > 1
        ]
        unresolved_refs = sorted(set(refs) - set(labels))
        unresolved_citations = sorted(set(citations) - set(bibitems))
        if duplicate_labels:
            errors.append(
                f"VALIDATION: повторяются LaTeX labels: {duplicate_labels[:10]}."
            )
        if unresolved_refs:
            warnings.append(
                f"VALIDATION: не разрешено ссылок: {len(unresolved_refs)}."
            )
        if unresolved_citations:
            warnings.append(
                f"VALIDATION: не разрешено цитат: {len(unresolved_citations)}."
            )
        expected = len(refs) + len(citations)
        resolved = expected - len(unresolved_refs) - len(unresolved_citations)
        reference_score = 1.0 if expected == 0 else max(0.0, resolved / expected)
        return {
            "labels": len(labels),
            "references": len(refs),
            "citations": len(citations),
            "bibitems": len(bibitems),
            "duplicate_labels": duplicate_labels,
            "unresolved_references": unresolved_refs,
            "unresolved_citations": unresolved_citations,
            "reference_score": round(reference_score, 4),
            "warnings": warnings,
            "errors": errors,
        }

    @staticmethod
    def _asset_checks(project_dir: Path, article: ArticleIR) -> dict[str, Any]:
        warnings: list[str] = []
        errors: list[str] = []
        missing: list[str] = []
        mismatched_hashes: list[str] = []
        for asset in article.assets:
            path = Path(asset.path)
            if not path.is_absolute():
                path = project_dir / path
            if not path.exists():
                missing.append(asset.id)
                continue
            if asset.sha256:
                from paper_formatter.utils.files import sha256_file

                if sha256_file(path) != asset.sha256:
                    mismatched_hashes.append(asset.id)
        if missing:
            errors.append(f"VALIDATION: отсутствуют ресурсы: {missing[:10]}.")
        if mismatched_hashes:
            warnings.append(
                f"VALIDATION: изменились контрольные суммы ресурсов: {mismatched_hashes[:10]}."
            )
        total = len(article.assets)
        score = 1.0 if total == 0 else (total - len(missing)) / total
        return {
            "total": total,
            "missing": missing,
            "hash_mismatch": mismatched_hashes,
            "score": round(score, 4),
            "warnings": warnings,
            "errors": errors,
        }

    def _text_coverage(self, source_path: Path, article: ArticleIR) -> float | None:
        source_text = self._source_text(source_path)
        if not source_text:
            return None
        article_text_parts = [
            *(item.text for item in article.metadata.titles),
            *(item.text for item in article.metadata.subtitles),
            *(author.name for author in article.metadata.authors),
            *(item.name for item in article.metadata.affiliations),
            *(item.text for item in article.metadata.abstracts),
            *article.metadata.keywords,
        ]
        for block in article.body:
            if isinstance(block, SectionBlock):
                article_text_parts.append(block.title)
            elif isinstance(block, (ParagraphBlock, ListItemBlock)):
                article_text_parts.append("".join(run.text for run in block.runs))
            elif isinstance(block, TableBlock):
                article_text_parts.extend(cell for row in block.rows for cell in row)
        article_text_parts.extend(item.text for item in article.references)
        source_tokens = self._tokens(source_text)
        target_tokens = self._tokens("\n".join(article_text_parts))
        if not source_tokens:
            return None
        return round(len(source_tokens & target_tokens) / len(source_tokens), 4)

    @staticmethod
    def _source_text(source_path: Path) -> str:
        try:
            if source_path.suffix.lower() == ".docx":
                with zipfile.ZipFile(source_path) as archive:
                    root = etree.fromstring(archive.read("word/document.xml"))
                    return " ".join(root.xpath(".//*[local-name()='t']/text()"))
            if source_path.suffix.lower() == ".tex":
                text = ConversionValidator._read_tex_tree(source_path)
                text = re.sub(r"%.*$", "", text, flags=re.MULTILINE)
                text = re.sub(
                    r"(?m)^\s*\\(?:documentclass|usepackage|settopmatter|renewcommand|"
                    r"newcommand|providecommand|journal|date|pubyear|shorttitle|"
                    r"shortauthors|titlerunning|authorrunning|pagerange)\b[^\n]*$",
                    " ",
                    text,
                )
                text = re.sub(
                    r"\\includegraphics(?:\[[^\]]*\])?\s*\{[^}]*\}|"
                    r"\\(?:label|bibitem|bibliography|bibliographystyle)\s*\{[^}]*\}",
                    " ",
                    text,
                )
                text = re.sub(
                    r"\\begin\s*\{(?:equation|align|gather|multline|displaymath)\*?\}"
                    r".*?\\end\s*\{(?:equation|align|gather|multline|displaymath)\*?\}",
                    " ",
                    text,
                    flags=re.DOTALL,
                )
                text = re.sub(r"\$[^$]*\$", " ", text)
                text = re.sub(r"\\(?:begin|end)\s*\{[^}]+\}", " ", text)
                text = re.sub(r"\\[A-Za-z@]+\*?(?:\[[^\]]*\])?", " ", text)
                return re.sub(r"[{}$&_^~\\]", " ", text)
            if source_path.suffix.lower() == ".pdf":
                import pymupdf

                document = pymupdf.open(source_path)
                text = "\n".join(page.get_text() for page in document)
                document.close()
                return text
        except Exception:
            return ""
        return ""

    @staticmethod
    def _read_tex_tree(source_path: Path, seen: set[Path] | None = None) -> str:
        source_path = source_path.resolve()
        seen = seen or set()
        if source_path in seen or not source_path.exists():
            return ""
        seen.add(source_path)
        text = source_path.read_text(encoding="utf-8", errors="replace")
        text = re.sub(r"(?<!\\)%.*$", "", text, flags=re.MULTILINE)

        def replace(match: re.Match[str]) -> str:
            child = source_path.parent / match.group(1).strip()
            if not child.suffix:
                child = child.with_suffix(".tex")
            return ConversionValidator._read_tex_tree(child, seen)

        return re.sub(r"\\(?:input|include)\s*\{([^}]+)\}", replace, text)

    @staticmethod
    def _compile_log_warnings(log_path: Path | None) -> list[str]:
        if not log_path or not log_path.exists():
            return []
        text = log_path.read_text(encoding="utf-8", errors="replace")
        result: list[str] = []
        overfull = [
            float(value)
            for value in re.findall(
                r"Overfull \\[hv]box \(([0-9.]+)pt too wide\)",
                text,
            )
        ]
        if overfull and max(overfull) >= 5.0:
            result.append(
                "VALIDATION: обнаружено существенное переполнение "
                f"до {max(overfull):.2f} pt."
            )
        if re.search(r"Missing character:", text):
            result.append(
                "VALIDATION: журнал сообщает об отсутствующих глифах в PDF."
            )
        if re.search(r"undefined (?:references|citations)", text, re.IGNORECASE):
            result.append("VALIDATION: компилятор сообщил о неразрешённых ссылках.")
        if re.search(r"LaTeX Error:", text):
            result.append("VALIDATION: журнал содержит LaTeX Error.")
        return result

    @staticmethod
    def _pdf_info(pdf_path: Path | None) -> dict[str, Any]:
        if not pdf_path or not pdf_path.exists():
            return {"exists": False, "pages": None}
        try:
            import pymupdf

            document = pymupdf.open(pdf_path)
            pages = document.page_count
            sizes = [
                {
                    "width_pt": round(document[index].rect.width, 2),
                    "height_pt": round(document[index].rect.height, 2),
                }
                for index in range(min(3, pages))
            ]
            document.close()
            return {"exists": True, "pages": pages, "sample_page_sizes": sizes}
        except Exception:
            return {"exists": True, "pages": None}

    @staticmethod
    def _compare_count(
        warnings: list[str],
        label: str,
        source_count: int | None,
        target_count: int,
        tolerance: int = 0,
    ) -> None:
        if source_count is None or source_count <= 0:
            return
        if target_count + tolerance < source_count:
            warnings.append(
                f"VALIDATION: число {label} в ArticleIR меньше исходного: "
                f"{target_count} < {source_count}."
            )

    @staticmethod
    def _structure_score(
        source_counts: dict[str, int | None], article_counts: dict[str, int]
    ) -> float:
        ratios: list[float] = []
        mappings = [
            ("tables", "tables"),
            ("formulas", "equations"),
            ("drawings", "figures"),
        ]
        for source_key, target_key in mappings:
            source = source_counts.get(source_key)
            if source:
                target = article_counts[target_key]
                if source_key == "formulas":
                    target += article_counts["formula_images"]
                if source_key == "drawings":
                    target += article_counts["formula_images"]
                ratios.append(min(1.0, target / source))
        return round(sum(ratios) / len(ratios), 4) if ratios else 1.0

    @staticmethod
    def _pandoc_disagreements(
        package_analysis: dict[str, Any] | None,
        article_counts: dict[str, int],
    ) -> list[str]:
        if not package_analysis:
            return []
        audit = package_analysis.get("pandoc_audit") or {}
        if not audit.get("success"):
            return []
        counts = audit.get("counts") or {}
        warnings: list[str] = []
        mappings = [
            ("Table", "tables", "таблиц"),
            ("Image", "figures", "рисунков"),
            ("Math", "equations", "формул"),
        ]
        for pandoc_key, article_key, label in mappings:
            pandoc_count = int(counts.get(pandoc_key, 0))
            article_count = article_counts[article_key]
            if pandoc_count and abs(pandoc_count - article_count) > max(1, pandoc_count * 0.1):
                warnings.append(
                    "VALIDATION: собственный парсер и Pandoc расходятся по числу "
                    f"{label}: {article_count} против {pandoc_count}."
                )
        return warnings

    @staticmethod
    def _tokens(text: str) -> set[str]:
        return set(re.findall(r"[0-9A-Za-zА-Яа-яЁё]{2,}", text.lower()))
