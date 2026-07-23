from __future__ import annotations

from collections import Counter
from io import BytesIO
import re
from typing import Any, Iterable

from .schema import BLOCK_CATALOG, get_document_blocks, normalize_template_rules


class DocumentTemplateEngineError(ValueError):
    pass


_UDC_RE = re.compile(r"^\s*(?:удк|udc)\b", re.I)
_SUPERVISOR_RE = re.compile(
    r"^\s*(?:научн\w*\s+руководител\w*|scientific\s+supervisor)\b",
    re.I,
)
_REFERENCES_RE = re.compile(
    r"^\s*(?:список\s+(?:использованн\w+\s+)?литератур\w*|литература|references)\s*[:.]?\s*$",
    re.I,
)
_ABSTRACT_RE = re.compile(r"^\s*(?:аннотация|abstract)\s*[:.]", re.I)
_KEYWORDS_RE = re.compile(r"^\s*(?:ключевые\s+слова|keywords)\s*[:.]", re.I)
_INSTITUTION_RE = re.compile(
    r"\b(?:университет|институт|академи|колледж|кафедр|лаборатор|"
    r"university|institute|academy|organization|organisation)\b",
    re.I,
)
_AUTHOR_TOKEN_RE = re.compile(
    r"(?:[A-ZА-ЯЁ]\s*\.\s*){1,3}[A-ZА-ЯЁ][A-ZА-ЯЁ-]{1,}|"
    r"[A-ZА-ЯЁ][A-ZА-ЯЁ-]{2,}\s+(?:[A-ZА-ЯЁ]\s*\.\s*){1,3}"
)
_WORD_RE = re.compile(r"[0-9A-Za-zА-Яа-яЁё][0-9A-Za-zА-Яа-яЁё'’-]*")


def _normalize_text(value: Any) -> str:
    text = str(value or "").casefold().replace("ё", "е")
    text = re.sub(r"[^0-9a-zа-я]+", " ", text)
    return " ".join(text.split())


def _as_float(value: Any) -> float | None:
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def _different(left: Any, right: Any, tolerance: float) -> bool:
    left_value = _as_float(left)
    right_value = _as_float(right)
    return (
        left_value is not None
        and right_value is not None
        and abs(left_value - right_value) > tolerance
    )


def _dominant(values: Iterable[Any]) -> Any:
    cleaned = [value for value in values if value not in (None, "")]
    return Counter(cleaned).most_common(1)[0][0] if cleaned else None


def _metadata_values(metadata: dict[str, Any] | None) -> dict[str, str]:
    metadata = metadata or {}
    return {
        "udc": str(metadata.get("udc") or "").strip(),
        "title": str(metadata.get("title") or "").strip(),
        "authors": str(
            metadata.get("authors")
            or metadata.get("document_authors")
            or ""
        ).strip(),
        "supervisor": str(metadata.get("supervisor") or "").strip(),
        "institution": str(
            metadata.get("institution")
            or metadata.get("organizations")
            or ""
        ).strip(),
        "city_country": str(metadata.get("city_country") or "").strip(),
        "abstract": str(metadata.get("abstract") or "").strip(),
        "keywords": str(metadata.get("keywords") or "").strip(),
    }


def _alignment_name(value: Any) -> str:
    if value is None:
        return ""
    try:
        numeric = int(value)
    except (TypeError, ValueError):
        numeric = None
    return {
        0: "left",
        1: "center",
        2: "right",
        3: "justify",
        4: "distribute",
    }.get(numeric, str(value).split()[0].casefold())


def _document_defaults(document) -> dict[str, Any]:
    from docx.oxml.ns import qn

    result = {
        "font_size_pt": None,
        "first_line_indent_cm": None,
        "line_spacing": 1.0,
    }
    doc_defaults = document.styles.element.find(qn("w:docDefaults"))
    if doc_defaults is None:
        return result
    run_default = doc_defaults.find(qn("w:rPrDefault"))
    run_properties = run_default.find(qn("w:rPr")) if run_default is not None else None
    size = run_properties.find(qn("w:sz")) if run_properties is not None else None
    if size is not None:
        value = _as_float(size.get(qn("w:val")))
        if value is not None:
            result["font_size_pt"] = value / 2
    paragraph_default = doc_defaults.find(qn("w:pPrDefault"))
    paragraph_properties = (
        paragraph_default.find(qn("w:pPr"))
        if paragraph_default is not None
        else None
    )
    indent = (
        paragraph_properties.find(qn("w:ind"))
        if paragraph_properties is not None
        else None
    )
    if indent is not None:
        first_line = _as_float(indent.get(qn("w:firstLine")))
        if first_line is not None:
            result["first_line_indent_cm"] = round(first_line / 567.0, 2)
    return result


def _resolve_paragraph_value(paragraph, attribute: str, normal_style, defaults):
    direct = getattr(paragraph.paragraph_format, attribute)
    if direct is not None:
        return direct
    style = getattr(paragraph, "style", None)
    if style is not None:
        inherited = getattr(style.paragraph_format, attribute)
        if inherited is not None:
            return inherited
    normal_value = getattr(normal_style.paragraph_format, attribute)
    if normal_value is not None:
        return normal_value
    if attribute == "line_spacing":
        return defaults.get("line_spacing")
    if attribute == "first_line_indent":
        value = defaults.get("first_line_indent_cm")
        if value is not None:
            from docx.shared import Cm

            return Cm(value)
    return None


def _resolve_run_font(run, paragraph, normal_style, defaults) -> tuple[str, float | None]:
    name = run.font.name
    size = run.font.size
    style = getattr(paragraph, "style", None)
    if not name and style is not None:
        name = style.font.name
    if size is None and style is not None:
        size = style.font.size
    if not name:
        name = normal_style.font.name
    if size is None:
        size = normal_style.font.size
    resolved_size = (
        round(size.pt, 1)
        if size is not None
        else defaults.get("font_size_pt")
    )
    return name or "", resolved_size


def _looks_like_authors(text: str) -> bool:
    if len(text) > 300:
        return False
    return len(_AUTHOR_TOKEN_RE.findall(text.upper())) >= 1


def _looks_like_city_country(text: str) -> bool:
    if not (2 <= len(text) <= 100) or "," not in text:
        return False
    if re.search(r"[.!?;:]", text.rstrip(",").strip()):
        return False
    return len(_WORD_RE.findall(text)) <= 8


def _is_probable_title(text: str) -> bool:
    letters = [char for char in text if char.isalpha()]
    if not (10 <= len(text) <= 500) or not letters:
        return False
    uppercase_ratio = sum(char.isupper() for char in letters) / len(letters)
    return uppercase_ratio >= 0.8


def _assign_roles(document, metadata: dict[str, Any] | None = None) -> dict[int, str]:
    paragraphs = document.paragraphs
    nonempty = [index for index, paragraph in enumerate(paragraphs) if paragraph.text.strip()]
    roles: dict[int, str] = {}
    values = _metadata_values(metadata)

    for index in nonempty:
        text = paragraphs[index].text.strip()
        normalized = _normalize_text(text)
        if _UDC_RE.match(text):
            roles[index] = "udc"
        elif _SUPERVISOR_RE.match(text):
            roles[index] = "supervisor"
        elif _REFERENCES_RE.match(text):
            roles[index] = "references_heading"
        elif _ABSTRACT_RE.match(text):
            roles[index] = "abstract"
        elif _KEYWORDS_RE.match(text):
            roles[index] = "keywords"
        else:
            for role in (
                "title",
                "authors",
                "supervisor",
                "institution",
                "city_country",
                "abstract",
                "keywords",
            ):
                candidate = _normalize_text(values.get(role))
                if candidate and (candidate == normalized or (len(candidate) > 12 and candidate in normalized)):
                    roles[index] = role
                    break

    title_index = next((index for index, role in roles.items() if role == "title"), None)
    if title_index is None:
        udc_index = next((index for index, role in roles.items() if role == "udc"), None)
        candidates = [
            index
            for index in nonempty
            if (udc_index is None or index > udc_index)
            and index not in roles
            and _is_probable_title(paragraphs[index].text.strip())
        ]
        if candidates:
            title_index = candidates[0]
            roles[title_index] = "title"

    if title_index is not None:
        for index in nonempty:
            if index <= title_index or index in roles:
                continue
            text = paragraphs[index].text.strip()
            if _looks_like_authors(text):
                roles[index] = "authors"
            break

    references_heading = next(
        (index for index, role in roles.items() if role == "references_heading"),
        None,
    )
    front_limit = references_heading if references_heading is not None else len(paragraphs)
    for index in nonempty:
        if index in roles or index >= front_limit:
            continue
        text = paragraphs[index].text.strip()
        if _INSTITUTION_RE.search(text) and len(text) <= 250:
            roles[index] = "institution"
        elif _looks_like_city_country(text):
            roles[index] = "city_country"

    if references_heading is not None:
        for index in nonempty:
            if index > references_heading:
                roles[index] = "references"

    front_roles = {
        "udc",
        "title",
        "authors",
        "supervisor",
        "institution",
        "city_country",
        "abstract",
        "keywords",
    }
    front_indices = [index for index, role in roles.items() if role in front_roles]
    body_start_after = max(front_indices) if front_indices else -1
    body_end = references_heading if references_heading is not None else len(paragraphs)
    for index in nonempty:
        if body_start_after < index < body_end and index not in roles:
            roles[index] = "body"

    if not any(role == "body" for role in roles.values()):
        for index in nonempty:
            if index not in roles and (references_heading is None or index < references_heading):
                roles[index] = "body"
    return roles


def _issue(
    code: str,
    title: str,
    message: str,
    *,
    severity: str = "warning",
    suggestion: str = "",
    fixable: bool = False,
    location: str = "Оформление по шаблону",
) -> dict[str, Any]:
    return {
        "code": code,
        "title": title,
        "severity": severity,
        "message": message,
        "location": location,
        "context": "",
        "context_before": "",
        "context_highlight": "",
        "context_after": "",
        "suggestion": suggestion,
        "fixable": fixable,
    }


def _load_document(docx_bytes: bytes):
    try:
        from docx import Document
    except ImportError as exc:
        raise DocumentTemplateEngineError("Для обработки DOCX требуется python-docx.") from exc
    try:
        return Document(BytesIO(docx_bytes))
    except Exception as exc:
        raise DocumentTemplateEngineError("Не удалось открыть DOCX.") from exc


def _configured_block_map(rules: dict[str, Any]) -> dict[str, dict[str, Any]]:
    return {item["role"]: item for item in get_document_blocks(rules)}


def _block_entries(document, roles: dict[int, str], rules: dict[str, Any]) -> list[dict[str, Any]]:
    entries = []
    by_role = _configured_block_map(rules)
    for role, block in by_role.items():
        paragraph_numbers = [
            index + 1
            for index, assigned_role in roles.items()
            if assigned_role == role
            or (role == "references" and assigned_role == "references_heading")
        ]
        entries.append(
            {
                "role": role,
                "label": block.get("label") or BLOCK_CATALOG[role]["label"],
                "required": bool(block.get("required")),
                "found": bool(paragraph_numbers),
                "paragraph_numbers": paragraph_numbers[:20],
                "status": (
                    "found"
                    if paragraph_numbers
                    else ("missing" if block.get("required") else "optional_missing")
                ),
            }
        )
    return entries


def _document_metrics(document, roles: dict[int, str]) -> dict[str, Any]:
    normal_style = document.styles["Normal"]
    defaults = _document_defaults(document)
    fonts = []
    sizes = []
    line_spacings = []
    first_line_indents = []
    alignments = []
    body_paragraphs = [
        document.paragraphs[index]
        for index, role in roles.items()
        if role == "body"
    ]
    for paragraph in body_paragraphs:
        for run in paragraph.runs:
            if not run.text.strip():
                continue
            font_name, font_size = _resolve_run_font(
                run,
                paragraph,
                normal_style,
                defaults,
            )
            fonts.append(font_name)
            sizes.append(font_size)
        spacing = _resolve_paragraph_value(
            paragraph,
            "line_spacing",
            normal_style,
            defaults,
        )
        if isinstance(spacing, (int, float)):
            line_spacings.append(round(float(spacing), 2))
        indent = _resolve_paragraph_value(
            paragraph,
            "first_line_indent",
            normal_style,
            defaults,
        )
        if indent is not None:
            first_line_indents.append(round(indent.cm, 2))
        alignment = paragraph.alignment
        if alignment is None and paragraph.style is not None:
            alignment = paragraph.style.paragraph_format.alignment
        if alignment is None:
            alignment = normal_style.paragraph_format.alignment
        alignments.append(_alignment_name(alignment))

    margins = {}
    if document.sections:
        section = document.sections[0]
        margins = {
            "top": round(section.top_margin.cm, 2),
            "right": round(section.right_margin.cm, 2),
            "bottom": round(section.bottom_margin.cm, 2),
            "left": round(section.left_margin.cm, 2),
        }
    return {
        "margins_cm": margins,
        "font_family": _dominant(fonts),
        "font_size_pt": _dominant(sizes),
        "line_spacing": _dominant(line_spacings),
        "first_line_indent_cm": _dominant(first_line_indents),
        "alignment": _dominant(alignments),
        "body_paragraphs": len(body_paragraphs),
        "word_count": len(
            _WORD_RE.findall("\n".join(paragraph.text for paragraph in document.paragraphs))
        ),
    }


def check_docx_against_template(
    docx_bytes: bytes,
    rules: dict[str, Any],
    *,
    metadata: dict[str, Any] | None = None,
) -> dict[str, Any]:
    normalized_rules = normalize_template_rules(rules)
    document = _load_document(docx_bytes)
    roles = _assign_roles(document, metadata)
    metrics = _document_metrics(document, roles)
    block_map = _configured_block_map(normalized_rules)
    block_entries = _block_entries(document, roles, normalized_rules)
    metadata_values = _metadata_values(metadata)
    issues: list[dict[str, Any]] = []

    for entry in block_entries:
        if entry["required"] and not entry["found"]:
            role = entry["role"]
            can_insert = bool(metadata_values.get(role))
            issues.append(
                _issue(
                    f"template_missing_block_{role}",
                    f"Не найден блок «{entry['label']}»",
                    "Шаблон требует этот элемент документа, но он не распознан как отдельный блок.",
                    suggestion=(
                        "Система может вставить значение из заполненных полей заявки."
                        if can_insert
                        else "Укажите значение в документе или в полях заявки; научный текст система не придумывает."
                    ),
                    fixable=can_insert,
                )
            )

    normalized_paragraphs = {
        _normalize_text(paragraph.text).strip(" .:")
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    }
    required_sections = (
        (normalized_rules.get("structure") or {}).get("required_sections") or []
    )
    missing_sections = [
        str(section)
        for section in required_sections
        if _normalize_text(section).strip(" .:") not in normalized_paragraphs
    ]
    for section in missing_sections:
        issues.append(
            _issue(
                "template_missing_section",
                f"Нет раздела «{section}»",
                "Обязательный заголовок раздела из выбранного шаблона не найден.",
                suggestion="Добавьте раздел и заполните его собственным научным содержанием.",
            )
        )
    metrics["required_sections"] = list(required_sections)
    metrics["missing_sections"] = missing_sections

    limits = normalized_rules.get("limits") or {}
    minimum_words = _as_float(limits.get("min_words"))
    maximum_words = _as_float(limits.get("max_words"))
    word_count = metrics["word_count"]
    if minimum_words is not None and word_count and word_count < minimum_words:
        issues.append(
            _issue(
                "template_min_words",
                "Объём меньше требования",
                f"Найдено {word_count} слов, требуется не менее {int(minimum_words)}.",
                suggestion="Дополните материал самостоятельно; система не создаёт научное содержание.",
            )
        )
    if maximum_words is not None and word_count > maximum_words:
        issues.append(
            _issue(
                "template_max_words",
                "Объём больше требования",
                f"Найдено {word_count} слов, допускается не более {int(maximum_words)}.",
                suggestion="Сократите материал вручную, сохранив научный смысл.",
            )
        )

    page_rules = normalized_rules.get("page") or {}
    expected_margins = page_rules.get("margins_cm") or {}
    for key, label in (
        ("top", "верхнее"),
        ("right", "правое"),
        ("bottom", "нижнее"),
        ("left", "левое"),
    ):
        actual = metrics["margins_cm"].get(key)
        expected = expected_margins.get(key)
        if _different(actual, expected, 0.12):
            issues.append(
                _issue(
                    f"template_margin_{key}",
                    f"Неверное {label} поле",
                    f"В документе {actual} см, шаблон требует {expected} см.",
                    suggestion="Конструктор выставит поле автоматически.",
                    fixable=True,
                )
            )

    body_rules = normalized_rules.get("body") or {}
    comparisons = (
        ("font_family", "Основной шрифт отличается", 0, "шрифт"),
        ("font_size_pt", "Размер основного шрифта отличается", 0.2, "размер"),
        ("line_spacing", "Межстрочный интервал отличается", 0.05, "интервал"),
        ("first_line_indent_cm", "Абзацный отступ отличается", 0.08, "отступ"),
    )
    for key, title, tolerance, noun in comparisons:
        expected = body_rules.get(key)
        actual = metrics.get(key)
        if expected in (None, "") or actual in (None, ""):
            continue
        differs = (
            str(actual).casefold() != str(expected).casefold()
            if key == "font_family"
            else _different(actual, expected, tolerance)
        )
        if differs:
            issues.append(
                _issue(
                    f"template_{key}",
                    title,
                    f"Распознано: {actual}; требуется: {expected}.",
                    suggestion=f"Конструктор исправит {noun} только в основном тексте.",
                    fixable=True,
                )
            )

    expected_alignment = str(body_rules.get("alignment") or "").casefold()
    if (
        expected_alignment
        and metrics.get("alignment")
        and expected_alignment != str(metrics["alignment"]).casefold()
    ):
        issues.append(
            _issue(
                "template_alignment",
                "Выравнивание основного текста отличается",
                f"Распознано: {metrics['alignment']}; требуется: {expected_alignment}.",
                suggestion="Конструктор исправит выравнивание только в основном тексте.",
                fixable=True,
            )
        )

    title_indices = [index for index, role in roles.items() if role == "title"]
    title_block = block_map.get("title") or {}
    constraints = title_block.get("constraints") or {}
    if title_indices and constraints.get("uppercase"):
        title_text = document.paragraphs[title_indices[0]].text.strip()
        if title_text != title_text.upper():
            issues.append(
                _issue(
                    "template_title_uppercase",
                    "Название должно быть прописными буквами",
                    "Название распознано, но содержит строчные буквы.",
                    suggestion="Конструктор может изменить регистр названия без изменения смысла.",
                    fixable=True,
                )
            )

    normal_style = document.styles["Normal"]
    defaults = _document_defaults(document)
    for role, block in block_map.items():
        if role == "body":
            continue
        expected_style = block.get("style") or {}
        role_indices = [index for index, assigned in roles.items() if assigned == role]
        if not role_indices:
            continue
        expected_alignment = str(expected_style.get("alignment") or "").casefold()
        expected_indent = expected_style.get("first_line_indent_cm")
        actual_alignments = []
        actual_indents = []
        for index in role_indices:
            paragraph = document.paragraphs[index]
            alignment = paragraph.alignment
            if alignment is None and paragraph.style is not None:
                alignment = paragraph.style.paragraph_format.alignment
            if alignment is None:
                alignment = normal_style.paragraph_format.alignment
            actual_alignments.append(_alignment_name(alignment or 0))
            indent = _resolve_paragraph_value(
                paragraph,
                "first_line_indent",
                normal_style,
                defaults,
            )
            actual_indents.append(round(indent.cm, 2) if indent is not None else 0)
        actual_alignment = _dominant(actual_alignments)
        actual_indent = _dominant(actual_indents)
        label = block.get("label") or BLOCK_CATALOG[role]["label"]
        if expected_alignment and actual_alignment != expected_alignment:
            issues.append(
                _issue(
                    f"template_block_{role}_alignment",
                    f"Неверно выровнен блок «{label}»",
                    f"Распознано: {actual_alignment}; требуется: {expected_alignment}.",
                    suggestion="Конструктор исправит только этот блок.",
                    fixable=True,
                )
            )
        if expected_indent is not None and _different(actual_indent, expected_indent, 0.08):
            issues.append(
                _issue(
                    f"template_block_{role}_indent",
                    f"Неверный отступ блока «{label}»",
                    f"Распознано: {actual_indent} см; требуется: {expected_indent} см.",
                    suggestion="Конструктор исправит только этот блок.",
                    fixable=True,
                )
            )

    return {
        "schema_version": "2.0",
        "rules": normalized_rules,
        "issues": issues,
        "metrics": metrics,
        "blocks": block_entries,
        "role_assignments": [
            {
                "paragraph_number": index + 1,
                "role": role,
                "text": document.paragraphs[index].text[:160],
            }
            for index, role in sorted(roles.items())
        ],
        "can_build": True,
        "content_policy": normalized_rules.get("document", {}).get("content_policy", {}),
    }


def build_docx_plan(
    docx_bytes: bytes,
    rules: dict[str, Any],
    *,
    metadata: dict[str, Any] | None = None,
) -> dict[str, Any]:
    report = check_docx_against_template(docx_bytes, rules, metadata=metadata)
    missing = [entry for entry in report["blocks"] if entry["status"] == "missing"]
    operations = [
        "Сохранить исходный научный текст, таблицы, рисунки и ссылки.",
        "Применить параметры страницы из выбранного шаблона.",
        "Оформить титульный блок отдельно от основного текста.",
        "Применить абзацный отступ и выравнивание только к основному тексту.",
        "Сохранить список литературы отдельным блоком без шаблонных заглушек.",
    ]
    return {
        **report,
        "missing_blocks": missing,
        "operations": operations,
        "engine": {
            "name": "document_template_engine",
            "version": "2.0",
            "framework_agnostic": True,
        },
    }


def _set_run_font(run, *, font_family: str, font_size: float | None):
    from docx.oxml.ns import qn
    from docx.shared import Pt

    if font_family:
        run.font.name = font_family
        fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
        for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            fonts.set(qn(key), font_family)
    if font_size is not None:
        run.font.size = Pt(font_size)


def _apply_paragraph_style(paragraph, style: dict[str, Any], *, body_rules: dict[str, Any]):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm

    merged = {
        key: value
        for key, value in body_rules.items()
        if key in {
            "font_family",
            "font_size_pt",
            "line_spacing",
            "first_line_indent_cm",
            "alignment",
        }
    }
    merged.update({key: value for key, value in (style or {}).items() if value not in (None, "")})
    font_family = str(merged.get("font_family") or "").strip()
    font_size = _as_float(merged.get("font_size_pt"))
    line_spacing = _as_float(merged.get("line_spacing"))
    first_line_indent = _as_float(merged.get("first_line_indent_cm"))
    alignment = str(merged.get("alignment") or "").casefold()
    alignment_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "по левому краю": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "центр": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "по правому краю": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "justified": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "по ширине": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    if line_spacing is not None:
        paragraph.paragraph_format.line_spacing = line_spacing
    if first_line_indent is not None:
        paragraph.paragraph_format.first_line_indent = Cm(first_line_indent)
    if alignment in alignment_map:
        paragraph.alignment = alignment_map[alignment]
    for run in paragraph.runs:
        _set_run_font(run, font_family=font_family, font_size=font_size)
        if merged.get("bold") is not None:
            run.bold = bool(merged["bold"])
        if merged.get("italic") is not None:
            run.italic = bool(merged["italic"])


def _insert_paragraph_before(anchor, text: str):
    from docx.text.paragraph import Paragraph
    from docx.oxml import OxmlElement

    element = OxmlElement("w:p")
    anchor._p.addprevious(element)
    paragraph = Paragraph(element, anchor._parent)
    paragraph.add_run(text)
    return paragraph


def _insert_supplied_metadata(document, rules: dict[str, Any], metadata: dict[str, Any] | None):
    values = _metadata_values(metadata)
    roles = _assign_roles(document, metadata)
    existing_roles = set(roles.values())
    blocks = get_document_blocks(rules)
    missing = [
        block
        for block in blocks
        if block.get("required")
        and block["role"] not in existing_roles
        and values.get(block["role"])
        and block["role"] not in {"body", "references"}
    ]
    if not missing:
        return []

    body_indices = [index for index, role in roles.items() if role == "body"]
    anchor = document.paragraphs[min(body_indices)] if body_indices else None
    inserted = []
    for block in missing:
        if anchor is None:
            paragraph = document.add_paragraph(values[block["role"]])
        else:
            paragraph = _insert_paragraph_before(anchor, values[block["role"]])
        inserted.append((block["role"], paragraph))
    return inserted


def _reorder_front_matter(document, roles: dict[int, str], rules: dict[str, Any]) -> bool:
    front_roles = {
        "udc",
        "title",
        "authors",
        "supervisor",
        "institution",
        "city_country",
        "abstract",
        "keywords",
    }
    anchor_indices = [
        index
        for index, role in roles.items()
        if role in {"body", "references_heading", "references"}
    ]
    if not anchor_indices:
        return False
    anchor_index = min(anchor_indices)
    front_indices = [
        index
        for index, paragraph in enumerate(document.paragraphs[:anchor_index])
        if paragraph.text.strip()
    ]
    if not front_indices or any(roles.get(index) not in front_roles for index in front_indices):
        return False
    order = list((rules.get("document") or {}).get("block_order") or [])
    rank = {role: index for index, role in enumerate(order)}
    desired = sorted(
        front_indices,
        key=lambda index: (rank.get(roles[index], len(rank)), index),
    )
    if desired == front_indices:
        return False
    paragraphs = document.paragraphs
    anchor_element = paragraphs[anchor_index]._p
    desired_elements = [paragraphs[index]._p for index in desired]
    for element in desired_elements:
        anchor_element.addprevious(element)
    return True


def _apply_text_constraints(paragraph, constraints: dict[str, Any]) -> bool:
    original = paragraph.text
    updated = original
    if constraints.get("uppercase"):
        updated = updated.upper()
    if constraints.get("terminal_period_allowed") is False:
        updated = updated.rstrip()
        if updated.endswith("."):
            updated = updated[:-1].rstrip()
    if updated == original or not paragraph.runs:
        return False
    paragraph.runs[0].text = updated
    for run in paragraph.runs[1:]:
        run.text = ""
    return True


def build_docx_from_template(
    docx_bytes: bytes,
    rules: dict[str, Any],
    *,
    metadata: dict[str, Any] | None = None,
) -> tuple[bytes, list[str], dict[str, Any]]:
    normalized_rules = normalize_template_rules(rules)
    document = _load_document(docx_bytes)
    changes: list[str] = []

    inserted = _insert_supplied_metadata(document, normalized_rules, metadata)
    if inserted:
        changes.append(
            "добавлены заполненные поля: "
            + ", ".join(BLOCK_CATALOG[role]["label"] for role, _paragraph in inserted)
        )

    try:
        from docx.enum.section import WD_ORIENT
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt
    except ImportError as exc:
        raise DocumentTemplateEngineError("Для обработки DOCX требуется python-docx.") from exc

    page_rules = normalized_rules.get("page") or {}
    margins = page_rules.get("margins_cm") or {}
    for section in document.sections:
        for key, attribute, label in (
            ("top", "top_margin", "верхнее поле"),
            ("right", "right_margin", "правое поле"),
            ("bottom", "bottom_margin", "нижнее поле"),
            ("left", "left_margin", "левое поле"),
        ):
            value = _as_float(margins.get(key))
            if value is not None:
                setattr(section, attribute, Cm(value))
                changes.append(f"{label}: {value:g} см")
        orientation = str(page_rules.get("orientation") or "").casefold()
        if orientation in {"landscape", "альбомная", "album"}:
            if section.page_height > section.page_width:
                section.page_width, section.page_height = section.page_height, section.page_width
            section.orientation = WD_ORIENT.LANDSCAPE
        elif orientation in {"portrait", "книжная"}:
            if section.page_width > section.page_height:
                section.page_width, section.page_height = section.page_height, section.page_width
            section.orientation = WD_ORIENT.PORTRAIT

    body_rules = normalized_rules.get("body") or {}
    font_family = str(body_rules.get("font_family") or "").strip()
    font_size = _as_float(body_rules.get("font_size_pt"))
    normal_style = document.styles["Normal"]
    if font_family:
        normal_style.font.name = font_family
        fonts = normal_style._element.get_or_add_rPr().get_or_add_rFonts()
        for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            fonts.set(qn(key), font_family)
        changes.append(f"основной шрифт: {font_family}")
    if font_size is not None:
        normal_style.font.size = Pt(font_size)
        changes.append(f"размер основного шрифта: {font_size:g} пт")

    roles = _assign_roles(document, metadata)
    if _reorder_front_matter(document, roles, normalized_rules):
        changes.append("блоки титульной части расставлены в порядке шаблона")
        roles = _assign_roles(document, metadata)
    block_map = _configured_block_map(normalized_rules)
    for index, paragraph in enumerate(document.paragraphs):
        if not paragraph.text.strip():
            continue
        role = roles.get(index, "")
        if role == "references_heading":
            role_style = {
                "alignment": "center",
                "first_line_indent_cm": 0,
                "bold": True,
                "font_family": font_family,
                "font_size_pt": font_size,
                "line_spacing": body_rules.get("line_spacing"),
            }
            _apply_paragraph_style(paragraph, role_style, body_rules={})
            continue
        block = block_map.get(role)
        if block:
            if role == "title" and _apply_text_constraints(
                paragraph,
                block.get("constraints") or {},
            ):
                changes.append("регистр и пунктуация названия приведены к шаблону")
            role_body_rules = body_rules if role in {"body", "abstract", "keywords"} else {
                "font_family": font_family,
                "font_size_pt": font_size,
                "line_spacing": body_rules.get("line_spacing"),
            }
            _apply_paragraph_style(
                paragraph,
                block.get("style") or {},
                body_rules=role_body_rules,
            )
        else:
            # Unknown elements keep their local alignment/indent. Typography is
            # normalized, because this does not change structural meaning.
            for run in paragraph.runs:
                _set_run_font(run, font_family=font_family, font_size=font_size)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        _set_run_font(run, font_family=font_family, font_size=font_size)

    if body_rules.get("line_spacing") not in (None, ""):
        changes.append(f"интервал основного текста: {body_rules['line_spacing']}")
    if body_rules.get("first_line_indent_cm") not in (None, ""):
        changes.append(
            f"абзацный отступ основного текста: {body_rules['first_line_indent_cm']} см"
        )
    changes.append("титульный блок оформлен отдельно от основного текста")

    output = BytesIO()
    document.save(output)
    built_bytes = output.getvalue()
    plan = build_docx_plan(built_bytes, normalized_rules, metadata=metadata)
    return built_bytes, list(dict.fromkeys(changes)), plan
